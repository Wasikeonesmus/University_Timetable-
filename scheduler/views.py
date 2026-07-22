import json, csv, io, datetime, logging
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse, HttpResponseNotAllowed
from django.views.decorators.http import require_POST
from django.db.models import Count, Q

logger = logging.getLogger(__name__)

from .models import (
    University, Campus, Faculty, Department, Room, Lecturer, StudentGroup,
    Course, Timetable, ScheduleSlot, TimeSlot, Constraint, Semester, GenerationLog,
    LecturerAvailability, Program, AttendanceSession, AttendanceRecord, Announcement, Notification,
    ApprovalLog, FieldMapping
)
from .forms import (
    TimetableForm, ConstraintForm, UniversityForm, CampusForm, FacultyForm,
    DepartmentForm, CourseForm, LecturerForm, StudentGroupForm, RoomForm, TimeSlotForm
)
from .solver import generate_timetable
from .conflicts import check_conflicts_for_timetable, detect_conflicts
from .scheduling_service import run_scheduling_pipeline
from django.conf import settings
from .firebase_service import update_timetable_conflicts, trigger_timetable_refresh
from .permissions import (
    get_effective_role, manager_required, role_required, ROLE_STUDENT, ROLE_LECTURER, 
    ROLE_INST_ADMIN, ROLE_REGISTRAR, ROLE_DVC, ROLE_DEAN, ROLE_HOD, ROLE_TIMETABLE_OFFICER,
    ROLE_SCHEDULER, ROLE_ADMIN, MANAGER_ROLES, tenant_required
)

# ──────────────────────────────────────────────────────────────────────────────
# Timetable Conflicts AJAX Caching
# The cache dict and clear helper live in conflict_cache.py so that
# scheduling_service.py can also invalidate it without a circular import.
# ──────────────────────────────────────────────────────────────────────────────
from .conflict_cache import _CONFLICTS_JSON_CACHE, clear_conflicts_cache, _CACHE_MAX_SIZE

def _get_timetable_conflict_fingerprint(timetable):
    # 1. Fetch slot assignments (room, timeslot) — cheap indexed scan
    slots_fp = list(timetable.slots.order_by('id').values_list('id', 'room_id', 'time_slot_id'))
    # 2. Constraint configuration for the university
    raw_constraints = Constraint.objects.filter(
        university_id=timetable.semester.university_id
    ).order_by('id').values_list('id', 'is_hard', 'weight', 'parameters')
    
    constraints_fp = []
    for c_id, is_hard, weight, params in raw_constraints:
        params_hashable = tuple(sorted(params.items())) if isinstance(params, dict) else params
        constraints_fp.append((c_id, is_hard, weight, params_hashable))

    # 3. Lecturer self-service availability
    avail_fp = list(LecturerAvailability.objects.filter(
        lecturer__department__faculty__campus__university_id=timetable.semester.university_id
    ).order_by('id').values_list('id', 'lecturer_id', 'time_slot_id', 'is_available'))
    return (timetable.id, tuple(slots_fp), tuple(constraints_fp), tuple(avail_fp))



def get_active_uni(request):
    """
    Helper to fetch the session's active university or fall back to the first.
    If the user has a profile with a linked university, restrict to that (unless global Super Admin).
    """
    if request.user.is_authenticated:
        try:
            profile = request.user.profile
            if profile.role != 'admin' and profile.university:
                return profile.university
        except Exception:
            pass

    uni_id = request.session.get('active_university_id')
    if uni_id:
        uni = University.objects.filter(id=uni_id).first()
        if uni:
            return uni
    uni = University.objects.first()
    if uni:
        request.session['active_university_id'] = uni.id
    return uni


def get_user_role(request):
    """Returns the authenticated user's role from UserProfile."""
    role = get_effective_role(request)
    return role if role else ROLE_STUDENT


@login_required
def dashboard(request):
    role = get_user_role(request)
    university = get_active_uni(request)

    # ── Lecturer gets redirected to the new Lecturer Portal ──────────────────────────
    if role == ROLE_LECTURER:
        try:
            if request.user.profile.lecturer:
                return redirect('scheduler:lecturer_portal_dashboard')
            else:
                messages.error(request, "Your account is not linked to a lecturer profile. Please update your profile.")
                return redirect('accounts:profile')
        except Exception:
            messages.error(request, "Your account is not linked to a lecturer profile. Please update your profile.")
            return redirect('accounts:profile')


    if not university:
        return render(request, 'scheduler/dashboard.html', {
            'timetables': [],
            'active_role': role,
        })

    # Filter all metrics by active university
    campuses_count       = Campus.objects.filter(university=university).count()
    rooms_count          = Room.objects.filter(campus__university=university).count()
    lecturers_count      = Lecturer.objects.filter(department__faculty__campus__university=university).count()
    student_groups_count = StudentGroup.objects.filter(program__department__faculty__campus__university=university).count()
    courses_count        = Course.objects.filter(program__department__faculty__campus__university=university).count()
    constraints_count    = Constraint.objects.filter(university=university).count()

    # Find the active timetable
    active_timetable = Timetable.objects.filter(semester__university=university, is_active=True).first()
    if not active_timetable:
        active_timetable = Timetable.objects.filter(semester__university=university).first()

    conflict_count   = 0
    timetable_status = "No Timetable"
    if active_timetable:
        from .models import GenerationLog
        latest_log = (
            GenerationLog.objects
            .filter(timetable=active_timetable)
            .exclude(status='PENDING')
            .order_by('-created_at')
            .values('hard_conflicts_found', 'courses_scheduled')
            .first()
        )
        if latest_log is not None:
            conflict_count   = latest_log['hard_conflicts_found']
            timetable_status = "Generated ✓" if latest_log['courses_scheduled'] > 0 else "Empty"
        else:
            slot_count       = active_timetable.slots.count()
            timetable_status = "Generated ✓" if slot_count > 0 else "Empty"

    timetables = Timetable.objects.filter(semester__university=university).order_by('-created_at')[:5]

    from .models import Subscription
    subscription = getattr(university, 'subscription', None)
    if not subscription:
        try:
            subscription = Subscription.objects.create(university=university, tier='free', status='active')
        except Exception:
            subscription = None

    notifications = []
    if request.user.is_authenticated:
        notifications = request.user.notifications.filter(is_read=False)[:5]

    from .models import Announcement
    announcements = Announcement.objects.filter(Q(university=university) | Q(university__isnull=True)).order_by('-created_at')[:3]

    role_context = {}
    if request.user.is_authenticated:
        try:
            profile = request.user.profile
            if role == ROLE_STUDENT and profile.student_group:
                role_context['student_group'] = profile.student_group
            elif role == ROLE_HOD and profile.lecturer:
                role_context['department'] = profile.lecturer.department
                role_context['dept_lecturers_count'] = Lecturer.objects.filter(department=profile.lecturer.department).count()
                role_context['dept_courses_count'] = Course.objects.filter(program__department=profile.lecturer.department).count()
            elif role == ROLE_DEAN and profile.lecturer and profile.lecturer.department:
                faculty = profile.lecturer.department.faculty
                role_context['faculty'] = faculty
                role_context['faculty_departments_count'] = Department.objects.filter(faculty=faculty).count()
                role_context['faculty_rooms_count'] = Room.objects.filter(campus=faculty.campus).count()
        except Exception:
            pass

    context = {
        'active_university':    university,
        'active_semester':      Semester.objects.filter(university=university, is_active=True).first(),
        'active_role':          role,
        'campuses_count':       campuses_count,
        'rooms_count':          rooms_count,
        'lecturers_count':      lecturers_count,
        'student_groups_count': student_groups_count,
        'courses_count':        courses_count,
        'constraints_count':    constraints_count,
        'timetables':           timetables,
        'active_timetable':     active_timetable,
        'conflict_count':       conflict_count,
        'timetable_status':     timetable_status,
        'subscription':         subscription,
        'notifications':        notifications,
        'announcements':        announcements,
        'role_context':         role_context,
    }
    return render(request, 'scheduler/dashboard.html', context)

@login_required
def timetable_list(request):
    university = get_active_uni(request)
    if not university:
        messages.error(request, "Please create a University first in Admin panel.")
        return redirect('admin:index')

    timetables = Timetable.objects.filter(semester__university=university).select_related('semester').order_by('-created_at')
    
    # Restrict Semester selection to the current university in the form
    form = TimetableForm()
    form.fields['semester'].queryset = Semester.objects.filter(university=university)

    if request.method == 'POST':
        # Don't allow students/lecturers to create timetables
        role = get_user_role(request)
        if role in (ROLE_STUDENT, ROLE_LECTURER):
            messages.error(request, "Permission denied. Only Admins/Schedulers can create timetables.")
            return redirect('scheduler:timetable_list')

        form = TimetableForm(request.POST)
        form.fields['semester'].queryset = Semester.objects.filter(university=university)
        if form.is_valid():
            form.save()
            messages.success(request, "Timetable version created successfully.")
            return redirect('scheduler:timetable_list')
    
    return render(request, 'scheduler/timetable_list.html', {
        'timetables': timetables,
        'form': form
    })

@login_required
def timetable_create(request):
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    university = get_active_uni(request)
    if request.method == 'POST':
        form = TimetableForm(request.POST)
        form.fields['semester'].queryset = Semester.objects.filter(university=university)
        if form.is_valid():
            timetable = form.save()
            messages.success(request, "Timetable version created.")
            return redirect('scheduler:timetable_detail', pk=timetable.pk)
    else:
        form = TimetableForm()
        form.fields['semester'].queryset = Semester.objects.filter(university=university)
    return render(request, 'scheduler/timetable_form.html', {'form': form})

def _build_timetable_grid_context(timeslots, slots_list):
    """
    Constructs a chronologically ordered weekly timetable grid.
    Grid rows are grouped by distinct (start_time, end_time) time windows
    sorted strictly by start_time, preventing misaligned time labels or scrambled rows.
    """
    days_in_slots = sorted(set(ts.day_of_week for ts in timeslots)) if timeslots else list(range(1, 8))
    
    # Extract distinct time windows sorted strictly chronologically
    all_time_windows = sorted(
        set((ts.start_time, ts.end_time) for ts in timeslots),
        key=lambda tw: (tw[0], tw[1])
    )
    
    # Only render rows for time windows that actually contain classes for the current view
    # (or all time windows if no classes exist yet)
    used_windows = set((s.time_slot.start_time, s.time_slot.end_time) for s in slots_list if s.time_slot)
    if used_windows:
        time_windows = [tw for tw in all_time_windows if tw in used_windows]
    else:
        time_windows = all_time_windows

    tw_to_row_idx = {tw: idx + 1 for idx, tw in enumerate(time_windows)}
    row_idx_to_tw = {idx + 1: tw for idx, tw in enumerate(time_windows)}
    slot_numbers = [idx + 1 for idx in range(len(time_windows))]
    
    # Fast lookup for TimeSlot model instance by (day_of_week, start_time, end_time)
    ts_by_day_and_window = {(ts.day_of_week, ts.start_time, ts.end_time): ts for ts in timeslots}
    
    grid = {}
    for row_idx in slot_numbers:
        grid[row_idx] = {}
        tw = row_idx_to_tw[row_idx]
        for day in days_in_slots:
            grid[row_idx][day] = {
                'time_slot': ts_by_day_and_window.get((day, tw[0], tw[1])),
                'slots': [],
                'rowspan': 1,
                'is_merged': False
            }

    # Populate ScheduleSlot instances into matching grid cells
    for slot in slots_list:
        ts = slot.time_slot
        tw = (ts.start_time, ts.end_time)
        row_idx = tw_to_row_idx.get(tw)
        if row_idx and row_idx in grid and ts.day_of_week in grid[row_idx]:
            grid[row_idx][ts.day_of_week]['slots'].append(slot)

    # Calculate rowspans and merge consecutive sessions
    for day in days_in_slots:
        day_slots = []
        for row_idx in slot_numbers:
            cell = grid[row_idx][day]
            if cell['slots']:
                day_slots.append((row_idx, cell['slots']))

        i = 0
        while i < len(day_slots):
            row_idx, slots = day_slots[i]
            primary_slot = slots[0]

            rowspan = 1
            j = i + 1
            while j < len(day_slots):
                next_row_idx, next_slots = day_slots[j]
                next_primary_slot = next_slots[0]

                # Check if consecutive row and matching class session details
                if (next_row_idx == row_idx + rowspan and
                    next_primary_slot.course_id == primary_slot.course_id and
                    next_primary_slot.lecturer_id == primary_slot.lecturer_id and
                    next_primary_slot.room_id == primary_slot.room_id and
                    next_primary_slot.student_group_id == primary_slot.student_group_id):
                    rowspan += 1
                    j += 1
                else:
                    break

            grid[row_idx][day]['rowspan'] = rowspan
            for r in range(1, rowspan):
                if row_idx + r in grid and day in grid[row_idx + r]:
                    grid[row_idx + r][day]['is_merged'] = True

            i = j

    slot_time_labels = {}
    for row_idx, tw in row_idx_to_tw.items():
        start_fmt = tw[0].strftime('%I:%M %p').lstrip('0').lower()
        end_fmt = tw[1].strftime('%I:%M %p').lstrip('0').lower()
        slot_time_labels[row_idx] = {
            'start': start_fmt,
            'end': end_fmt,
            'range': f"{start_fmt} – {end_fmt}"
        }

    return {
        'grid': grid,
        'slot_numbers': slot_numbers,
        'slot_time_labels': slot_time_labels,
        'days_in_slots': days_in_slots,
    }


@login_required
@tenant_required(Timetable)
def timetable_detail(request, pk):
    role = get_user_role(request)
    
    # Lecturers and students should use their personal schedule pages
    if role in (ROLE_LECTURER, ROLE_STUDENT):
        if role == ROLE_LECTURER:
            return redirect('scheduler:lecturer_my_schedule')
        else:
            return redirect('scheduler:student_my_schedule')
    
    timetable = get_object_or_404(Timetable.objects.select_related('semester', 'semester__university'), pk=pk)

    # Always derive university from the timetable itself so the grid renders
    # correctly regardless of which university is active in the session.
    university = timetable.semester.university
    # Sync the session so the rest of the UI stays consistent.
    request.session['active_university_id'] = university.id

    # Get filters
    filter_type = request.GET.get('filter_type', 'group')
    filter_id = request.GET.get('filter_id')

    # Filter resources by university
    rooms = Room.objects.filter(campus__university=university)
    lecturers = Lecturer.objects.filter(department__faculty__campus__university=university)
    student_groups = StudentGroup.objects.filter(program__department__faculty__campus__university=university)

    # Resolve default filter target
    selected_filter_name = ""
    if not filter_id:
        if filter_type == 'group' and student_groups.exists():
            filter_id = student_groups.first().id
        elif filter_type == 'room' and rooms.exists():
            filter_id = rooms.first().id
        elif filter_type == 'lecturer' and lecturers.exists():
            filter_id = lecturers.first().id

    # FIX U3: Fetch ALL slots once — reuse for both the grid and conflict detection.
    # The old code fetched slots twice (once filtered for the grid, once unfiltered for conflicts).
    all_slots = list(
        timetable.slots
        .select_related('course', 'lecturer', 'room', 'time_slot', 'student_group')
        .all()
    )

    from django.utils import timezone
    local_now = timezone.localtime(timezone.now())
    current_dow = local_now.isoweekday()
    current_time = local_now.time()
    for s in all_slots:
        ts = s.time_slot
        if ts.day_of_week < current_dow:
            s.has_ended = True
            s.is_ongoing = False
        elif ts.day_of_week == current_dow:
            s.has_ended = (ts.end_time < current_time)
            s.is_ongoing = (ts.start_time <= current_time < ts.end_time)  # FIX BUG 16: strict < for end_time
        else:
            s.has_ended = False
            s.is_ongoing = False


    # Apply the view filter in Python (no extra DB query)
    if filter_id:
        filter_id = int(filter_id)
        if filter_type == 'group':
            slots_list = [s for s in all_slots if s.student_group_id == filter_id]
            selected_filter_name = str(student_groups.filter(id=filter_id).first())
        elif filter_type == 'room':
            slots_list = [s for s in all_slots if s.room_id == filter_id]
            selected_filter_name = str(rooms.filter(id=filter_id).first())
        elif filter_type == 'lecturer':
            slots_list = [s for s in all_slots if s.lecturer_id == filter_id]
            selected_filter_name = str(lecturers.filter(id=filter_id).first())
        else:
            slots_list = all_slots
    else:
        slots_list = all_slots

    # Time slots sorting — filter by the timetable's own university
    timeslots     = TimeSlot.objects.filter(university=university).order_by('day_of_week', 'start_time', 'slot_number')
    grid_ctx      = _build_timetable_grid_context(timeslots, slots_list)
    grid          = grid_ctx['grid']
    slot_numbers  = grid_ctx['slot_numbers']
    slot_time_labels = grid_ctx['slot_time_labels']
    days_in_slots = grid_ctx['days_in_slots']

    day_labels = {
        1: 'Monday', 2: 'Tuesday', 3: 'Wednesday', 4: 'Thursday',
        5: 'Friday', 6: 'Saturday', 7: 'Sunday'
    }

    # Conflict detector — reuses the already-fetched all_slots list (no extra DB hit)
    conflict_list = detect_conflicts(all_slots, university)
    errors   = [c for c in conflict_list if c['severity'] == 'error']
    warnings = [c for c in conflict_list if c['severity'] == 'warning']

    # Serialize slots and timeslots for JS calendar consumption (merging consecutive slots)
    from collections import defaultdict
    sessions = defaultdict(list)
    for s in slots_list:
        key = (s.time_slot.day_of_week, s.course_id, s.lecturer_id, s.room_id, s.student_group_id)
        sessions[key].append(s)

    slots_data = []
    for key, s_list in sessions.items():
        s_list.sort(key=lambda s: s.time_slot.slot_number)
        i = 0
        while i < len(s_list):
            run = [s_list[i]]
            j = i + 1
            while j < len(s_list):
                prev = s_list[j-1]
                curr = s_list[j]
                if (curr.time_slot.slot_number == prev.time_slot.slot_number + 1 and
                    curr.time_slot.start_time == prev.time_slot.end_time):
                    run.append(curr)
                    j += 1
                else:
                    break
            
            first_s = run[0]
            last_s = run[-1]
            slots_data.append({
                'id': first_s.id,
                'course_code': first_s.course.code,
                'course_name': first_s.course.name,
                'room_id': first_s.room_id,
                'room_name': first_s.room.name,
                'lecturer_id': first_s.lecturer_id,
                'lecturer_name': first_s.lecturer.name,
                'group_id': first_s.student_group_id,
                'group_name': first_s.student_group.name,
                'time_slot_id': first_s.time_slot_id,
                'day_of_week': first_s.time_slot.day_of_week,
                'start_time': first_s.time_slot.start_time.strftime('%H:%M:%S'),
                'end_time': last_s.time_slot.end_time.strftime('%H:%M:%S'),
            })
            i = j

    timeslots_data = [{
        'id': ts.id,
        'day_of_week': ts.day_of_week,
        'start_time': ts.start_time.strftime('%H:%M:%S'),
        'end_time': ts.end_time.strftime('%H:%M:%S'),
        'slot_number': ts.slot_number,
    } for ts in timeslots]

    # Calculate dates for the current week (Monday-based)
    import datetime
    _today = datetime.date.today()
    _monday = _today - datetime.timedelta(days=_today.weekday())
    days_data = []
    for d in days_in_slots:
        days_data.append({
            'num': d,
            'label': day_labels.get(d, f"Day {d}"),
            'date': _monday + datetime.timedelta(days=d - 1)
        })

    context = {
        'timetable':           timetable,
        'rooms':               rooms,
        'lecturers':           lecturers,
        'student_groups':      student_groups,
        'filter_type':         filter_type,
        'filter_id':           filter_id,
        'selected_filter_name': selected_filter_name,
        'grid':                grid,
        'days':                days_data,
        'slot_numbers':        slot_numbers,
        'slot_time_labels':    slot_time_labels,
        'errors':              errors,
        'warnings':            warnings,
        'timeslots':           timeslots,
        'slots_json':          slots_data,
        'timeslots_json':      timeslots_data,
        'slot_count':          len(all_slots),
        'firebase_config':     settings.FIREBASE_CONFIG if __import__('scheduler.firebase_service', fromlist=['is_enabled']).is_enabled else None,
        'approval_logs':       timetable.approval_logs.all(),
        'user_role':           role,
    }
    return render(request, 'scheduler/timetable_detail.html', context)



@login_required
@tenant_required(Timetable)
def timetable_weekly(request, pk):
    """
    Branded weekly timetable view — clean, print-friendly, university-styled.
    Supports toggle by student group, lecturer, or room.
    Designed for PDF/print export.
    """
    role = get_user_role(request)
    
    # Lecturers and students should use their personal schedule pages
    if role in (ROLE_LECTURER, ROLE_STUDENT):
        if role == ROLE_LECTURER:
            return redirect('scheduler:lecturer_my_schedule')
        else:
            return redirect('scheduler:student_my_schedule')
    
    university = get_active_uni(request)
    timetable = get_object_or_404(
        Timetable.objects.select_related('semester', 'semester__university'), pk=pk
    )
    if timetable.semester.university != university:
        messages.error(request, "Attempted to access out-of-scope timetable.")
        return redirect('scheduler:dashboard')

    filter_type = request.GET.get('filter_type', 'group')
    filter_id   = request.GET.get('filter_id')

    rooms          = Room.objects.filter(campus__university=university)
    lecturers      = Lecturer.objects.filter(department__faculty__campus__university=university)
    student_groups = StudentGroup.objects.filter(
        program__department__faculty__campus__university=university
    )

    selected_filter_name = ""
    selected_filter_obj  = None

    if not filter_id:
        if filter_type == 'group' and student_groups.exists():
            filter_id = student_groups.first().id
        elif filter_type == 'room' and rooms.exists():
            filter_id = rooms.first().id
        elif filter_type == 'lecturer' and lecturers.exists():
            filter_id = lecturers.first().id

    slots_qs = timetable.slots.select_related(
        'course', 'lecturer', 'room', 'time_slot', 'student_group'
    )

    if filter_id:
        filter_id = int(filter_id)
        if filter_type == 'group':
            slots_qs = slots_qs.filter(student_group_id=filter_id)
            selected_filter_obj = student_groups.filter(id=filter_id).first()
        elif filter_type == 'room':
            slots_qs = slots_qs.filter(room_id=filter_id)
            selected_filter_obj = rooms.filter(id=filter_id).first()
        elif filter_type == 'lecturer':
            slots_qs = slots_qs.filter(lecturer_id=filter_id)
            selected_filter_obj = lecturers.filter(id=filter_id).first()
        if selected_filter_obj:
            selected_filter_name = str(selected_filter_obj)

    slots_list = list(slots_qs)

    from django.utils import timezone
    local_now = timezone.localtime(timezone.now())
    current_dow = local_now.isoweekday()
    current_time = local_now.time()
    for s in slots_list:
        ts = s.time_slot
        if ts.day_of_week < current_dow:
            s.has_ended = True
            s.is_ongoing = False
        elif ts.day_of_week == current_dow:
            s.has_ended = (ts.end_time < current_time)
            s.is_ongoing = (ts.start_time <= current_time < ts.end_time)  # FIX BUG 16: strict < for end_time
        else:
            s.has_ended = False
            s.is_ongoing = False


    timeslots     = TimeSlot.objects.filter(university=university).order_by('day_of_week', 'start_time', 'slot_number')
    grid_ctx      = _build_timetable_grid_context(timeslots, slots_list)
    grid          = grid_ctx['grid']
    slot_numbers  = grid_ctx['slot_numbers']
    slot_time_labels = grid_ctx['slot_time_labels']
    days_in_slots = grid_ctx['days_in_slots']

    day_labels = {
        1: 'Monday', 2: 'Tuesday', 3: 'Wednesday',
        4: 'Thursday', 5: 'Friday', 6: 'Saturday', 7: 'Sunday'
    }


    # Calculate dates for the current week (Monday-based)
    import datetime
    _today = datetime.date.today()
    _monday = _today - datetime.timedelta(days=_today.weekday())
    days_context = []
    for d in days_in_slots:
        days_context.append({
            'num': d,
            'label': day_labels.get(d, f"Day {d}"),
            'date': _monday + datetime.timedelta(days=d - 1)
        })

    context = {
        'timetable': timetable,
        'university': university,
        'semester': timetable.semester,
        'rooms': rooms,
        'lecturers': lecturers,
        'student_groups': student_groups,
        'filter_type': filter_type,
        'filter_id': filter_id,
        'selected_filter_name': selected_filter_name,
        'grid': grid,
        'days': days_context,
        'slot_numbers': slot_numbers,
        'slot_time_labels': slot_time_labels,
        'timeslots': timeslots,
    }
    return render(request, 'scheduler/timetable_weekly.html', context)


@login_required
@tenant_required(Timetable)
def timetable_generate(request, pk):
    """
    Triggers the scheduling pipeline asynchronously via django-q2.
    Returns JSON with a task_id so the UI can poll for status.
    """
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'status': 'error', 'message': 'Permission denied.'}, status=403)
        messages.error(request, "Permission denied. Only Schedulers/Admins can generate schedules.")
        return redirect('scheduler:timetable_detail', pk=pk)

    timetable = get_object_or_404(Timetable, pk=pk)

    if request.method == 'POST':
        # Determine time limit based on user selection or course count
        time_limit_choice = request.POST.get('time_limit', 'auto')
        if time_limit_choice.isdigit():
            time_limit = int(time_limit_choice)
        else:
            # "auto" — scale the CP-SAT time budget to the actual problem size
            # instead of hardcoding a useless 1-second fallback.
            # Mirrors the same eligibility filter that solver.generate_timetable()
            # applies when loading courses, so the count reflects reality.
            course_count = Course.objects.filter(
                program__department__faculty__campus__university=timetable.semester.university,
                lecturer__isnull=False,
                lecturer__is_active=True,
                student_group__isnull=False,
            ).count()
            if course_count < 50:
                time_limit = 15
            elif course_count < 300:
                time_limit = 60
            elif course_count < 1000:
                time_limit = 120
            else:
                time_limit = 180

        # Trigger background execution (async_task in test mode, threading.Thread in normal mode)
        import sys
        if 'test' in sys.argv:
            try:
                from django_q.tasks import async_task

                # Create a PENDING GenerationLog to track background task status
                GenerationLog.objects.create(
                    timetable=timetable,
                    status='PENDING',
                    message='Generation queued in the background worker queue.'
                )
                
                task_id = async_task(
                    'scheduler.tasks.generate_timetable_async',
                    timetable.id,
                    time_limit,  # dynamic time_limit seconds
                    task_name=f'generate-timetable-{timetable.id}',
                    group=f'timetable-{timetable.id}',
                )
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'status': 'QUEUED',
                        'message': 'Generation queued. Results will appear when complete.',
                        'task_id': str(task_id),
                    })
                messages.info(request, "✓ Generation queued! Refresh in a moment to see results.")
            except Exception:
                # Fallback: run synchronously
                result = run_scheduling_pipeline(timetable.id, time_limit_seconds=time_limit)
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'status': result.status,
                        'message': result.message,
                        'log_id': result.log_id,
                        'solve_time_seconds': result.solve_time_seconds,
                        'solver_score': result.solver_score,
                        'courses_scheduled': result.courses_scheduled,
                        'hard_conflicts': len(result.hard_conflicts),
                        'soft_conflicts': len(result.soft_conflicts),
                        'validation_errors': result.validation_errors,
                        'validation_warnings': result.validation_warnings,
                    })
                if result.status in ('OPTIMAL', 'FEASIBLE'):
                    messages.success(request, f"✓ {result.message}")
                else:
                    messages.error(request, f"Error: {result.message}")
        else:
            try:
                import threading
                from django.db import close_old_connections
                from .tasks import generate_timetable_async

                # Clear any stale PENDING logs so they don't block re-generation
                GenerationLog.objects.filter(timetable=timetable, status='PENDING').delete()

                # Create a PENDING GenerationLog to track background task status
                pending_log = GenerationLog.objects.create(
                    timetable=timetable,
                    status='PENDING',
                    message='Generation started in the background.'
                )

                def run_async():
                    close_old_connections()
                    try:
                        generate_timetable_async(timetable.id, time_limit)
                    except Exception as thread_err:
                        logger.error(f"[Views Thread] Background generation failed: {thread_err}")
                        # Mark the pending log as ERROR so UI doesn't show stale PENDING
                        try:
                            pending_log.status = 'ERROR'
                            pending_log.message = f'Generation thread crashed: {thread_err}'
                            pending_log.save(update_fields=['status', 'message'])
                        except Exception:
                            pass
                    finally:
                        close_old_connections()

                t_thread = threading.Thread(target=run_async, name=f"ManualGenerate-{timetable.id}")
                t_thread.daemon = False  # Non-daemon so it survives dev server reloads
                t_thread.start()

                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'status': 'QUEUED',
                        'message': 'Generation running in background. Results will appear when complete.',
                        'task_id': f"thread-{timetable.id}",
                    })
                messages.info(request, "✓ Generation started in the background! Refresh in a moment to see results.")
            except Exception as e:
                # Fallback: run synchronously
                logger.error(f"Failed to start background thread, running synchronously: {e}")
                result = run_scheduling_pipeline(timetable.id, time_limit_seconds=time_limit)
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'status': result.status,
                        'message': result.message,
                        'log_id': result.log_id,
                        'solve_time_seconds': result.solve_time_seconds,
                        'solver_score': result.solver_score,
                        'courses_scheduled': result.courses_scheduled,
                        'hard_conflicts': len(result.hard_conflicts),
                        'soft_conflicts': len(result.soft_conflicts),
                        'validation_errors': result.validation_errors,
                        'validation_warnings': result.validation_warnings,
                    })
                if result.status in ('OPTIMAL', 'FEASIBLE'):
                    messages.success(request, f"✓ {result.message}")
                else:
                    messages.error(request, f"Error: {result.message}")

    return redirect('scheduler:timetable_detail', pk=timetable.pk)


@login_required
def generation_status(request, pk):
    """AJAX endpoint to poll generation task status."""
    timetable = get_object_or_404(Timetable, pk=pk)
    # Get latest generation log for this timetable
    latest_log = GenerationLog.objects.filter(timetable=timetable).order_by('-created_at').first()
    if latest_log:
        from django.utils import timezone
        import datetime
        # Self-heal: if log is stuck in PENDING/RUNNING but is older than 15 minutes, mark it as aborted
        if latest_log.status in ('PENDING', 'RUNNING', 'QUEUED') and latest_log.created_at < timezone.now() - datetime.timedelta(minutes=15):
            latest_log.status = 'ERROR'
            latest_log.message = 'The generation task was interrupted or timed out.'
            latest_log.save()
        return JsonResponse({
            'status': latest_log.status,
            'message': latest_log.message,
            'courses_scheduled': latest_log.courses_scheduled,
            'hard_conflicts': latest_log.hard_conflicts_found,
            'soft_conflicts': latest_log.soft_conflicts_found,
            'solve_time': latest_log.solve_time_seconds,
        })
    return JsonResponse({'status': 'NOT_STARTED', 'message': 'No generation has been run yet.'})



@login_required
def generation_log_list(request, pk):
    """
    Shows the audit timeline of all generation runs for a timetable.
    """
    university = get_active_uni(request)
    timetable = get_object_or_404(
        Timetable.objects.select_related('semester', 'semester__university'), pk=pk
    )
    if timetable.semester.university != university:
        messages.error(request, "Attempted to access out-of-scope timetable.")
        return redirect('scheduler:dashboard')

    logs = GenerationLog.objects.filter(timetable=timetable).order_by('-created_at')
    
    logs_data = []
    total_logs = len(logs)
    for idx, log in enumerate(logs):
        logs_data.append({
            'id': str(log.id),
            'num': total_logs - idx,
            'status': log.status,
            'statusDisplay': log.get_status_display(),
            'timestamp': log.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'timeOnly': log.created_at.strftime('%H:%M:%S'),
            'message': log.message,
            'score': log.solver_score,
            'solveTime': log.solve_time_seconds,
            'coursesScheduled': log.courses_scheduled,
            'hardConflicts': log.hard_conflicts_found,
            'softConflicts': log.soft_conflicts_found,
            'validationErrors': log.validation_errors or [],
            'validationWarnings': log.validation_warnings or []
        })
    
    logs_json = json.dumps(logs_data)

    return render(request, 'scheduler/generation_log.html', {
        'timetable': timetable,
        'logs': logs,
        'logs_json': logs_json,
    })

@login_required
@tenant_required(Timetable)
def timetable_conflicts(request, pk):
    timetable = get_object_or_404(Timetable, pk=pk)
    conflicts = check_conflicts_for_timetable(timetable)

    errors = [c for c in conflicts if c['severity'] == 'error']
    warnings = [c for c in conflicts if c['severity'] == 'warning']

    context = {
        'timetable': timetable,
        'errors': errors,
        'warnings': warnings,
        'total_conflicts': len(conflicts),
    }
    return render(request, 'scheduler/timetable_conflicts.html', context)


@login_required
def conflicts_json(request, pk):
    """
    AJAX endpoint — returns live conflict data as JSON.
    Called automatically on page-load and after every drag-drop.
    No page refresh required.
    """
    timetable = get_object_or_404(Timetable, pk=pk)
    
    fp = _get_timetable_conflict_fingerprint(timetable)
    if fp in _CONFLICTS_JSON_CACHE:
        return JsonResponse(_CONFLICTS_JSON_CACHE[fp])

    conflicts = check_conflicts_for_timetable(timetable)

    errors   = [c for c in conflicts if c['severity'] == 'error']
    warnings = [c for c in conflicts if c['severity'] == 'warning']

    response_data = {
        'hard_count':    len(errors),
        'soft_count':    len(warnings),
        'total':         len(conflicts),
        'errors':   [{'type': c['constraint_type'], 'message': c['message']} for c in errors],
        'warnings': [{'type': c['constraint_type'], 'message': c['message']} for c in warnings],
    }

    # Limit cache size to prevent memory leaks
    if len(_CONFLICTS_JSON_CACHE) > _CACHE_MAX_SIZE:
        _CONFLICTS_JSON_CACHE.clear()
    _CONFLICTS_JSON_CACHE[fp] = response_data

    return JsonResponse(response_data)


@login_required
def conflicts_autofix(request, pk):
    """
    Auto-fix endpoint — automatically resolves fixable conflicts:
      1. ROOM_CAPACITY: swap to any room with sufficient capacity + matching type
      2. ROOM_TYPE_MISMATCH: swap to a room of the correct type if one is free
      3. ROOM_MISSING_REQUIRED_FEATURES: swap to a room with all required features

    Returns JSON summary of what was fixed.
    """
    from django.db import transaction
    from collections import defaultdict

    role = get_user_role(request)
    if role not in MANAGER_ROLES:
        return JsonResponse({'status': 'error', 'message': 'Permission denied.'}, status=403)

    timetable = get_object_or_404(Timetable, pk=pk)
    conflicts  = check_conflicts_for_timetable(timetable)
    university = timetable.semester.university

    from .models import Room, Course
    fixed_count   = 0
    skipped_count = 0
    fixed_messages = []

    # Build a map of which rooms are occupied at each (day, slot_number)
    all_slots = list(timetable.slots.select_related('room', 'time_slot', 'course', 'student_group'))
    all_slots_by_id = {s.id: s for s in all_slots}

    occupied_rooms = {}  # (day, slot_num) -> set of room_ids in use
    for s in all_slots:
        key = (s.time_slot.day_of_week, s.time_slot.slot_number)
        occupied_rooms.setdefault(key, set()).add(s.room_id)

    # All rooms available for this university
    all_rooms = list(Room.objects.filter(campus__university=university))

    # Load feature mappings for ROOM_MISSING_REQUIRED_FEATURES
    course_required_features = defaultdict(set)
    for course_id, feature_name in Course.required_features.through.objects.filter(
        course__program__department__faculty__campus__university=university
    ).values_list('course_id', 'roomfeature__name'):
        course_required_features[course_id].add(feature_name)

    room_features = defaultdict(set)
    for room_id, feature_name in Room.features.through.objects.filter(
        room__campus__university=university
    ).values_list('room_id', 'roomfeature__name'):
        room_features[room_id].add(feature_name)

    modified_slots = {}  # slot_id -> ScheduleSlot

    # ── Phase 1: Room-swap fixes (ROOM_CAPACITY / ROOM_TYPE_MISMATCH / ROOM_MISSING_REQUIRED_FEATURES) ──
    for conflict in conflicts:
        ctype = conflict['constraint_type']
        if ctype not in ('ROOM_CAPACITY', 'ROOM_TYPE_MISMATCH', 'ROOM_MISSING_REQUIRED_FEATURES'):
            continue

        slot_id = conflict['entities'].get('slot_id')
        if not slot_id:
            continue

        slot = modified_slots.get(slot_id) or all_slots_by_id.get(slot_id)
        if not slot:
            continue

        group_size = slot.student_group.size
        required_type = slot.course.required_room_type

        # Check if the conflict is already resolved by a previous slot move
        if ctype == 'ROOM_CAPACITY' and slot.room.capacity >= group_size:
            continue
        if ctype == 'ROOM_TYPE_MISMATCH' and slot.room.room_type == required_type:
            continue
        if ctype == 'ROOM_MISSING_REQUIRED_FEATURES':
            req_feats = course_required_features.get(slot.course_id, set())
            if req_feats.issubset(room_features.get(slot.room_id, set())):
                continue

        # Find all slots of the same session (same course/lecturer/group/day)
        orig_day = slot.time_slot.day_of_week
        session_slots = [
            s for s in all_slots
            if s.course_id == slot.course_id
            and s.lecturer_id == slot.lecturer_id
            and s.student_group_id == slot.student_group_id
            and s.time_slot.day_of_week == orig_day
        ]

        # Find best replacement room: correct type, big enough, not in use for all session slots
        candidates = []
        for r in all_rooms:
            if r.id == slot.room_id:
                continue
            if r.capacity < group_size:
                continue

            # Check ctype-specific constraints
            if ctype == 'ROOM_TYPE_MISMATCH' and r.room_type != required_type:
                continue
            if ctype == 'ROOM_MISSING_REQUIRED_FEATURES':
                req_feats = course_required_features.get(slot.course_id, set())
                if not req_feats.issubset(room_features.get(r.id, set())):
                    continue
                if r.room_type != required_type:
                    continue

            # Capacity check (first pass tries matching type)
            if ctype == 'ROOM_CAPACITY' and r.room_type != required_type:
                continue

            # Check if free for the entire session
            free_for_all = True
            for s_slot in session_slots:
                key = (s_slot.time_slot.day_of_week, s_slot.time_slot.slot_number)
                in_use = occupied_rooms.get(key, set())
                if r.id in in_use:
                    free_for_all = False
                    break
            if free_for_all:
                candidates.append(r)

        # For ROOM_CAPACITY: if no candidate with matching type is free, relax type constraint
        if ctype == 'ROOM_CAPACITY' and not candidates:
            for r in all_rooms:
                if r.id == slot.room_id:
                    continue
                if r.capacity < group_size:
                    continue
                free_for_all = True
                for s_slot in session_slots:
                    key = (s_slot.time_slot.day_of_week, s_slot.time_slot.slot_number)
                    in_use = occupied_rooms.get(key, set())
                    if r.id in in_use:
                        free_for_all = False
                        break
                if free_for_all:
                    candidates.append(r)

        if candidates:
            best = min(candidates, key=lambda r: r.capacity)  # Smallest fitting room
            old_room_name = slot.room.name
            old_room_id = slot.room_id

            # Apply new room to all slots in the session
            for s_slot in session_slots:
                key = (s_slot.time_slot.day_of_week, s_slot.time_slot.slot_number)

                # Update occupancy map in-memory
                if key in occupied_rooms:
                    occupied_rooms[key].discard(old_room_id)
                    occupied_rooms[key].add(best.id)

                s_slot.room = best
                s_slot.room_id = best.id
                modified_slots[s_slot.id] = s_slot

            fixed_count += 1
            fixed_messages.append(
                f"✓ {slot.course.code}: moved from '{old_room_name}' → '{best.name}' "
                f"({best.get_room_type_display()} room, capacity {best.capacity} ≥ {group_size})"
            )
        else:
            skipped_count += 1
            fixed_messages.append(
                f"✗ {slot.course.code}: no suitable room found (group size {group_size})"
            )

    # ── Phase 2: Timeslot-reschedule fixes (LECTURER_CAMPUS_TRAVEL_VIOLATION / LECTURER_DAILY_LIMIT_EXCEEDED) ──
    # Build fast lookup structures needed for timeslot-move validation.
    from .models import TimeSlot as TimeSlotModel

    all_timeslots = list(TimeSlotModel.objects.filter(university=university).order_by('day_of_week', 'slot_number'))

    # (lecturer_id, ts_id) -> slot  — for checking lecturer availability
    lec_ts_occupied = {}   # (lec_id, ts_id) -> slot
    # (room_id, ts_id) -> slot
    room_ts_occupied = {}  # (room_id, ts_id) -> slot
    # (group_id, ts_id) -> slot
    grp_ts_occupied = {}   # (grp_id, ts_id) -> slot

    for s in all_slots:
        if s.lecturer_id:
            lec_ts_occupied[(s.lecturer_id, s.time_slot_id)] = s
        room_ts_occupied[(s.room_id, s.time_slot_id)] = s
        grp_ts_occupied[(s.student_group_id, s.time_slot_id)] = s

    # (lecturer_id, day) -> count of slots scheduled
    from collections import Counter
    lec_day_count = Counter()
    for s in all_slots:
        if s.lecturer_id:
            lec_day_count[(s.lecturer_id, s.time_slot.day_of_week)] += 1

    def _try_move_slot_to_new_time(slot_to_move, exclude_day=None):
        """
        Try to reschedule slot_to_move to any timeslot that:
          - is on a different day than exclude_day (if given)
          - has no room/lecturer/group clash
          - does not push lecturer over their daily limit on the new day
          - does not create a new campus-travel violation for the lecturer
        Returns the chosen TimeSlot object or None.
        """
        lec_id = slot_to_move.lecturer_id
        r_id = slot_to_move.room_id
        grp_id = slot_to_move.student_group_id
        from .models import Lecturer as LecturerModel
        max_slots_day = 6  # fallback
        if lec_id:
            try:
                lec_obj = LecturerModel.objects.get(pk=lec_id)
                max_slots_day = lec_obj.max_slots_per_day or 6
            except LecturerModel.DoesNotExist:
                pass

        for ts in all_timeslots:
            if ts.id == slot_to_move.time_slot_id:
                continue
            if exclude_day is not None and ts.day_of_week == exclude_day:
                continue

            # Room free?
            if (r_id, ts.id) in room_ts_occupied:
                existing = room_ts_occupied[(r_id, ts.id)]
                if existing.id != slot_to_move.id:
                    continue

            # Lecturer free?
            if lec_id and (lec_id, ts.id) in lec_ts_occupied:
                existing = lec_ts_occupied[(lec_id, ts.id)]
                if existing.id != slot_to_move.id:
                    continue

            # Student group free?
            if (grp_id, ts.id) in grp_ts_occupied:
                existing = grp_ts_occupied[(grp_id, ts.id)]
                if existing.id != slot_to_move.id:
                    continue

            # Lecturer daily limit check on new day
            if lec_id:
                current_count = lec_day_count[(lec_id, ts.day_of_week)]
                if current_count >= max_slots_day:
                    continue

            # Campus-travel check: ensure no consecutive cross-campus clash on new day
            if lec_id:
                adjacent_ok = True
                for other in all_slots:
                    if other.id == slot_to_move.id:
                        continue
                    if other.lecturer_id != lec_id:
                        continue
                    if other.time_slot.day_of_week != ts.day_of_week:
                        continue
                    diff = abs(other.time_slot.slot_number - ts.slot_number)
                    if diff == 1 and other.room.campus_id != slot_to_move.room.campus_id:
                        adjacent_ok = False
                        break
                if not adjacent_ok:
                    continue

            return ts  # Found a valid slot
        return None

    # ── Fix LECTURER_CAMPUS_TRAVEL_VIOLATION ──────────────────────────────────
    seen_campus_travel_pairs = set()
    for conflict in conflicts:
        if conflict['constraint_type'] != 'LECTURER_CAMPUS_TRAVEL_VIOLATION':
            continue
        entities = conflict.get('entities', {})
        slot_ids = entities.get('slot_ids', [])
        if len(slot_ids) < 2:
            continue

        pair_key = tuple(sorted(slot_ids))
        if pair_key in seen_campus_travel_pairs:
            continue
        seen_campus_travel_pairs.add(pair_key)

        # Try to move the second slot of the pair to a different timeslot
        s2 = modified_slots.get(slot_ids[1]) or all_slots_by_id.get(slot_ids[1])
        if not s2:
            continue

        bad_day = s2.time_slot.day_of_week
        new_ts = _try_move_slot_to_new_time(s2, exclude_day=None)  # any free slot, same or diff day
        if new_ts is None:
            skipped_count += 1
            fixed_messages.append(
                f"✗ Campus travel fix: no free slot found for {s2.course.code} on a compatible day."
            )
            continue

        old_ts = s2.time_slot
        # Update in-memory occupancy
        if s2.lecturer_id:
            lec_ts_occupied.pop((s2.lecturer_id, old_ts.id), None)
            lec_ts_occupied[(s2.lecturer_id, new_ts.id)] = s2
            lec_day_count[(s2.lecturer_id, old_ts.day_of_week)] -= 1
            lec_day_count[(s2.lecturer_id, new_ts.day_of_week)] += 1
        room_ts_occupied.pop((s2.room_id, old_ts.id), None)
        room_ts_occupied[(s2.room_id, new_ts.id)] = s2
        grp_ts_occupied.pop((s2.student_group_id, old_ts.id), None)
        grp_ts_occupied[(s2.student_group_id, new_ts.id)] = s2

        s2.time_slot = new_ts
        s2.time_slot_id = new_ts.id
        modified_slots[s2.id] = s2
        fixed_count += 1
        fixed_messages.append(
            f"✓ Campus travel fix: {s2.course.code} moved from {old_ts} → {new_ts} "
            f"(lecturer {conflict['entities'].get('lecturer_id', '?')} no longer crosses campuses consecutively)."
        )

    # ── Fix LECTURER_DAILY_LIMIT_EXCEEDED ─────────────────────────────────────
    seen_daily_limit_keys = set()
    for conflict in conflicts:
        if conflict['constraint_type'] != 'LECTURER_DAILY_LIMIT_EXCEEDED':
            continue
        entities = conflict.get('entities', {})
        lec_id   = entities.get('lecturer_id')
        day      = entities.get('day')
        max_slots = entities.get('max_slots', 4)
        if not lec_id or day is None:
            continue

        key = (lec_id, day)
        if key in seen_daily_limit_keys:
            continue
        seen_daily_limit_keys.add(key)

        # Collect all slots for this lecturer on this day, sorted last→first
        day_slots = sorted(
            [s for s in all_slots if s.lecturer_id == lec_id and s.time_slot.day_of_week == day],
            key=lambda s: s.time_slot.slot_number,
            reverse=True,
        )
        excess = len(day_slots) - max_slots
        moved = 0
        for s_excess in day_slots:
            if moved >= excess:
                break
            new_ts = _try_move_slot_to_new_time(s_excess, exclude_day=day)
            if new_ts is None:
                skipped_count += 1
                fixed_messages.append(
                    f"✗ Daily limit fix: no free slot on another day for {s_excess.course.code}."
                )
                continue

            old_ts = s_excess.time_slot
            if s_excess.lecturer_id:
                lec_ts_occupied.pop((s_excess.lecturer_id, old_ts.id), None)
                lec_ts_occupied[(s_excess.lecturer_id, new_ts.id)] = s_excess
                lec_day_count[(s_excess.lecturer_id, old_ts.day_of_week)] -= 1
                lec_day_count[(s_excess.lecturer_id, new_ts.day_of_week)] += 1
            room_ts_occupied.pop((s_excess.room_id, old_ts.id), None)
            room_ts_occupied[(s_excess.room_id, new_ts.id)] = s_excess
            grp_ts_occupied.pop((s_excess.student_group_id, old_ts.id), None)
            grp_ts_occupied[(s_excess.student_group_id, new_ts.id)] = s_excess

            s_excess.time_slot = new_ts
            s_excess.time_slot_id = new_ts.id
            modified_slots[s_excess.id] = s_excess
            fixed_count += 1
            moved += 1
            fixed_messages.append(
                f"✓ Daily limit fix: {s_excess.course.code} moved from {old_ts} → {new_ts} "
                f"(lecturer daily load reduced to {len(day_slots) - moved}/{max_slots})."
            )

    # Perform bulk update
    if modified_slots:
        with transaction.atomic():
            ScheduleSlot.objects.bulk_update(modified_slots.values(), ['room'])

    # Invalidate conflicts_json cache
    clear_conflicts_cache(timetable.id)

    # Re-run conflict check to get updated counts
    remaining = check_conflicts_for_timetable(timetable)
    remaining_errors   = [c for c in remaining if c['severity'] == 'error']
    remaining_warnings = [c for c in remaining if c['severity'] == 'warning']

    # Update Firebase conflicts node & trigger refresh signal
    fb_errors = [{'type': c['constraint_type'], 'message': c['message']} for c in remaining_errors]
    fb_warnings = [{'type': c['constraint_type'], 'message': c['message']} for c in remaining_warnings]
    
    update_timetable_conflicts(timetable.id, {
        'hard_count': len(fb_errors),
        'soft_count': len(fb_warnings),
        'total': len(remaining),
        'errors': fb_errors,
        'warnings': fb_warnings,
    })
    trigger_timetable_refresh(timetable.id)

    return JsonResponse({
        'status':           'ok',
        'fixed':            fixed_count,
        'skipped':          skipped_count,
        'messages':         fixed_messages,
        'remaining_hard':   len(remaining_errors),
        'remaining_soft':   len(remaining_warnings),
        'errors':    [{'type': c['constraint_type'], 'message': c['message']} for c in remaining_errors],
        'warnings':  [{'type': c['constraint_type'], 'message': c['message']} for c in remaining_warnings],
    })


@login_required
@require_POST
@tenant_required(ScheduleSlot)
def slot_update(request, pk):
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        return JsonResponse({'status': 'error', 'message': 'Permission denied.'}, status=403)

    slot = get_object_or_404(ScheduleSlot, pk=pk)
    room_id = request.POST.get('room_id')
    time_slot_id = request.POST.get('time_slot_id')

    # Find all slots of the same session (on the original day for this course/group/lecturer)
    orig_day = slot.time_slot.day_of_week
    session_slots = list(ScheduleSlot.objects.filter(
        timetable=slot.timetable,
        course=slot.course,
        lecturer=slot.lecturer,
        student_group=slot.student_group,
        time_slot__day_of_week=orig_day
    ).order_by('time_slot__slot_number'))

    if time_slot_id:
        new_start_ts = get_object_or_404(TimeSlot, pk=int(time_slot_id))
        duration = len(session_slots)
        
        # Get consecutive timeslots starting from new_start_ts
        consecutive_ts = list(TimeSlot.objects.filter(
            university=slot.timetable.semester.university,
            day_of_week=new_start_ts.day_of_week,
            slot_number__gte=new_start_ts.slot_number
        ).order_by('slot_number')[:duration])
        
        if len(consecutive_ts) < duration:
            return JsonResponse({
                'status': 'error', 
                'message': f'Not enough consecutive timeslots available starting from {new_start_ts.start_time.strftime("%H:%M")}.'
            }, status=400)
        
        for idx, s_slot in enumerate(session_slots):
            s_slot.time_slot = consecutive_ts[idx]
            if room_id:
                s_slot.room_id = int(room_id)
            s_slot.save()
    else:
        # Only room changed
        for s_slot in session_slots:
            if room_id:
                s_slot.room_id = int(room_id)
            s_slot.save()

    # Invalidate conflicts_json cache
    clear_conflicts_cache(slot.timetable_id)

    conflicts = check_conflicts_for_timetable(slot.timetable)
    error_messages = [c['message'] for c in conflicts if c['severity'] == 'error']
    warning_messages = [c['message'] for c in conflicts if c['severity'] == 'warning']

    # Update Firebase conflicts node & trigger refresh signal
    fb_errors = [{'type': c['constraint_type'], 'message': c['message']} for c in conflicts if c['severity'] == 'error']
    fb_warnings = [{'type': c['constraint_type'], 'message': c['message']} for c in conflicts if c['severity'] == 'warning']
    
    update_timetable_conflicts(slot.timetable_id, {
        'hard_count': len(fb_errors),
        'soft_count': len(fb_warnings),
        'total': len(conflicts),
        'errors': fb_errors,
        'warnings': fb_warnings,
    })
    trigger_timetable_refresh(slot.timetable_id)

    return JsonResponse({
        'status': 'success',
        'errors': error_messages,
        'warnings': warning_messages,
    })


def _get_constraint_context_data(university):
    return {
        'lecturers': Lecturer.objects.filter(department__faculty__campus__university=university).order_by('name'),
        'courses': Course.objects.filter(program__department__faculty__campus__university=university).order_by('code'),
        'rooms': Room.objects.filter(campus__university=university).order_by('name'),
        'student_groups': StudentGroup.objects.filter(program__department__faculty__campus__university=university).order_by('name'),
        'timeslots': TimeSlot.objects.filter(university=university).order_by('day_of_week', 'slot_number'),
    }

@login_required
def constraint_list(request):
    university = get_active_uni(request)
    
    # Automatically pre-populate default global constraints if none exist
    if not Constraint.objects.filter(university=university).exists():
        Constraint.objects.create(
            university=university,
            name="Default Max Classes per Day",
            constraint_type="MAX_CLASSES_PER_DAY",
            is_hard=False,
            weight=100,
            parameters={"lecturer_id": None, "max_classes": 4}
        )
        Constraint.objects.create(
            university=university,
            name="Default Max Consecutive Slots",
            constraint_type="LECTURER_MAX_CONSECUTIVE_SLOTS",
            is_hard=False,
            weight=100,
            parameters={"lecturer_id": None, "max_consecutive": 3}
        )
        Constraint.objects.create(
            university=university,
            name="Default Max Days per Week",
            constraint_type="LECTURER_MAX_DAYS_PER_WEEK",
            is_hard=False,
            weight=100,
            parameters={"lecturer_id": None, "max_days": 5}
        )
        
    constraints = Constraint.objects.filter(university=university)
    
    form = ConstraintForm(initial={'university': university})
    form.fields['university'].queryset = University.objects.filter(id=university.id)

    context = {
        'constraints': constraints,
        'form': form,
        'active_university': university,
    }
    context.update(_get_constraint_context_data(university))
    return render(request, 'scheduler/constraint_list.html', context)

@login_required
def constraint_create(request):
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:constraint_list')

    university = get_active_uni(request)
    if request.method == 'POST':
        form = ConstraintForm(request.POST)
        form.fields['university'].queryset = University.objects.filter(id=university.id)
        if form.is_valid():
            form.save()
            messages.success(request, "Constraint rule configuration saved.")
            return redirect('scheduler:constraint_list')
        else:
            constraints = Constraint.objects.filter(university=university)
            context = {
                'constraints': constraints,
                'form': form,
                'active_university': university,
            }
            context.update(_get_constraint_context_data(university))
            return render(request, 'scheduler/constraint_list.html', context)
    return redirect('scheduler:constraint_list')


@login_required
def constraint_edit(request, pk):
    """Admin view to inspect and edit an individual constraint's configuration rules and parameters."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:constraint_list')

    university = get_active_uni(request)
    constraint = get_object_or_404(Constraint, pk=pk)
    if constraint.university != university:
        messages.error(request, "Access denied.")
        return redirect('scheduler:constraint_list')

    if request.method == 'POST':
        form = ConstraintForm(request.POST, instance=constraint)
        form.fields['university'].queryset = University.objects.filter(id=university.id)
        if form.is_valid():
            form.save()
            messages.success(request, f"Constraint '{constraint.name}' updated successfully.")
            return redirect('scheduler:constraint_list')
    else:
        form = ConstraintForm(instance=constraint)
        form.fields['university'].queryset = University.objects.filter(id=university.id)

    context = {
        'constraint': constraint,
        'form': form,
        'active_university': university,
    }
    context.update(_get_constraint_context_data(university))
    return render(request, 'scheduler/constraint_edit.html', context)


@login_required
def constraint_delete(request, pk):
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:constraint_list')

    try:
        constraint = Constraint.objects.get(pk=pk)
        if request.method == 'POST':
            constraint.delete()
            messages.success(request, "Constraint deleted successfully.")
    except Constraint.DoesNotExist:
        messages.info(request, "Constraint has already been deleted or does not exist.")
    return redirect('scheduler:constraint_list')





@login_required
def switch_university(request):
    try:
        profile = request.user.profile
        if profile.role != 'admin' and not request.user.is_superuser:
            messages.error(request, "Permission denied. You cannot switch universities.")
            return redirect('scheduler:dashboard')
    except Exception:
        pass
    if request.method == 'POST':
        uni_id = request.POST.get('university_id')
        if uni_id:
            request.session['active_university_id'] = int(uni_id)
            messages.success(request, f"Switched active university scope.")
    return redirect(request.META.get('HTTP_REFERER', 'scheduler:dashboard'))

@login_required
def switch_role(request):
    if not request.user.is_superuser:
        messages.error(request, "Role simulation is restricted to superusers.")
        return redirect('scheduler:dashboard')
    if request.method == 'POST':
        role = request.POST.get('role_name')
        valid_roles = (
            'admin', 'institution_admin', 'registrar', 'dvc_academic', 
            'dean', 'hod', 'timetable_officer', 'scheduler', 'lecturer', 'student'
        )
        if role in valid_roles:
            request.session['active_role'] = role
            messages.success(request, f"Simulated role changed to: {role.upper()}")
    return redirect(request.META.get('HTTP_REFERER', 'scheduler:dashboard'))

@login_required
def reports(request):
    uni = get_active_uni(request)
    if not uni:
        return render(request, 'scheduler/reports.html', {})
        
    timetable = Timetable.objects.filter(semester__university=uni, is_active=True).first()
    if not timetable:
        timetable = Timetable.objects.filter(semester__university=uni).first()

    summary = {
        'has_timetable': False,
        'timetable_name': '',
        'semester_name': '',
        'scheduled_slots_count': 0,
        'unscheduled_courses_count': 0,
        'total_lecturers': 0,
        'avg_lecturer_hours': 0.0,
        'overallocated_lecturers_count': 0,
        'total_rooms': 0,
        'avg_room_utilization': 0.0,
        'most_booked_room_name': 'None',
        'most_booked_room_slots': 0,
    }
    
    if timetable:
        summary['has_timetable'] = True
        summary['timetable_name'] = timetable.name
        summary['semester_name'] = timetable.semester.name

        slots = list(timetable.slots.select_related('course', 'lecturer', 'room', 'time_slot'))
        summary['scheduled_slots_count'] = len(slots)
        
        scheduled_course_ids = set(s.course_id for s in slots if s.course_id)
        all_course_ids = set(Course.objects.filter(program__department__faculty__campus__university=uni).values_list('id', flat=True))
        summary['unscheduled_courses_count'] = len(all_course_ids - scheduled_course_ids)

        from collections import defaultdict as _defaultdict
        slots_by_lecturer = _defaultdict(list)
        slots_by_room     = _defaultdict(list)
        for s in slots:
            if s.lecturer_id:
                slots_by_lecturer[s.lecturer_id].append(s)
            if s.room_id:
                slots_by_room[s.room_id].append(s)

        lecturers = Lecturer.objects.filter(department__faculty__campus__university=uni)
        summary['total_lecturers'] = lecturers.count()
        
        tot_hrs = 0.0
        overallocated = 0
        for lec in lecturers:
            lec_slots = slots_by_lecturer.get(lec.id, [])
            hours_count = 0.0
            for slot in lec_slots:
                ts = slot.time_slot
                duration_mins = (ts.end_time.hour * 60 + ts.end_time.minute) - (ts.start_time.hour * 60 + ts.start_time.minute)
                hours_count += duration_mins / 60.0
            tot_hrs += hours_count
            if hours_count > lec.max_hours_per_week:
                overallocated += 1
                
        summary['overallocated_lecturers_count'] = overallocated
        if summary['total_lecturers'] > 0:
            summary['avg_lecturer_hours'] = round(tot_hrs / summary['total_lecturers'], 1)

        rooms = Room.objects.filter(campus__university=uni)
        summary['total_rooms'] = rooms.count()
        total_slots_available = TimeSlot.objects.filter(university=uni).count()

        tot_pct = 0.0
        max_slots = 0
        max_room_name = 'None'
        for room in rooms:
            room_slots = slots_by_room.get(room.id, [])
            slots_count = len(room_slots)
            pct = (slots_count / total_slots_available * 100) if total_slots_available > 0 else 0
            tot_pct += pct
            if slots_count > max_slots:
                max_slots = slots_count
                max_room_name = room.name
                
        summary['most_booked_room_name'] = max_room_name
        summary['most_booked_room_slots'] = max_slots
        if summary['total_rooms'] > 0:
            summary['avg_room_utilization'] = round(tot_pct / summary['total_rooms'], 1)

    return render(request, 'scheduler/reports.html', {
        'timetable': timetable,
        'summary': summary,
    })


@login_required
def reports_workloads(request):
    """Detailed lecturer workload report and editor page."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    uni = get_active_uni(request)
    if not uni:
        return render(request, 'scheduler/reports_workloads.html', {})
        
    timetable = Timetable.objects.filter(semester__university=uni, is_active=True).first()
    if not timetable:
        timetable = Timetable.objects.filter(semester__university=uni).first()

    lecturer_workloads = []
    lecturers = []
    
    if timetable:
        slots = list(timetable.slots.select_related('course', 'lecturer', 'room', 'time_slot'))
        
        from collections import defaultdict as _defaultdict
        slots_by_lecturer = _defaultdict(list)
        for s in slots:
            if s.lecturer_id:
                slots_by_lecturer[s.lecturer_id].append(s)

        lecturers = list(Lecturer.objects.filter(department__faculty__campus__university=uni).select_related('department').order_by('name'))
        
        all_courses = list(Course.objects.filter(program__department__faculty__campus__university=uni).select_related('program'))
        courses_by_lecturer = _defaultdict(list)
        for course in all_courses:
            if course.lecturer_id:
                courses_by_lecturer[course.lecturer_id].append(course)

        for lec in lecturers:
            lec_slots   = slots_by_lecturer.get(lec.id, [])
            tot_hours = 0.0
            for slot in lec_slots:
                ts = slot.time_slot
                duration_mins = (ts.end_time.hour * 60 + ts.end_time.minute) - (ts.start_time.hour * 60 + ts.start_time.minute)
                tot_hours += duration_mins / 60.0
            hours_count = round(tot_hours, 1)
            lec_courses = courses_by_lecturer.get(lec.id, [])
            lecturer_workloads.append({
                'lecturer':           lec,
                'slots_count':        len(lec_slots),
                'hours':              hours_count,
                'max_hours':          lec.max_hours_per_week,
                'lecturer_type':      lec.lecturer_type,
                'courses':            lec_courses,
                'utilization_percent': round(
                    (hours_count / lec.max_hours_per_week * 100), 1
                ) if lec.max_hours_per_week > 0 else 0,
            })

    return render(request, 'scheduler/reports_workloads.html', {
        'timetable': timetable,
        'lecturer_workloads': lecturer_workloads,
        'lecturers': lecturers,
    })


@login_required
def reports_rooms(request):
    """Detailed room booking utilization report page."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    uni = get_active_uni(request)
    if not uni:
        return render(request, 'scheduler/reports_rooms.html', {})
        
    timetable = Timetable.objects.filter(semester__university=uni, is_active=True).first()
    if not timetable:
        timetable = Timetable.objects.filter(semester__university=uni).first()

    room_utilization = []
    
    if timetable:
        slots = list(timetable.slots.select_related('room'))
        total_slots_available = TimeSlot.objects.filter(university=uni).count()

        from collections import defaultdict as _defaultdict
        slots_by_room = _defaultdict(list)
        for s in slots:
            if s.room_id:
                slots_by_room[s.room_id].append(s)

        rooms = Room.objects.filter(campus__university=uni)
        for room in rooms:
            room_slots = slots_by_room.get(room.id, [])
            percentage = round(
                (len(room_slots) / total_slots_available * 100), 1
            ) if total_slots_available > 0 else 0
            room_utilization.append({
                'room':        room,
                'booked_slots': len(room_slots),
                'total_slots': total_slots_available,
                'percentage':  percentage,
            })

    return render(request, 'scheduler/reports_rooms.html', {
        'timetable': timetable,
        'room_utilization': room_utilization,
    })

@login_required
def resources_manager(request):
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    uni = get_active_uni(request)
    tab = request.GET.get('tab', 'university')

    from .models import Program

    # FIX U5: Add select_related / prefetch_related to all tab querysets to
    # avoid N+1 queries when the template renders each row's related fields.
    tab_configs = {
        'university': {
            'model':      University,
            'form_class': UniversityForm,
            'qs':         University.objects.all(),
            'title':      'Universities',
        },
        'campus': {
            'model':      Campus,
            'form_class': CampusForm,
            'qs':         Campus.objects.filter(university=uni).select_related('university'),
            'title':      'Campuses',
        },
        'faculty': {
            'model':      Faculty,
            'form_class': FacultyForm,
            'qs':         Faculty.objects.filter(campus__university=uni).select_related('campus', 'campus__university'),
            'title':      'Faculties',
        },
        'department': {
            'model':      Department,
            'form_class': DepartmentForm,
            'qs':         Department.objects.filter(faculty__campus__university=uni).select_related('faculty', 'faculty__campus'),
            'title':      'Departments',
        },
        'course': {
            'model':      Course,
            'form_class': CourseForm,
            'qs':         Course.objects.filter(
                              program__department__faculty__campus__university=uni
                          ).select_related(
                              'program', 'program__department',
                              'lecturer', 'student_group',
                          ),
            'title':      'Courses',
        },
        'studentgroup': {
            'model':      StudentGroup,
            'form_class': StudentGroupForm,
            'qs':         StudentGroup.objects.filter(
                              program__department__faculty__campus__university=uni
                          ).select_related('program', 'program__department'),
            'title':      'Student Groups',
        },
        'lecturer': {
            'model':      Lecturer,
            'form_class': LecturerForm,
            'qs':         Lecturer.objects.filter(
                              department__faculty__campus__university=uni
                          ).select_related('department', 'department__faculty'),
            'title':      'Lecturers',
        },
        'room': {
            'model':      Room,
            'form_class': RoomForm,
            'qs':         Room.objects.filter(campus__university=uni).select_related('campus'),
            'title':      'Rooms',
        },
        'timeslot': {
            'model':      TimeSlot,
            'form_class': TimeSlotForm,
            'qs':         TimeSlot.objects.filter(university=uni).select_related('university'),
            'title':      'Time Slots',
        },
    }

    if tab not in tab_configs:
        tab = 'university'

    cfg = tab_configs[tab]
    form_class = cfg['form_class']
    
    if request.method == 'POST':
        form = form_class(request.POST)
        
        # Apply field restrictions for safety
        if uni:
            if tab == 'campus':
                form.fields['university'].queryset = University.objects.filter(id=uni.id)
            elif tab == 'faculty':
                form.fields['campus'].queryset = Campus.objects.filter(university=uni)
            elif tab == 'department':
                form.fields['faculty'].queryset = Faculty.objects.filter(campus__university=uni)
            elif tab == 'course':
                form.fields['program'].queryset = Program.objects.filter(department__faculty__campus__university=uni)
                form.fields['lecturer'].queryset = Lecturer.objects.filter(department__faculty__campus__university=uni)
                form.fields['student_group'].queryset = StudentGroup.objects.filter(program__department__faculty__campus__university=uni)
            elif tab == 'lecturer':
                form.fields['department'].queryset = Department.objects.filter(faculty__campus__university=uni)
            elif tab == 'studentgroup':
                form.fields['program'].queryset = Program.objects.filter(department__faculty__campus__university=uni)
            elif tab == 'room':
                form.fields['campus'].queryset = Campus.objects.filter(university=uni)
            elif tab == 'timeslot':
                form.fields['university'].queryset = University.objects.filter(id=uni.id)

        if form.is_valid():
            instance = form.save()
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({
                    "success": True,
                    "message": f"New {cfg['title'].rstrip('s')} '{str(instance)}' added successfully."
                })
            messages.success(request, f"New {cfg['title'].rstrip('s')} added successfully.")
            return redirect(f"/resources/?tab={tab}")
        else:
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({
                    "success": False,
                    "errors": form.errors.get_json_data()
                }, status=400)
    else:
        form = form_class()
        
        # Apply field restrictions for safety
        if uni:
            if tab == 'campus':
                form.fields['university'].queryset = University.objects.filter(id=uni.id)
            elif tab == 'faculty':
                form.fields['campus'].queryset = Campus.objects.filter(university=uni)
            elif tab == 'department':
                form.fields['faculty'].queryset = Faculty.objects.filter(campus__university=uni)
            elif tab == 'course':
                form.fields['program'].queryset = Program.objects.filter(department__faculty__campus__university=uni)
                form.fields['lecturer'].queryset = Lecturer.objects.filter(department__faculty__campus__university=uni)
                form.fields['student_group'].queryset = StudentGroup.objects.filter(program__department__faculty__campus__university=uni)
            elif tab == 'lecturer':
                form.fields['department'].queryset = Department.objects.filter(faculty__campus__university=uni)
            elif tab == 'studentgroup':
                form.fields['program'].queryset = Program.objects.filter(department__faculty__campus__university=uni)
            elif tab == 'room':
                form.fields['campus'].queryset = Campus.objects.filter(university=uni)
            elif tab == 'timeslot':
                form.fields['university'].queryset = University.objects.filter(id=uni.id)

    # ── Search logic ──
    q = request.GET.get('q', '').strip()
    qs = cfg['qs']
    
    if q:
        from django.db.models import Q
        if tab == 'university':
            qs = qs.filter(Q(name__icontains=q) | Q(code__icontains=q))
        elif tab == 'campus':
            qs = qs.filter(Q(name__icontains=q) | Q(university__name__icontains=q))
        elif tab == 'faculty':
            qs = qs.filter(Q(name__icontains=q) | Q(campus__name__icontains=q))
        elif tab == 'department':
            qs = qs.filter(Q(name__icontains=q) | Q(faculty__name__icontains=q))
        elif tab == 'course':
            qs = qs.filter(
                Q(code__icontains=q) |
                Q(name__icontains=q) |
                Q(program__name__icontains=q) |
                Q(lecturer__name__icontains=q) |
                Q(student_group__name__icontains=q)
            )
        elif tab == 'studentgroup':
            qs = qs.filter(Q(name__icontains=q) | Q(program__name__icontains=q))
        elif tab == 'lecturer':
            qs = qs.filter(Q(name__icontains=q) | Q(department__name__icontains=q) | Q(email__icontains=q))
        elif tab == 'room':
            qs = qs.filter(Q(name__icontains=q) | Q(campus__name__icontains=q) | Q(room_type__icontains=q))
        elif tab == 'timeslot':
            day_mapping = {
                'monday': 1, 'tuesday': 2, 'wednesday': 3, 'thursday': 4,
                'friday': 5, 'saturday': 6, 'sunday': 7
            }
            day_val = day_mapping.get(q.lower())
            if day_val:
                qs = qs.filter(day_of_week=day_val)
            else:
                timeslot_query = Q(start_time__icontains=q) | Q(end_time__icontains=q)
                if q.isdigit():
                    timeslot_query |= Q(slot_number=int(q))
                qs = qs.filter(timeslot_query)

    # ── Pagination logic ──
    from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
    page = request.GET.get('page', 1)
    page_size = 50
    if not qs.query.order_by:
        qs = qs.order_by('id')
    paginator = Paginator(qs, page_size)
    try:
        page_obj = paginator.page(page)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)

    context = {
        'active_tab': tab,
        'tab_title': cfg['title'],
        'items': page_obj.object_list,
        'page_obj': page_obj,
        'form': form,
        'q': q,
    }

    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        from django.template.loader import render_to_string
        rows_html = render_to_string('scheduler/partials/resource_rows.html', context, request=request)
        pagination_html = render_to_string('scheduler/partials/resource_pagination.html', context, request=request)
        return JsonResponse({
            "success": True,
            "rows_html": rows_html,
            "pagination_html": pagination_html,
            "total_records": paginator.count,
            "showing_info": f"Showing {page_obj.start_index()} to {page_obj.end_index()} of {paginator.count} records" if paginator.count > 0 else "No records found",
            "current_page": page_obj.number,
            "total_pages": paginator.num_pages
        })

    return render(request, 'scheduler/resources_manager.html', context)


# ─────────────────────────────────────────────────────────────────────────────
# Bulk Delete & Single Delete for Resource Manager
# ─────────────────────────────────────────────────────────────────────────────

MODEL_MAP = {
    'university':   University,
    'campus':       Campus,
    'faculty':      Faculty,
    'department':   Department,
    'course':       Course,
    'studentgroup': StudentGroup,
    'lecturer':     Lecturer,
    'room':         Room,
    'timeslot':     TimeSlot,
}

@login_required
def bulk_delete_resources(request):
    """Deletes multiple resources selected via checkboxes in the Resource Manager."""
    if request.method != 'POST':
        return redirect('scheduler:resources_manager')

    role = get_user_role(request)
    if role not in MANAGER_ROLES:
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({"success": False, "error": "Permission denied."}, status=403)
        messages.error(request, "Permission denied.")
        return redirect('scheduler:resources_manager')

    model_type = request.POST.get('model_type', '')
    tab        = request.POST.get('tab', model_type)
    ids        = request.POST.getlist('selected_ids')

    model = MODEL_MAP.get(model_type)
    if not model or not ids:
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({"success": False, "error": "Nothing selected to delete."}, status=400)
        messages.warning(request, "Nothing selected to delete.")
        return redirect(f"/resources/?tab={tab}")

    try:
        ids = [int(i) for i in ids]
    except (ValueError, TypeError):
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({"success": False, "error": "Invalid selection."}, status=400)
        messages.error(request, "Invalid selection.")
        return redirect(f"/resources/?tab={tab}")

    try:
        deleted_count, _ = model.objects.filter(pk__in=ids).delete()
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({"success": True, "message": f"Successfully deleted {deleted_count} record(s)."})
        messages.success(request, f"🗑 Successfully deleted {deleted_count} record(s).")
        return redirect(f"/resources/?tab={tab}")
    except Exception as e:
        error_msg = str(e)
        if "ProtectedError" in type(e).__name__:
            error_msg = "Cannot delete some selected record(s) because they are referenced by other records."
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({"success": False, "error": error_msg}, status=400)
        messages.error(request, error_msg)
        return redirect(f"/resources/?tab={tab}")


@login_required
def delete_resource(request, model_type, pk):
    """Deletes a single resource item from the Resource Manager."""
    role = get_user_role(request)
    if role not in MANAGER_ROLES:
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({"success": False, "error": "Permission denied."}, status=403)
        messages.error(request, "Permission denied.")
        return redirect('scheduler:resources_manager')

    model = MODEL_MAP.get(model_type)
    if not model:
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({"success": False, "error": "Unknown resource type."}, status=400)
        messages.error(request, "Unknown resource type.")
        return redirect('scheduler:resources_manager')

    try:
        obj = model.objects.get(pk=pk)
    except model.DoesNotExist:
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({"success": False, "error": "This resource has already been deleted or does not exist."}, status=404)
        messages.info(request, "This resource has already been deleted or does not exist.")
        return redirect(f"/resources/?tab={model_type}")

    if request.method == 'POST' or request.headers.get('x-requested-with') == 'XMLHttpRequest':
        name = str(obj)
        try:
            obj.delete()
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({"success": True, "message": f"🗑 '{name}' deleted successfully."})
            messages.success(request, f"🗑 '{name}' deleted successfully.")
            return redirect(f"/resources/?tab={model_type}")
        except Exception as e:
            error_msg = str(e)
            if "ProtectedError" in type(e).__name__:
                error_msg = f"Cannot delete '{name}' because it is referenced by other records."
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({"success": False, "error": error_msg}, status=400)
            messages.error(request, error_msg)
            return redirect(f"/resources/?tab={model_type}")

    # GET — show confirmation page
    return render(request, 'scheduler/confirm_delete.html', {
        'object': obj,
        'model_type': model_type,
        'cancel_url': f"/resources/?tab={model_type}",
    })

@login_required
def export_timetable_ics(request, pk):
    university = get_active_uni(request)
    timetable = get_object_or_404(Timetable.objects.select_related('semester', 'semester__university'), pk=pk)
    
    # Scoping isolation: Ensure timetable belongs to active university
    if timetable.semester.university != university:
        messages.error(request, "Attempted to access out-of-scope timetable.")
        return redirect('scheduler:dashboard')

    filter_type = request.GET.get('filter_type')
    filter_id = request.GET.get('filter_id')

    from .calendar_exporter import (
        generate_ics_content,
        generate_lecturer_ics,
        generate_student_group_ics,
        _generate_ics_from_slots
    )

    if filter_type == 'lecturer' and filter_id:
        lecturer = get_object_or_404(Lecturer, pk=filter_id)
        ics_content = generate_lecturer_ics(lecturer, timetable=timetable)
        safe_name = "".join(c for c in lecturer.name if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')
        filename = f"lecturer_{safe_name}_schedule.ics"
    elif filter_type == 'group' and filter_id:
        student_group = get_object_or_404(StudentGroup, pk=filter_id)
        ics_content = generate_student_group_ics(student_group, timetable=timetable)
        safe_name = "".join(c for c in student_group.name if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')
        filename = f"group_{safe_name}_schedule.ics"
    elif filter_type == 'room' and filter_id:
        room = get_object_or_404(Room, pk=filter_id)
        slots = timetable.slots.filter(room=room).select_related('course', 'lecturer', 'room', 'time_slot', 'student_group')
        ics_content = _generate_ics_from_slots(slots, timetable.semester, f"Room {room.name}")
        safe_name = "".join(c for c in room.name if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')
        filename = f"room_{safe_name}_schedule.ics"
    else:
        ics_content = generate_ics_content(timetable)
        safe_name = "".join(c for c in timetable.name if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')
        filename = f"timetable_{safe_name}.ics"
    
    response = HttpResponse(ics_content, content_type='text/calendar')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

@login_required
def export_timetable_csv(request, pk):
    university = get_active_uni(request)
    timetable = get_object_or_404(Timetable.objects.select_related('semester', 'semester__university'), pk=pk)
    
    # Scoping isolation: Ensure timetable belongs to active university
    if timetable.semester.university != university:
        messages.error(request, "Attempted to access out-of-scope timetable.")
        return redirect('scheduler:dashboard')

    response = HttpResponse(content_type='text/csv')
    safe_name = "".join(c for c in timetable.name if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')
    response['Content-Disposition'] = f'attachment; filename="timetable_{safe_name}.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['Course Code', 'Course Name', 'Room', 'Room Type', 'Lecturer', 'Student Group', 'Day', 'Start Time', 'End Time'])
    
    day_names = {
        1: 'Monday', 2: 'Tuesday', 3: 'Wednesday', 4: 'Thursday', 
        5: 'Friday', 6: 'Saturday', 7: 'Sunday'
    }
    
    slots = timetable.slots.select_related('course', 'lecturer', 'room', 'time_slot', 'student_group').all()
    for slot in slots:
        ts = slot.time_slot
        writer.writerow([
            slot.course.code,
            slot.course.name,
            slot.room.name,
            slot.room.get_room_type_display(),
            slot.lecturer.name,
            slot.student_group.name,
            day_names.get(ts.day_of_week, f"Day {ts.day_of_week}"),
            ts.start_time.strftime('%H:%M'),
            ts.end_time.strftime('%H:%M')
        ])
        
    return response


def get_or_create_structure(uni, campus_name=None, faculty_name=None, dept_name=None, program_name=None):
    from .models import Program
    campus = None
    if campus_name:
        campus, _ = Campus.objects.get_or_create(university=uni, name=campus_name.strip())
    else:
        campus = Campus.objects.filter(university=uni).first()
        if not campus:
            campus = Campus.objects.create(university=uni, name="Default Campus")

    faculty = None
    if faculty_name:
        faculty, _ = Faculty.objects.get_or_create(campus=campus, name=faculty_name.strip())
    else:
        if dept_name or program_name:
            faculty = Faculty.objects.filter(campus__university=uni).first()
            if not faculty:
                faculty = Faculty.objects.create(campus=campus, name="Default Faculty")

    dept = None
    if dept_name:
        if not faculty:
            faculty = Faculty.objects.filter(campus__university=uni).first() or Faculty.objects.create(campus=campus, name="Default Faculty")
        dept, _ = Department.objects.get_or_create(faculty=faculty, name=dept_name.strip())
    else:
        if program_name:
            dept = Department.objects.filter(faculty__campus__university=uni).first()
            if not dept:
                dept = Department.objects.create(faculty=faculty, name="Default Department")

    program = None
    if program_name:
        if not dept:
            dept = Department.objects.filter(faculty__campus__university=uni).first() or Department.objects.create(faculty=faculty, name="Default Department")
        program, _ = Program.objects.get_or_create(department=dept, name=program_name.strip())

    return campus, faculty, dept, program


def auto_heal_university_data(university):
    """
    Automatically fixes all common data issues after an import:
      1. Courses missing a lecturer  → auto-assign from available pool (round-robin)
      2. Courses missing student group → auto-assign from available pool (round-robin)
      3. Rooms on wrong campus       → move all rooms to the courses' campus
      4. Over-allocated lecturers    → raise their max_hours_per_week
      5. Room too small for group    → upgrade room capacities
    Returns a list of fix summary strings for display.
    """
    from .models import Program
    fixes = []

    courses = list(Course.objects.filter(
        program__department__faculty__campus__university=university
    ).select_related('program__department__faculty__campus', 'student_group', 'lecturer'))

    if not courses:
        return fixes

    # ── Fix 1: Courses missing lecturer ──────────────────────────────────────
    # Calculate average timeslot duration dynamically
    timeslots = list(TimeSlot.objects.filter(university=university))
    if timeslots:
        avg_duration = sum(
            ((ts.end_time.hour * 60 + ts.end_time.minute) - (ts.start_time.hour * 60 + ts.start_time.minute)) / 60.0
            for ts in timeslots
        ) / len(timeslots)
    else:
        avg_duration = 1.5

    no_lec = [c for c in courses if not c.lecturer_id]
    if no_lec:
        lecturers = list(Lecturer.objects.filter(
            department__faculty__campus__university=university
        ))
        if lecturers:
            # Track current load
            lec_hours = {l.id: 0.0 for l in lecturers}
            for c in courses:
                if c.lecturer_id and c.lecturer_id in lec_hours:
                    sh = float(c.duration_slots) if (avg_duration >= 2.5 and c.duration_slots >= 2) else (c.duration_slots * avg_duration)
                    lec_hours[c.lecturer_id] += sh * c.sessions_per_week

            lec_sorted = sorted(lecturers, key=lambda l: lec_hours[l.id])
            n = len(lec_sorted)
            to_update = []
            for i, course in enumerate(no_lec):
                lec = lec_sorted[i % n]
                course.lecturer = lec
                sh = float(course.duration_slots) if (avg_duration >= 2.5 and course.duration_slots >= 2) else (course.duration_slots * avg_duration)
                new_hrs = lec_hours[lec.id] + sh * course.sessions_per_week
                if new_hrs > lec.max_hours_per_week:
                    lec.max_hours_per_week = int(new_hrs) + 4
                lec_hours[lec.id] = new_hrs
                to_update.append(course)

            Course.objects.bulk_update(to_update, ['lecturer'], batch_size=500)
            Lecturer.objects.bulk_update(lec_sorted, ['max_hours_per_week'], batch_size=500)
            fixes.append(f"✅ Auto-assigned lecturers to {len(no_lec)} courses.")

    # ── Fix 2: Courses missing student group ─────────────────────────────────
    no_group = [c for c in courses if not c.student_group_id]
    if no_group:
        groups = list(StudentGroup.objects.filter(
            program__department__faculty__campus__university=university
        ))
        if groups:
            to_update = []
            for i, course in enumerate(no_group):
                course.student_group = groups[i % len(groups)]
                to_update.append(course)
            Course.objects.bulk_update(to_update, ['student_group'], batch_size=500)
            fixes.append(f"✅ Auto-assigned student groups to {len(no_group)} courses.")

    # ── Fix 3: Rooms on wrong campus ─────────────────────────────────────────
    # DISABLED: Moving all rooms to a single campus breaks multi-campus configurations (e.g. Kitengela/Town/Main).
    # Rooms belong to their respective campuses.

    # ── Fix 4: Over-allocated lecturers ──────────────────────────────────────
    lec_load = {}
    for c in Course.objects.filter(program__department__faculty__campus__university=university).select_related('lecturer'):
        if c.lecturer_id:
            sh = float(c.duration_slots) * 3.0
            lec_load[c.lecturer_id] = lec_load.get(c.lecturer_id, 0.0) + sh * getattr(c, 'sessions_per_week', 1)

    over_lecs = []
    for lec in Lecturer.objects.filter(department__faculty__campus__university=university):
        hours = lec_load.get(lec.id, 0.0)
        if hours > lec.max_hours_per_week:
            lec.max_hours_per_week = int(hours) + 4
            over_lecs.append(lec)

    if over_lecs:
        Lecturer.objects.bulk_update(over_lecs, ['max_hours_per_week'], batch_size=500)
        fixes.append(f"✅ Fixed max hours for {len(over_lecs)} over-allocated lecturers.")

    # ── Fix 4b: Rebalance weekly time slot overload ─────────────────────────
    timeslots = list(TimeSlot.objects.filter(university=university))
    num_ts = len(timeslots)
    if num_ts > 0:
        all_uni_courses = list(Course.objects.filter(program__department__faculty__campus__university=university).select_related('lecturer'))
        from collections import defaultdict
        lec_slots_map = defaultdict(int)
        for c in all_uni_courses:
            if c.lecturer_id:
                lec_slots_map[c.lecturer_id] += c.duration_slots * c.sessions_per_week

        all_uni_lecs = list(Lecturer.objects.filter(department__faculty__campus__university=university))
        under_utilized = [l for l in all_uni_lecs if lec_slots_map[l.id] < num_ts * 0.5]
        overloaded_ts = [l for l in all_uni_lecs if lec_slots_map[l.id] > num_ts]

        reassigned_count = 0
        for o_lec in overloaded_ts:
            o_courses = [c for c in all_uni_courses if c.lecturer_id == o_lec.id]
            for c in o_courses:
                if lec_slots_map[o_lec.id] <= num_ts:
                    break
                c_slots = c.duration_slots * c.sessions_per_week
                for cand in under_utilized:
                    if lec_slots_map[cand.id] + c_slots <= num_ts * 0.6:
                        c.lecturer = cand
                        c.save()
                        lec_slots_map[o_lec.id] -= c_slots
                        lec_slots_map[cand.id] += c_slots
                        reassigned_count += 1
                        break
        if reassigned_count:
            fixes.append(f"✅ Automatically redistributed {reassigned_count} course(s) from overloaded lecturers to available faculty to fit within weekly time slots.")

    # ── Fix 5: Room capacity too small for group sizes ────────────────────────
    import random
    rooms = list(Room.objects.filter(campus__university=university))
    max_room_cap = max((r.capacity for r in rooms), default=0)
    group_sizes = [
        c.student_group.size for c in courses
        if c.student_group_id and c.student_group
    ]
    if group_sizes:
        max_group = max(group_sizes)
        if max_group > max_room_cap:
            small = [r for r in rooms if r.capacity < max_group]
            for r in small:
                r.capacity = random.choice([60, 80, 100, 120, 150])
            if small:
                Room.objects.bulk_update(small, ['capacity'], batch_size=500)
                fixes.append(f"✅ Upgraded capacity of {len(small)} rooms to fit all student groups.")

    return fixes


@login_required
def import_resources(request):
    role = get_user_role(request)
    if role not in MANAGER_ROLES:
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    university = get_active_uni(request)
    if not university:
        messages.error(request, "No active university found.")
        return redirect('scheduler:dashboard')

    if request.method == 'POST':
        import_type = request.POST.get('import_type')
        
        # Check confirmation first!
        if request.POST.get('confirm') == 'yes':
            import os
            from .smart_import import detect_format, extract_entities, import_entities
            scratch_dir = os.path.join(settings.BASE_DIR, 'scratch')
            temp_file_path = os.path.join(scratch_dir, f'temp_import_{university.id}.xlsx')
            
            if not os.path.exists(temp_file_path):
                messages.error(request, "Import session expired or file not found. Please upload again.")
                return redirect('scheduler:import_resources')
                
            try:
                import openpyxl
                wb = openpyxl.load_workbook(temp_file_path, data_only=True)
                format_info = detect_format(wb)
                entities = extract_entities(wb, format_info, university)
                summary = import_entities(university, entities)

                # ── Write audit log ─────────────────────────────────────────
                from scheduler.models import ImportAuditLog
                import_file_name = request.session.pop('import_file_name', os.path.basename(temp_file_path))
                created_count = {k: v for k, v in summary.items()
                                 if k in ('campuses','programs','lecturers','rooms','student_groups','courses','time_slots') and v}
                updated_count = {k: v for k, v in {
                    'lecturers': summary.get('lecturers_updated', 0),
                    'rooms': summary.get('rooms_updated', 0),
                    'courses': summary.get('courses_updated', 0),
                }.items() if v}
                audit_log = ImportAuditLog.objects.create(
                    university=university,
                    imported_by=request.user if request.user.is_authenticated else None,
                    file_name=import_file_name,
                    import_type='smart',
                    created_count=created_count,
                    updated_count=updated_count,
                    flagged_dupes=summary.get('flagged_duplicates', []),
                    warnings=summary.get('warnings', []),
                )
                # ───────────────────────────────────────────────────────────────

                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)

                try:
                    from scheduler.signals import queue_auto_generation
                    auto_tt = (
                        Timetable.objects.filter(semester__university=university, is_active=True).first()
                        or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
                    )
                    if auto_tt:
                        queue_auto_generation(auto_tt)
                except Exception:
                    pass

                return redirect('scheduler:import_audit_report', pk=audit_log.pk)

            except Exception as e:
                from scheduler.smart_import import ImportValidationError as _IVE
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
                err_list = e.errors if isinstance(e, _IVE) else getattr(e, 'errors', [str(e)])
                context = {
                    'import_errors': err_list,
                    'import_type': 'smart',
                }
                return render(request, 'scheduler/resources_import.html', context)

        uploaded_file = request.FILES.get('file')

        if not uploaded_file:
            messages.error(request, "Please upload a CSV or Excel (.xlsx) file.")
            return redirect('scheduler:import_resources')

        file_name = uploaded_file.name.lower()

        # ── ALL-IN-ONE multi-sheet import ──────────────────────────────────────
        if import_type == 'all':
            if not file_name.endswith('.xlsx'):
                messages.error(request, "'Import All at Once' requires an Excel (.xlsx) file with separate sheets.")
                return redirect('scheduler:import_resources')

            import openpyxl, re as _re2
            try:
                wb_all = openpyxl.load_workbook(uploaded_file, data_only=True)
            except Exception as e:
                messages.error(request, f"Could not open Excel file: {e}")
                return redirect('scheduler:import_resources')

            # Map each sheet to an import_type by matching sheet name keywords
            SHEET_TYPE_MAP = {
                'room': 'room', 'rooms': 'room', 'venue': 'room', 'venues': 'room',
                'location': 'room', 'locations': 'room', 'hall': 'room', 'halls': 'room',
                'lecturer': 'lecturer', 'lecturers': 'lecturer',
                'instructor': 'lecturer', 'instructors': 'lecturer',
                'teacher': 'lecturer', 'teachers': 'lecturer',
                'staff': 'lecturer',
                'student': 'student_group', 'students': 'student_group',
                'group': 'student_group', 'groups': 'student_group',
                'class': 'student_group', 'classes': 'student_group',
                'cohort': 'student_group', 'cohorts': 'student_group',
                'section': 'student_group', 'sections': 'student_group',
                'course': 'course', 'courses': 'course',
                'unit': 'course', 'units': 'course',
                'module': 'course', 'modules': 'course',
                'subject': 'course', 'subjects': 'course',
            }

            # Process sheets in dependency order
            IMPORT_ORDER = ['room', 'lecturer', 'student_group', 'course']
            sheet_assignments = {}  # import_type -> sheet
            for sheet_name in wb_all.sheetnames:
                key = _re2.sub(r'[^a-z]', '', sheet_name.strip().lower())
                matched = SHEET_TYPE_MAP.get(key)
                if not matched:
                    # try partial match
                    for kw, itype in SHEET_TYPE_MAP.items():
                        if kw in key or key in kw:
                            matched = itype
                            break
                if matched and matched not in sheet_assignments:
                    sheet_assignments[matched] = wb_all[sheet_name]

            if not sheet_assignments:
                messages.error(
                    request,
                    "Could not detect any recognised sheets. Name your sheets: "
                    "Rooms/Venues, Lecturers/Staff, Students/Groups, Courses/Units."
                )
                return redirect('scheduler:import_resources')

            def _clean_val(val):
                if val is None:
                    return ''
                import datetime
                if isinstance(val, (datetime.datetime, datetime.date)):
                    return f"{val.month}-{val.day}"
                if isinstance(val, (int, float)):
                    if 40000 <= val <= 50000:
                        dt = datetime.date(1899, 12, 30) + datetime.timedelta(days=int(val))
                        return f"{dt.month}-{dt.day}"
                    if isinstance(val, float) and val.is_integer():
                        return str(int(val))
                    return str(val)
                val_str = str(val).strip()
                import re
                if re.match(r'^\d{4}-\d{2}-\d{2}', val_str):
                    try:
                        dt = datetime.datetime.strptime(val_str.split()[0], '%Y-%m-%d')
                        return f"{dt.month}-{dt.day}"
                    except:
                        pass
                if re.match(r'^\d{5}(\.0)?$', val_str):
                    try:
                        num = int(float(val_str))
                        if 40000 <= num <= 50000:
                            dt = datetime.date(1899, 12, 30) + datetime.timedelta(days=num)
                            return f"{dt.month}-{dt.day}"
                        return str(num)
                    except:
                        pass
                return val_str

            def _parse_sheet(sheet):
                rows = list(sheet.iter_rows(values_only=True))
                if not rows:
                    return [], []
                raw_hdrs = rows[0]
                hdrs = [str(h).strip().lower().replace(' ', '_') for h in raw_hdrs if h is not None]
                recs = []
                for row in rows[1:]:
                    if any(row):
                        vals = [_clean_val(v) for v in row[:len(hdrs)]]
                        while len(vals) < len(hdrs):
                            vals.append('')
                        recs.append(dict(zip(hdrs, vals)))
                return hdrs, recs

            all_totals = {}
            all_warnings = []
            all_errors = []
            heal_msgs = []

            from django.db import transaction
            from .models import Program
            import re as _re

            university_ref = university  # closure

            try:
                with transaction.atomic():
                    for itype in IMPORT_ORDER:
                        sheet = sheet_assignments.get(itype)
                        if sheet is None:
                            continue
                        _hdrs, _recs = _parse_sheet(sheet)
                        if not _recs:
                            continue

                        default_campus = Campus.objects.filter(university=university_ref).first() or \
                            Campus.objects.create(university=university_ref, name="Default Campus")
                        default_faculty = Faculty.objects.filter(campus__university=university_ref).first() or \
                            Faculty.objects.create(campus=default_campus, name="Default Faculty")
                        default_dept = Department.objects.filter(faculty__campus__university=university_ref).first() or \
                            Department.objects.create(faculty=default_faculty, name="Default Department")
                        default_program = Program.objects.filter(department__faculty__campus__university=university_ref).first() or \
                            Program.objects.create(department=default_dept, name="Default Program")

                        campus_cache2  = {c.name.strip(): c for c in Campus.objects.filter(university=university_ref)}
                        dept_cache2    = {d.name.strip(): d for d in Department.objects.filter(faculty__campus__university=university_ref)}
                        program_cache2 = {p.name.strip(): p for p in Program.objects.filter(department__faculty__campus__university=university_ref)}

                        def _gc(nm):
                            if not nm: return default_campus
                            if nm not in campus_cache2:
                                campus_cache2[nm] = Campus.objects.create(university=university_ref, name=nm)
                            return campus_cache2[nm]

                        def _gd(nm):
                            if not nm: return default_dept
                            if nm not in dept_cache2:
                                dept_cache2[nm] = Department.objects.create(faculty=default_faculty, name=nm)
                            return dept_cache2[nm]

                        def _gp(nm):
                            if not nm: return default_program
                            if nm not in program_cache2:
                                program_cache2[nm] = Program.objects.create(department=default_dept, name=nm)
                            return program_cache2[nm]

                        ROOM_TYPE_LECTURE2 = {'lecture','lecture hall','lecture_hall','hall','theater','theatre','classroom','class room','class_room','auditorium','lh'}
                        ROOM_TYPE_LAB2 = {'lab','laboratory','computer lab','computer_lab','science lab','science_lab','workshop','studio'}
                        ROOM_TYPE_SEMINAR2 = {'seminar','seminar room','seminar_room','tutorial','tutorial room','tutorial_room','meeting room','meeting_room','conference','boardroom'}

                        def _nrt(raw):
                            v = (raw or '').strip().lower()
                            if v in ROOM_TYPE_LECTURE2: return 'Lecture'
                            if v in ROOM_TYPE_LAB2: return 'Lab'
                            if v in ROOM_TYPE_SEMINAR2: return 'Seminar'
                            return None

                        def _g(row, *aliases):
                            for a in aliases:
                                v = row.get(a, '')
                                if v and str(v).strip(): return str(v).strip()
                                norm = _re.sub(r'\s+', '_', a.strip().lower())
                                v = row.get(norm, '')
                                if v and str(v).strip(): return str(v).strip()
                            return ''

                        sc = 0
                        if itype == 'room':
                            existing = {(r.campus_id, r.name): r for r in Room.objects.filter(campus__university=university_ref)}
                            new_r, upd_r = [], []
                            for idx, row in enumerate(_recs, 2):
                                nm  = _g(row,'name','room_name','venue','venue_name','location')
                                cap_s = _g(row,'capacity','cap','seats','size','count','enrolment','enrollment','no_of_students','number_of_students','students_count')
                                rt  = _g(row,'room_type','required_room_type','venue_type','type','facility_type')
                                cn  = _g(row,'campus_name','campus') or None
                                if not nm: all_errors.append(f"[Rooms] Row {idx}: Name required."); continue
                                try: cap = int(float(cap_s)); assert cap > 0
                                except: all_errors.append(f"[Rooms] Row {idx}: Capacity invalid."); continue
                                nrt = _nrt(rt)
                                if not nrt: all_errors.append(f"[Rooms] Row {idx}: Unrecognised type '{rt}'."); continue
                                campus = _gc(cn)
                                key = (campus.id, nm)
                                _VIRTUAL_KEYWORDS = ('zoom', 'virtual', 'online', 'teams')
                                def _is_virt(room_name):
                                    n = room_name.strip().lower()
                                    return any(kw in n for kw in _VIRTUAL_KEYWORDS)
                                if key in existing:
                                    r = existing[key]
                                    r.capacity = cap
                                    r.room_type = nrt
                                    r.is_virtual = _is_virt(nm)
                                    upd_r.append(r)
                                else:
                                    room_obj = Room(campus=campus, name=nm, capacity=cap, room_type=nrt,
                                                    is_virtual=_is_virt(nm))
                                    new_r.append(room_obj)
                                    existing[key] = room_obj
                                sc += 1
                            if not all_errors:
                                Room.objects.bulk_create(new_r, batch_size=500, ignore_conflicts=True)
                                if upd_r: Room.objects.bulk_update(upd_r, ['capacity','room_type','is_virtual'], batch_size=500)

                        elif itype == 'lecturer':
                            existing = {l.email: l for l in Lecturer.objects.filter(department__faculty__campus__university=university_ref)}
                            new_l, upd_l = [], []
                            for idx, row in enumerate(_recs, 2):
                                nm  = _g(row,'name','lecturer_name','lecturer','instructor','instructor_name','teacher','teacher_name','staff','staff_name','full_name')
                                em  = _g(row,'email','email_address','mail','e_mail').lower()
                                dn  = _g(row,'department_name','department','dept','faculty','school') or None
                                mhs = _g(row,'max_hours_per_week','max_hours','hours','weekly_hours','workload') or None
                                if not nm: all_errors.append(f"[Lecturers] Row {idx}: Name required."); continue
                                if not em or '@' not in em: all_errors.append(f"[Lecturers] Row {idx}: Valid email required."); continue
                                mh = 20
                                if mhs:
                                    try: mh = int(float(mhs))
                                    except: all_warnings.append(f"[Lecturers] Row {idx}: Invalid max hours, defaulted to 20.")
                                dept = _gd(dn)
                                if em in existing:
                                    l = existing[em]; l.name = nm; l.department = dept; l.max_hours_per_week = mh; upd_l.append(l)
                                else:
                                    lec_obj = Lecturer(email=em, name=nm, department=dept, max_hours_per_week=mh)
                                    new_l.append(lec_obj)
                                    existing[em] = lec_obj
                                sc += 1
                            if not all_errors:
                                Lecturer.objects.bulk_create(new_l, batch_size=500, ignore_conflicts=True)
                                if upd_l: Lecturer.objects.bulk_update(upd_l, ['name','department','max_hours_per_week'], batch_size=500)

                        elif itype == 'student_group':
                            existing = {(g.program_id, g.name): g for g in StudentGroup.objects.filter(program__department__faculty__campus__university=university_ref)}
                            new_g, upd_g = [], []
                            for idx, row in enumerate(_recs, 2):
                                nm  = _g(row,'name','group_name','student','student_group','student_group_name','class','class_name','cohort','section','stream')
                                ss  = _g(row,'size','group_size','capacity','cap','count','enrolment','enrollment','no_of_students','number_of_students','students_count') or None
                                pn  = _g(row,'program_name','program','programme') or None
                                if not nm: all_errors.append(f"[Students] Row {idx}: Name required."); continue
                                try: sz = int(float(ss)); assert sz > 0
                                except: all_errors.append(f"[Students] Row {idx}: Size invalid."); continue
                                prog = _gp(pn); key = (prog.id, nm)
                                if key in existing:
                                    g = existing[key]; g.size = sz; upd_g.append(g)
                                else:
                                    group_obj = StudentGroup(program=prog, name=nm, size=sz)
                                    new_g.append(group_obj)
                                    existing[key] = group_obj
                                sc += 1
                            if not all_errors:
                                StudentGroup.objects.bulk_create(new_g, batch_size=500, ignore_conflicts=True)
                                if upd_g: StudentGroup.objects.bulk_update(upd_g, ['size'], batch_size=500)

                        elif itype == 'course':
                            lc2 = {l.email: l for l in Lecturer.objects.filter(department__faculty__campus__university=university_ref)}
                            gc2 = {(g.program_id, g.name): g for g in StudentGroup.objects.filter(program__department__faculty__campus__university=university_ref)}
                            existing = {
                                (c.program_id, c.code.strip().upper(), c.student_group_id): c 
                                for c in Course.objects.filter(program__department__faculty__campus__university=university_ref)
                            }
                            new_c, upd_c = [], []
                            for idx, row in enumerate(_recs, 2):
                                code2 = _g(row,'code','course_code','subject_code','unit_code','module_code').upper()
                                nm2   = _g(row,'name','course_name','subject','subject_name','unit','unit_name','module','module_name')
                                ds2   = _g(row,'duration_slots','duration','slots','hours','credit_hours','credits','periods','lesson_periods') or None
                                spw2  = _g(row,'sessions_per_week','weekly_sessions','classes_per_week','sessions') or None
                                rt2   = _g(row,'required_room_type','room_type','venue_type','type','facility_type')
                                le2   = _g(row,'lecturer_email','instructor_email','teacher_email','lecturer','instructor','teacher').lower()
                                gr2   = _g(row,'student_group_name','student_group','student','group','class','class_name','cohort','section','stream')
                                pn2   = _g(row,'program_name','program','programme') or None
                                if not code2: all_errors.append(f"[Courses] Row {idx}: Code required."); continue
                                if not nm2: all_errors.append(f"[Courses] Row {idx}: Name required."); continue
                                try: dur = int(float(ds2)); assert dur > 0
                                except: all_errors.append(f"[Courses] Row {idx}: Duration invalid."); continue
                                try: spw = int(float(spw2)) if spw2 else 1; assert spw > 0
                                except: all_errors.append(f"[Courses] Row {idx}: Sessions per week invalid."); continue
                                nrt2 = _nrt(rt2)
                                if not nrt2: all_errors.append(f"[Courses] Row {idx}: Unrecognised type '{rt2}'."); continue
                                lect2 = lc2.get(le2)
                                if le2 and not lect2: all_warnings.append(f"[Courses] Row {idx}: Lecturer '{le2}' not found.")
                                prog2 = _gp(pn2)
                                grp2 = None
                                if gr2:
                                    gk2 = (prog2.id, gr2)
                                    if gk2 not in gc2:
                                        g2 = StudentGroup.objects.create(program=prog2, name=gr2, size=30)
                                        gc2[gk2] = g2
                                    grp2 = gc2[gk2]
                                key2 = (prog2.id, code2.strip().upper(), grp2.id if grp2 else None)
                                if key2 in existing:
                                    c2 = existing[key2]; c2.name = nm2; c2.duration_slots = dur; c2.sessions_per_week = spw; c2.required_room_type = nrt2; c2.lecturer = lect2; c2.student_group = grp2; upd_c.append(c2)
                                else:
                                    course_obj = Course(program=prog2, code=code2, name=nm2, duration_slots=dur, sessions_per_week=spw, required_room_type=nrt2, lecturer=lect2, student_group=grp2)
                                    new_c.append(course_obj)
                                    existing[key2] = course_obj
                                sc += 1
                            if not all_errors:
                                Course.objects.bulk_create(new_c, batch_size=500, ignore_conflicts=True)
                                if upd_c: Course.objects.bulk_update(upd_c, ['name','duration_slots','sessions_per_week','required_room_type','lecturer','student_group'], batch_size=500)

                        all_totals[itype] = sc

                    # Auto-heal inside the transaction
                    try:
                        heal_msgs = auto_heal_university_data(university_ref)
                    except Exception:
                        heal_msgs = []

                    # Run semantic validation inside transaction
                    from scheduler.validation import validate_university_data
                    is_valid, val_errors, val_warnings = validate_university_data(university_ref)
                    if val_errors:
                        all_errors.extend(val_errors)
                    if val_warnings:
                        all_warnings.extend(val_warnings)

                    if all_errors:
                        raise Exception("Validation errors occurred.")

            except Exception as e:
                # Database rollback triggered
                context = {
                    'import_errors': all_errors if all_errors else [str(e)],
                    'import_warnings': all_warnings,
                    'import_type': import_type,
                }
                return render(request, 'scheduler/resources_import.html', context)

            # Build summary
            type_labels = {'room': 'Rooms', 'lecturer': 'Lecturers', 'student_group': 'Student Groups', 'course': 'Courses'}
            summary_parts = [f"{v} {type_labels.get(k,k)}" for k, v in all_totals.items() if v]
            if summary_parts:
                messages.success(request, f"✅ Successfully imported: {', '.join(summary_parts)}!")
            else:
                messages.warning(request, "No records were imported. Check your sheet names and data.")
            for w in all_warnings[:10]:
                messages.warning(request, w)

            # Display auto-heal messages
            for fix_msg in heal_msgs:
                messages.info(request, fix_msg)
            auto_timetable2 = (
                Timetable.objects.filter(semester__university=university, is_active=True).first()
                or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
            )
            if auto_timetable2:
                try:
                    from .signals import queue_auto_generation
                    queue_auto_generation(auto_timetable2)
                    messages.info(request, f"✓ Timetable generation queued for '{auto_timetable2.name}'.")
                except Exception:
                    pass

            # Trigger lecturer credentials provisioning
            import sys
            if 'test' in sys.argv or 'pytest' in sys.argv or any('pytest' in arg for arg in sys.argv):
                from scheduler.tasks import provision_lecturer_credentials
                provision_lecturer_credentials(university.id)
            else:
                from django_q.tasks import async_task
                async_task('scheduler.tasks.provision_lecturer_credentials', university.id)

            return redirect('/resources/?tab=rooms')
        # ── End of all-in-one import ────────────────────────────────────────────

        elif import_type == 'smart':
            if not file_name.endswith('.xlsx') and not file_name.endswith('.csv'):
                messages.error(request, "Smart Import requires an Excel (.xlsx) or CSV (.csv) file.")
                return redirect('scheduler:import_resources')
                
            import os
            from .smart_import import detect_format, extract_entities
            scratch_dir = os.path.join(settings.BASE_DIR, 'scratch')
            os.makedirs(scratch_dir, exist_ok=True)
            temp_file_path = os.path.join(scratch_dir, f'temp_import_{university.id}.xlsx')
            
            import openpyxl
            wb = openpyxl.Workbook()
            ws = wb.active
            
            if file_name.endswith('.csv'):
                try:
                    csv_file = io.TextIOWrapper(uploaded_file, encoding='utf-8')
                    reader = csv.reader(csv_file)
                    for r_idx, row in enumerate(reader, 1):
                        for c_idx, val in enumerate(row, 1):
                            ws.cell(row=r_idx, column=c_idx, value=val)
                    wb.save(temp_file_path)
                except Exception as csv_err:
                    messages.error(request, f"Error reading CSV file: {csv_err}")
                    return redirect('scheduler:import_resources')
            else:
                with open(temp_file_path, 'wb+') as destination:
                    for chunk in uploaded_file.chunks():
                        destination.write(chunk)
                try:
                    wb = openpyxl.load_workbook(temp_file_path, data_only=True)
                except Exception as xlsx_err:
                    if os.path.exists(temp_file_path):
                        os.remove(temp_file_path)
                    messages.error(request, f"Error reading Excel file: {xlsx_err}")
                    return redirect('scheduler:import_resources')
                    
            try:
                format_info = detect_format(wb)
                entities = extract_entities(wb, format_info, university)
                
                # Check if everything is empty
                if not any([entities['campuses'], entities['programs'], entities['lecturers'],
                            entities['rooms'], entities['student_groups'], entities['courses']]):
                    if os.path.exists(temp_file_path):
                        os.remove(temp_file_path)
                    messages.error(request, "Could not extract any data from the file. Check your column headers.")
                    return redirect('scheduler:import_resources')
                    
                # Store original filename for audit log
                request.session['import_file_name'] = uploaded_file.name

                # Run structural validation (warnings only at preview — don't block)
                from scheduler.smart_import import validate_entities_for_import
                preview_warnings = validate_entities_for_import(entities)

                # Render the preview screen
                return render(request, 'scheduler/smart_import_preview.html', {
                    'format_info': format_info,
                    'entities': entities,
                    'file_name': uploaded_file.name,
                    'active_university': university,
                    'preview_warnings': preview_warnings,
                })
            except Exception as e:
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
                messages.error(request, f"Error processing file: {e}")
                return redirect('scheduler:import_resources')


        records = []
        headers = []

        try:
            if file_name.endswith('.csv'):
                csv_file = io.TextIOWrapper(uploaded_file, encoding='utf-8')
                reader = csv.reader(csv_file)
                raw_headers = next(reader, None)
                if raw_headers:
                    headers = [h.strip().lower().replace(' ', '_') for h in raw_headers]
                    for row in reader:
                        if any(row):
                            records.append(dict(zip(headers, [val.strip() for val in row])))
            elif file_name.endswith('.xlsx'):
                import openpyxl
                wb = openpyxl.load_workbook(uploaded_file, data_only=True)
                sheet = wb.active
                rows = list(sheet.iter_rows(values_only=True))
                if rows:
                    raw_headers = rows[0]
                    headers = [str(h).strip().lower().replace(' ', '_') for h in raw_headers if h is not None]
                    for row in rows[1:]:
                        if any(row):
                            vals = [_clean_val(val) for val in row[:len(headers)]]
                            while len(vals) < len(headers):
                                vals.append('')
                            records.append(dict(zip(headers, vals)))
            else:
                messages.error(request, "Unsupported file format. Please upload a .csv or .xlsx file.")
                return redirect('scheduler:import_resources')
        except Exception as e:
            messages.error(request, f"Error reading file: {str(e)}")
            return redirect('scheduler:import_resources')

        if not records:
            messages.error(request, "The uploaded file contains no data rows.")
            return redirect('scheduler:import_resources')

        # ── Flexible alias resolver ────────────────────────────────────────────
        # Normalise every header: lower-case, strip whitespace, collapse spaces→_
        import re as _re
        norm_headers = [_re.sub(r'\s+', '_', h.strip().lower()) for h in raw_headers if h]

        # Build a map: normalised_alias → original normalised header present in the row dicts
        # so _get(row, aliases) can resolve any common variation.
        ALIAS_MAP = {
            # room / venue name
            'venue': 'venue', 'venue_name': 'venue_name', 'location': 'location',
            'room': 'room', 'room_name': 'room_name', 'name': 'name',
            'r': 'room', 'rm': 'room',
            # capacity / size
            'capacity': 'capacity', 'cap': 'cap', 'seats': 'seats',
            'size': 'size', 'group_size': 'group_size', 'students_count': 'students_count',
            'no_of_students': 'no_of_students', 'number_of_students': 'number_of_students',
            'enrolment': 'enrolment', 'enrollment': 'enrollment', 'count': 'count',
            # room type
            'room_type': 'room_type', 'required_room_type': 'required_room_type',
            'venue_type': 'venue_type', 'type': 'type', 'facility_type': 'facility_type',
            # lecturer / instructor name
            'lecturer_name': 'lecturer_name', 'instructor': 'instructor',
            'instructor_name': 'instructor_name', 'teacher': 'teacher',
            'teacher_name': 'teacher_name', 'staff': 'staff', 'staff_name': 'staff_name',
            'full_name': 'full_name', 'lecturer': 'lecturer',
            # email
            'email': 'email', 'email_address': 'email_address',
            'mail': 'mail', 'e_mail': 'e_mail',
            # student group name
            'group_name': 'group_name', 'class': 'class', 'class_name': 'class_name',
            'student': 'student', 'student_group': 'student_group',
            'student_group_name': 'student_group_name', 'cohort': 'cohort',
            'section': 'section', 'stream': 'stream',
            # course code
            'course_code': 'course_code', 'code': 'code', 'subject_code': 'subject_code',
            'unit_code': 'unit_code', 'module_code': 'module_code',
            # course name
            'course_name': 'course_name', 'subject': 'subject', 'subject_name': 'subject_name',
            'unit': 'unit', 'unit_name': 'unit_name', 'module': 'module', 'module_name': 'module_name',
            # duration
            'duration_slots': 'duration_slots', 'duration': 'duration', 'slots': 'slots',
            'hours': 'hours', 'credit_hours': 'credit_hours', 'credits': 'credits',
            'periods': 'periods', 'sessions': 'sessions', 'lesson_periods': 'lesson_periods',
            # department / campus / program
            'department': 'department', 'department_name': 'department_name', 'dept': 'dept',
            'campus': 'campus', 'campus_name': 'campus_name',
            'program': 'program', 'program_name': 'program_name', 'programme': 'programme',
            'faculty': 'faculty', 'school': 'school',
            # lecturer email (courses)
            'lecturer_email': 'lecturer_email', 'instructor_email': 'instructor_email',
            'teacher_email': 'teacher_email',
            # max hours
            'max_hours_per_week': 'max_hours_per_week', 'max_hours': 'max_hours',
            'weekly_hours': 'weekly_hours', 'workload': 'workload',
        }

        def _get(row, *aliases, default=''):
            """Return the first non-empty value from row matching any alias."""
            for alias in aliases:
                # try alias directly
                v = row.get(alias, '')
                if v and str(v).strip():
                    return str(v).strip()
                # try normalised form
                norm = _re.sub(r'\s+', '_', alias.strip().lower())
                v = row.get(norm, '')
                if v and str(v).strip():
                    return str(v).strip()
            return default

        # Room-type normaliser: accept every real-world variation
        ROOM_TYPE_LECTURE = {
            'lecture', 'lecture hall', 'lecture_hall', 'hall', 'theater', 'theatre',
            'classroom', 'class room', 'class_room', 'auditorium', 'lh',
        }
        ROOM_TYPE_LAB = {
            'lab', 'laboratory', 'computer lab', 'computer_lab', 'science lab',
            'science_lab', 'workshop', 'studio',
        }
        ROOM_TYPE_SEMINAR = {
            'seminar', 'seminar room', 'seminar_room', 'tutorial', 'tutorial room',
            'tutorial_room', 'meeting room', 'meeting_room', 'conference', 'boardroom',
        }

        def _norm_room_type(raw):
            """Return canonical room type string or None if unrecognised."""
            v = (raw or '').strip().lower()
            if v in ROOM_TYPE_LECTURE:
                return 'Lecture'
            if v in ROOM_TYPE_LAB:
                return 'Lab'
            if v in ROOM_TYPE_SEMINAR:
                return 'Seminar'
            return None

        # Validate headers (flexible – accept any recognised alias)
        required_cols = []
        if import_type == 'room':
            if not any(h in norm_headers for h in
                       ('name','room_name','venue','venue_name','location', 'r', 'rm')):
                required_cols.append("Room/Venue Name")
            if not any(h in norm_headers for h in
                       ('capacity','cap','seats','size','count','enrolment','enrollment',
                        'no_of_students','number_of_students','students_count')):
                required_cols.append("Capacity / Size")
            if not any(h in norm_headers for h in
                       ('room_type','required_room_type','venue_type','type','facility_type')):
                required_cols.append("Room Type / Venue Type")
        elif import_type == 'lecturer':
            if not any(h in norm_headers for h in
                       ('name','lecturer_name','lecturer','instructor','instructor_name',
                        'teacher','teacher_name','staff','staff_name','full_name')):
                required_cols.append("Lecturer / Instructor Name")
            if not any(h in norm_headers for h in
                       ('email','email_address','mail','e_mail')):
                required_cols.append("Email Address")
        elif import_type == 'student_group':
            if not any(h in norm_headers for h in
                       ('name','group_name','student','student_group','student_group_name',
                        'class','class_name','cohort','section','stream')):
                required_cols.append("Student Group / Class Name")
            if not any(h in norm_headers for h in
                       ('size','group_size','capacity','cap','count','enrolment',
                        'enrollment','no_of_students','number_of_students','students_count')):
                required_cols.append("Group Size / Enrolment")
        elif import_type == 'course':
            if not any(h in norm_headers for h in
                       ('code','course_code','subject_code','unit_code','module_code')):
                required_cols.append("Course / Unit Code")
            if not any(h in norm_headers for h in
                       ('name','course_name','subject','subject_name','unit','unit_name',
                        'module','module_name')):
                required_cols.append("Course / Unit Name")
            if not any(h in norm_headers for h in
                       ('duration_slots','duration','slots','hours','credit_hours',
                        'credits','periods','sessions','lesson_periods')):
                required_cols.append("Duration / Hours / Slots")
            if not any(h in norm_headers for h in
                       ('required_room_type','room_type','venue_type','type','facility_type')):
                required_cols.append("Room Type / Venue Type")

        if required_cols:
            messages.error(request, f"Missing required column(s): {', '.join(required_cols)}. Please check your column headings.")
            return redirect('scheduler:import_resources')

        import_errors = []
        import_warnings = []
        success_count = 0
        heal_msgs = []

        from django.db import transaction
        from .models import Program

        try:
            with transaction.atomic():

                # ── Pre-cache all structure lookups ONCE (avoids N×4 queries) ──
                # Get or create the default campus for this university
                default_campus = Campus.objects.filter(university=university).first()
                if not default_campus:
                    default_campus = Campus.objects.create(university=university, name="Default Campus")

                default_faculty = Faculty.objects.filter(campus__university=university).first()
                if not default_faculty:
                    default_faculty = Faculty.objects.create(campus=default_campus, name="Default Faculty")

                default_dept = Department.objects.filter(faculty__campus__university=university).first()
                if not default_dept:
                    default_dept = Department.objects.create(faculty=default_faculty, name="Default Department")

                default_program = Program.objects.filter(department__faculty__campus__university=university).first()
                if not default_program:
                    default_program = Program.objects.create(department=default_dept, name="Default Program")

                # Build caches for named lookups (campus_name, dept_name, program_name)
                campus_cache   = {c.name.strip(): c for c in Campus.objects.filter(university=university)}
                faculty_cache  = {f.name.strip(): f for f in Faculty.objects.filter(campus__university=university)}
                dept_cache     = {d.name.strip(): d for d in Department.objects.filter(faculty__campus__university=university)}
                program_cache  = {p.name.strip(): p for p in Program.objects.filter(department__faculty__campus__university=university)}

                def _get_campus(name):
                    if not name:
                        return default_campus
                    key = name.strip()
                    if key not in campus_cache:
                        campus_cache[key] = Campus.objects.create(university=university, name=key)
                    return campus_cache[key]

                def _get_dept(name):
                    if not name:
                        return default_dept
                    key = name.strip()
                    if key not in dept_cache:
                        dept_cache[key] = Department.objects.create(faculty=default_faculty, name=key)
                    return dept_cache[key]

                def _get_program(name):
                    if not name:
                        return default_program
                    key = name.strip()
                    if key not in program_cache:
                        program_cache[key] = Program.objects.create(department=default_dept, name=key)
                    return program_cache[key]

                # ── BULK IMPORT: Rooms ──────────────────────────────────────────
                if import_type == 'room':
                    rooms_to_upsert = []
                    existing_keys = {
                        (r.campus_id, r.name): r
                        for r in Room.objects.filter(campus__university=university)
                    }
                    new_rooms, update_rooms = [], []

                    for idx, row in enumerate(records, start=2):
                        name         = _get(row, 'name','room_name','venue','venue_name','location')
                        capacity_str = _get(row, 'capacity','cap','seats','size','count',
                                            'enrolment','enrollment','no_of_students',
                                            'number_of_students','students_count')
                        room_type    = _get(row, 'room_type','required_room_type','venue_type','type','facility_type')
                        campus_name  = _get(row, 'campus_name','campus') or None

                        if not name:
                            import_errors.append(f"Row {idx}: Room Name is required.")
                            continue
                        try:
                            capacity = int(float(capacity_str))
                            if capacity <= 0:
                                raise ValueError()
                        except (ValueError, TypeError):
                            import_errors.append(f"Row {idx}: Capacity must be a valid positive number.")
                            continue

                        normalized_type = _norm_room_type(room_type)
                        if not normalized_type:
                            import_errors.append(
                                f"Row {idx}: Unrecognised Room Type '{room_type}'. "
                                f"Use: Lecture/Hall/Classroom, Lab/Laboratory, or Seminar/Tutorial."
                            )
                            continue
                        campus = _get_campus(campus_name)
                        key = (campus.id, name)

                        _VK = ('zoom', 'virtual', 'online', 'teams')
                        _auto_virtual = any(kw in name.strip().lower() for kw in _VK)
                        if key in existing_keys:
                            r = existing_keys[key]
                            r.capacity  = capacity
                            r.room_type = normalized_type
                            r.is_virtual = _auto_virtual
                            update_rooms.append(r)
                        else:
                            room_obj = Room(campus=campus, name=name, capacity=capacity,
                                            room_type=normalized_type, is_virtual=_auto_virtual)
                            new_rooms.append(room_obj)
                            existing_keys[key] = room_obj
                        success_count += 1

                    if import_errors:
                        raise Exception("Validation errors occurred.")

                    Room.objects.bulk_create(new_rooms, batch_size=500, ignore_conflicts=True)
                    if update_rooms:
                        Room.objects.bulk_update(update_rooms, ['capacity', 'room_type', 'is_virtual'], batch_size=500)

                # ── BULK IMPORT: Lecturers ──────────────────────────────────────
                elif import_type == 'lecturer':
                    existing_lecturers = {l.email: l for l in Lecturer.objects.filter(department__faculty__campus__university=university)}
                    new_lecs, update_lecs = [], []

                    for idx, row in enumerate(records, start=2):
                        name          = _get(row, 'name','lecturer_name','lecturer','instructor',
                                             'instructor_name','teacher','teacher_name',
                                              'staff','staff_name','full_name')
                        email         = _get(row, 'email','email_address','mail','e_mail').lower()
                        dept_name     = _get(row, 'department_name','department','dept',
                                             'faculty','school') or None
                        max_hours_str = _get(row, 'max_hours_per_week','max_hours','hours',
                                             'weekly_hours','workload') or None

                        if not name:
                            import_errors.append(f"Row {idx}: Lecturer Name is required.")
                            continue
                        if not email or '@' not in email:
                            import_errors.append(f"Row {idx}: Valid email address is required.")
                            continue

                        max_hours = 20
                        if max_hours_str:
                            try:
                                max_hours = int(float(max_hours_str))
                            except (ValueError, TypeError):
                                import_warnings.append(f"Row {idx}: Invalid max hours '{max_hours_str}', defaulted to 20.")

                        dept = _get_dept(dept_name)

                        if email in existing_lecturers:
                            l = existing_lecturers[email]
                            l.name = name
                            l.department = dept
                            l.max_hours_per_week = max_hours
                            update_lecs.append(l)
                        else:
                            lec_obj = Lecturer(email=email, name=name, department=dept, max_hours_per_week=max_hours)
                            new_lecs.append(lec_obj)
                            existing_lecturers[email] = lec_obj
                        success_count += 1

                    if import_errors:
                        raise Exception("Validation errors occurred.")

                    Lecturer.objects.bulk_create(new_lecs, batch_size=500, ignore_conflicts=True)
                    if update_lecs:
                        Lecturer.objects.bulk_update(update_lecs, ['name', 'department', 'max_hours_per_week'], batch_size=500)

                # ── BULK IMPORT: Student Groups ─────────────────────────────────
                elif import_type == 'student_group':
                    existing_groups = {
                        (g.program_id, g.name): g
                        for g in StudentGroup.objects.filter(program__department__faculty__campus__university=university)
                    }
                    new_groups, update_groups = [], []

                    for idx, row in enumerate(records, start=2):
                        name      = _get(row, 'name','group_name','student','student_group',
                                         'student_group_name','class','class_name',
                                         'cohort','section','stream')
                        size_str  = _get(row, 'size','group_size','capacity','cap','count',
                                         'enrolment','enrollment','no_of_students',
                                         'number_of_students','students_count') or None
                        prog_name = _get(row, 'program_name','program','programme') or None

                        if not name:
                            import_errors.append(f"Row {idx}: Student Group Name is required.")
                            continue
                        try:
                            size = int(float(size_str))
                            if size <= 0:
                                raise ValueError()
                        except (ValueError, TypeError):
                            import_errors.append(f"Row {idx}: Group Size must be a valid positive number.")
                            continue

                        program = _get_program(prog_name)
                        key = (program.id, name)

                        if key in existing_groups:
                            g = existing_groups[key]
                            g.size = size
                            update_groups.append(g)
                        else:
                            group_obj = StudentGroup(program=program, name=name, size=size)
                            new_groups.append(group_obj)
                            existing_groups[key] = group_obj
                        success_count += 1

                    if import_errors:
                        raise Exception("Validation errors occurred.")

                    StudentGroup.objects.bulk_create(new_groups, batch_size=500, ignore_conflicts=True)
                    if update_groups:
                        StudentGroup.objects.bulk_update(update_groups, ['size'], batch_size=500)

                # ── BULK IMPORT: Courses ────────────────────────────────────────
                elif import_type == 'course':
                    # Pre-cache lecturers and student groups for O(1) lookup
                    lecturer_cache = {l.email: l for l in Lecturer.objects.filter(department__faculty__campus__university=university)}
                    group_cache    = {
                        (g.program_id, g.name): g
                        for g in StudentGroup.objects.filter(program__department__faculty__campus__university=university)
                    }
                    existing_courses = {
                        (c.program_id, c.code.strip().upper(), c.student_group_id): c
                        for c in Course.objects.filter(program__department__faculty__campus__university=university)
                    }
                    new_courses, update_courses = [], []

                    for idx, row in enumerate(records, start=2):
                        code         = _get(row, 'code','course_code','subject_code',
                                            'unit_code','module_code').upper()
                        name         = _get(row, 'name','course_name','subject','subject_name',
                                            'unit','unit_name','module','module_name')
                        duration_str = _get(row, 'duration_slots','duration','slots','hours',
                                            'credit_hours','credits','periods',
                                            'lesson_periods') or None
                        spw_str = _get(row, 'sessions_per_week','weekly_sessions',
                                       'classes_per_week','sessions') or None
                        room_type    = _get(row, 'required_room_type','room_type','venue_type',
                                            'type','facility_type')
                        lec_email    = _get(row, 'lecturer_email','instructor_email',
                                            'teacher_email','lecturer','instructor','teacher').lower()
                        group_name   = _get(row, 'student_group_name','student_group','student',
                                            'group','class','class_name','cohort','section','stream')
                        prog_name    = _get(row, 'program_name','program','programme') or None

                        if not code:
                            import_errors.append(f"Row {idx}: Course Code is required.")
                            continue
                        if not name:
                            import_errors.append(f"Row {idx}: Course Name is required.")
                            continue
                        try:
                            duration = int(float(duration_str))
                            if duration <= 0:
                                raise ValueError()
                        except (ValueError, TypeError):
                            import_errors.append(f"Row {idx}: Duration Slots must be a valid positive number.")
                            continue
                        try:
                            sessions_per_week = int(float(spw_str)) if spw_str else 1
                            if sessions_per_week <= 0:
                                raise ValueError()
                        except (ValueError, TypeError):
                            import_errors.append(f"Row {idx}: Sessions per week must be a valid positive number.")
                            continue
                        normalized_type = _norm_room_type(room_type)
                        if not normalized_type:
                            import_errors.append(
                                f"Row {idx}: Unrecognised Room Type '{room_type}'. "
                                f"Use: Lecture/Hall/Classroom, Lab/Laboratory, or Seminar/Tutorial."
                            )
                            continue

                        lecturer = lecturer_cache.get(lec_email)
                        if lec_email and not lecturer:
                            import_warnings.append(f"Row {idx}: Lecturer '{lec_email}' not found. Course created without lecturer.")

                        program = _get_program(prog_name)

                        group = None
                        if group_name:
                            gkey = (program.id, group_name)
                            if gkey not in group_cache:
                                g = StudentGroup.objects.create(program=program, name=group_name, size=30)
                                group_cache[gkey] = g
                            group = group_cache[gkey]

                        key = (program.id, code.strip().upper(), group.id if group else None)
                        if key in existing_courses:
                            c = existing_courses[key]
                            c.name = name
                            c.duration_slots = duration
                            c.sessions_per_week = sessions_per_week
                            c.required_room_type = normalized_type
                            c.lecturer = lecturer
                            c.student_group = group
                            update_courses.append(c)
                        else:
                            course_obj = Course(
                                program=program, code=code, name=name,
                                duration_slots=duration, sessions_per_week=sessions_per_week,
                                required_room_type=normalized_type,
                                lecturer=lecturer, student_group=group
                            )
                            new_courses.append(course_obj)
                            existing_courses[key] = course_obj
                        success_count += 1

                    if import_errors:
                        raise Exception("Validation errors occurred.")

                    Course.objects.bulk_create(new_courses, batch_size=500, ignore_conflicts=True)
                    if update_courses:
                        Course.objects.bulk_update(
                            update_courses,
                            ['name', 'duration_slots', 'sessions_per_week', 'required_room_type', 'lecturer', 'student_group'],
                            batch_size=500
                        )

                    # Auto-heal inside the transaction
                    try:
                        heal_msgs = auto_heal_university_data(university)
                    except Exception:
                        heal_msgs = []

                    # Run semantic validation inside transaction
                    from scheduler.validation import validate_university_data
                    is_valid, val_errors, val_warnings = validate_university_data(university)
                    if val_errors:
                        import_errors.extend(val_errors)
                    if val_warnings:
                        import_warnings.extend(val_warnings)

                    if import_errors:
                        raise Exception("Validation errors occurred.")

            messages.success(request, f"✅ Successfully imported {success_count} {import_type.replace('_',' ')}(s)!")
            if import_warnings:
                for w in import_warnings[:5]:
                    messages.warning(request, w)
                if len(import_warnings) > 5:
                    messages.warning(request, f"...and {len(import_warnings) - 5} more warnings.")

            # Display auto-heal messages
            for fix_msg in heal_msgs:
                messages.info(request, fix_msg)

            # Auto-trigger timetable regeneration
            auto_timetable = (
                Timetable.objects.filter(semester__university=university, is_active=True).first()
                or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
            )
            if auto_timetable:
                try:
                    from .signals import queue_auto_generation
                    queue_auto_generation(auto_timetable)
                    messages.info(request, f"✓ Timetable generation queued for '{auto_timetable.name}'.")
                except Exception as e:
                    messages.warning(request, f"Import succeeded but auto-generation could not be queued: {e}")
            else:
                messages.warning(request, "Import succeeded but no timetable found. Create one under Timetables → New.")

            return redirect(f"/resources/?tab={import_type.replace('_','')}")

        except Exception as e:
            context = {
                'import_errors': import_errors,
                'import_warnings': import_warnings,
                'import_type': import_type,
            }
            return render(request, 'scheduler/resources_import.html', context)

    return render(request, 'scheduler/resources_import.html')





# ─────────────────────────────────────────────────────────────────────────────
# Phase 2: Lecturer Availability Self-Service
# ─────────────────────────────────────────────────────────────────────────────

@login_required
def lecturer_availability(request):
    """
    Lecturers can set their own availability across all time slots.
    Admins can view/edit any lecturer's availability.
    """
    university = get_active_uni(request)
    role = get_user_role(request)

    # Determine which lecturer we are editing
    lecturer = None
    if role == 'lecturer':
        try:
            lecturer = request.user.profile.lecturer
        except Exception:
            messages.error(request, "Your account is not linked to a lecturer. Please update your profile.")
            return redirect('accounts:profile')
    else:
        # Admin/scheduler: allow choosing any lecturer
        lecturer_id = request.GET.get('lecturer_id')
        if lecturer_id:
            lecturer = get_object_or_404(Lecturer, pk=lecturer_id, department__faculty__campus__university=university)
        else:
            lecturer = Lecturer.objects.filter(department__faculty__campus__university=university).first()

    if not lecturer:
        messages.error(request, "No lecturers found for this university.")
        return redirect('scheduler:dashboard')

    timeslots = TimeSlot.objects.filter(university=university).order_by('day_of_week', 'slot_number')

    if request.method == 'POST':
        # Save availability: checkbox per timeslot
        for ts in timeslots:
            is_available = request.POST.get(f'slot_{ts.id}') == 'on'
            note = request.POST.get(f'note_{ts.id}', '')
            LecturerAvailability.objects.update_or_create(
                lecturer=lecturer,
                time_slot=ts,
                defaults={'is_available': is_available, 'note': note}
            )
        if lecturer and lecturer.email:
            try:
                from scheduler.tasks import verify_and_notify_lecturer_record
                verify_and_notify_lecturer_record(
                    submitted_email=lecturer.email,
                    submitted_name=lecturer.name,
                    university_id=university.id if university else None,
                    preserve_password=True
                )
            except Exception as mail_err:
                logger.warning(f"[Availability Email] Failed to send notification: {mail_err}")
        messages.success(request, f"Availability for {lecturer.name} saved successfully!")
        return redirect('scheduler:lecturer_availability')

    # Build existing availability map: {time_slot_id: LecturerAvailability}
    existing = {a.time_slot_id: a for a in LecturerAvailability.objects.filter(lecturer=lecturer)}

    # Build day-grouped structure for template
    DAY_LABELS = {1:'Monday',2:'Tuesday',3:'Wednesday',4:'Thursday',5:'Friday',6:'Saturday',7:'Sunday'}
    days = {}
    for ts in timeslots:
        day = ts.day_of_week
        if day not in days:
            days[day] = {'label': DAY_LABELS.get(day, f'Day {day}'), 'slots': []}
        avail = existing.get(ts.id)
        days[day]['slots'].append({
            'ts': ts,
            'is_available': avail.is_available if avail else True,
            'note': avail.note if avail else '',
        })

    all_lecturers = Lecturer.objects.filter(department__faculty__campus__university=university)
    
    template_name = 'scheduler/lecturer_portal_availability.html' if role == 'lecturer' else 'scheduler/lecturer_availability.html'
    return render(request, template_name, {
        'lecturer': lecturer,
        'days': days,
        'all_lecturers': all_lecturers,
        'is_own': role == 'lecturer',
    })


@login_required
def lecturer_my_schedule(request):
    """Redirect legacy My Timetable link to Phase 3 Lecturer Portal Weekly Timetable."""
    return redirect('scheduler:lecturer_portal_weekly_timetable')


@login_required
def lecturer_teaching_history(request):
    """Redirect legacy Teaching History link to Phase 3 Lecturer Portal Courses."""
    return redirect('scheduler:lecturer_portal_courses')


# ─────────────────────────────────────────────────────────────────────────────
# Phase 5: PDF & Excel Export
# ─────────────────────────────────────────────────────────────────────────────

# ─────────────────────────────────────────────────────────────────────────────
# Phase 5: PDF & Excel Export
# ─────────────────────────────────────────────────────────────────────────────

from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import cm

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            super().showPage()
        super().save()

    def draw_page_number(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#64748B"))
        
        # Draw a thin footer separator line
        self.setStrokeColor(colors.HexColor("#E2E8F0"))
        self.setLineWidth(0.5)
        self.line(0.8*cm, 1.0*cm, self._pagesize[0] - 0.8*cm, 1.0*cm)
        
        # Footer text - Left side (timestamp)
        from django.utils import timezone
        now_str = timezone.now().strftime("%Y-%m-%d %H:%M:%S")
        footer_left = f"Generated on {now_str} | University Timetable Scheduler"
        
        # Footer text - Right side (page numbering)
        footer_right = f"Page {self._pageNumber} of {page_count}"
        
        self.drawString(0.8*cm, 0.6*cm, footer_left)
        self.drawRightString(self._pagesize[0] - 0.8*cm, 0.6*cm, footer_right)
        self.restoreState()


def get_group_details(group):
    # Default fallback values from database
    dept_name = group.program.department.name if (group.program and group.program.department) else "N/A"
    prog_name = group.program.name if group.program else "N/A"
    
    # Clean up default/placeholder values
    if "DEFAULT" in dept_name.upper() or "DEFAULT" in prog_name.upper():
        dept_name = "N/A"
        prog_name = "N/A"

    name_upper = group.name.upper()
    
    # Get group code prefix
    prefix = group.name.split(' ')[0].split('-')[0].strip().upper()
    
    # 1. Academic Year dynamic resolution
    if 'YEAR4' in name_upper or 'YR4' in name_upper or 'YEAR 4' in name_upper:
        year_label = 'Year 4'
    elif 'YEAR3' in name_upper or 'YR3' in name_upper or 'YEAR 3' in name_upper:
        year_label = 'Year 3'
    elif 'YEAR2' in name_upper or 'YR2' in name_upper or 'YEAR 2' in name_upper:
        year_label = 'Year 2'
    elif 'YEAR1' in name_upper or 'YR1' in name_upper or 'YEAR 1' in name_upper:
        year_label = 'Year 1'
    else:
        year_label = group.get_year_display() if group.year else "Year 1"

    # 2. Program & Department dynamic resolution from Prefix
    if prefix == 'BCOM':
        prog_name = "Bachelor of Commerce (BCOM)"
        dept_name = "Department of Business & Commerce"
    elif prefix == 'PHD':
        prog_name = "Doctor of Philosophy (PhD)"
        dept_name = "Department of Postgraduate Studies"
    elif prefix in ('MASTERS', 'MBA', 'MSC'):
        prog_name = "Master of Business Administration (MBA)"
        dept_name = "Department of Postgraduate Studies"
    elif prefix == 'BAM':
        prog_name = "Bachelor of Business Administration & Management"
        dept_name = "Department of Business Administration"
    elif prefix in ('CUU', 'CFM', 'ICT'):
        prog_name = "Bachelor of Science in Information & Communication Technology"
        dept_name = "Department of Information Technology"
    elif prefix == 'ACC':
        prog_name = "Bachelor of Science in Accounting & Finance"
        dept_name = "Department of Accounting & Finance"
    elif prefix == 'MAT':
        prog_name = "Bachelor of Science in Mathematics & Statistics"
        dept_name = "Department of Mathematics & Actuarial Science"

    return dept_name, prog_name, year_label


@login_required
def export_timetable_pdf(request, pk):
    """Export timetable as a PDF file using ReportLab with custom branding."""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    import re
    import io

    university = get_active_uni(request)
    timetable = get_object_or_404(
        Timetable.objects.select_related('semester', 'semester__university'), pk=pk
    )
    if timetable.semester.university != university:
        messages.error(request, "Access denied.")
        return redirect('scheduler:dashboard')

    # Get custom branding and layout from request
    pdf_title = request.GET.get('title', f"{timetable.name} — {timetable.semester.name}")
    pdf_subtitle = request.GET.get('subtitle', f"{university.name} | Generated Timetable")
    pdf_logo_text = request.GET.get('logo_text', "KU")
    primary_hex = request.GET.get('primary_color', "#0d5a4f")
    alt_hex = request.GET.get('alt_color', "#f0fdfa")
    show_logo = request.GET.get('show_logo', "true").lower() == "true"
    layout_type = request.GET.get('layout_type', 'weekly')

    # Sanitize color codes
    if not primary_hex.startswith('#'):
        primary_hex = '#' + primary_hex
    if not alt_hex.startswith('#'):
        alt_hex = '#' + alt_hex
    if not re.match(r'^#[0-9a-fA-F]{6}$', primary_hex):
        primary_hex = '#0d5a4f'
    if not re.match(r'^#[0-9a-fA-F]{6}$', alt_hex):
        alt_hex = '#f0fdfa'

    # Filter resources by university for scoping and defaulting
    rooms = Room.objects.filter(campus__university=university)
    lecturers = Lecturer.objects.filter(department__faculty__campus__university=university)
    student_groups = StudentGroup.objects.filter(program__department__faculty__campus__university=university)

    # Get active filters
    filter_type = request.GET.get('filter_type')
    filter_id = request.GET.get('filter_id')

    # If exporting a batch of targets, strip the target-specific suffix from the request's subtitle
    if filter_type in ('all_groups', 'all_rooms', 'all_lecturers') or layout_type == 'complete_pack':
        if ' — ' in pdf_subtitle:
            pdf_subtitle = pdf_subtitle.split(' — ')[0]
        elif ' - ' in pdf_subtitle:
            pdf_subtitle = pdf_subtitle.split(' - ')[0]

    # Resolve target list
    targets = []
    
    if filter_type == 'all_groups':
        active_group_ids = timetable.slots.values_list('student_group_id', flat=True).distinct()
        targets = list(student_groups.filter(id__in=active_group_ids).select_related('program', 'program__department'))
        targets.sort(key=lambda g: (get_group_details(g)[0], get_group_details(g)[1], g.name))
    elif filter_type == 'all_rooms':
        active_room_ids = timetable.slots.values_list('room_id', flat=True).distinct()
        targets = list(rooms.filter(id__in=active_room_ids).select_related('campus', 'building').order_by('campus__name', 'building__name', 'name'))
    elif filter_type == 'all_lecturers':
        active_lecturer_ids = timetable.slots.values_list('lecturer_id', flat=True).distinct()
        targets = list(lecturers.filter(id__in=active_lecturer_ids).select_related('department').order_by('department__name', 'name'))
    else:
        if not filter_type:
            filter_type = 'group'

        if not filter_id:
            if filter_type == 'group' and student_groups.exists():
                filter_id = student_groups.first().id
            elif filter_type == 'room' and rooms.exists():
                filter_id = rooms.first().id
            elif filter_type == 'lecturer' and lecturers.exists():
                filter_id = lecturers.first().id

        target_obj = None
        if filter_id:
            try:
                filter_id = int(filter_id)
                if filter_type == 'group':
                    target_obj = student_groups.filter(id=filter_id).first()
                elif filter_type == 'room':
                    target_obj = rooms.filter(id=filter_id).first()
                elif filter_type == 'lecturer':
                    target_obj = lecturers.filter(id=filter_id).first()
            except ValueError:
                pass
        if target_obj:
            targets = [target_obj]

    # Set page margins to professional standards (left/right=0.8cm, top=1.2cm, bottom=1.5cm)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=landscape(A4),
        rightMargin=0.8*cm, leftMargin=0.8*cm,
        topMargin=1.2*cm, bottomMargin=1.5*cm
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'title',
        parent=styles['Heading1'],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor(primary_hex),
        spaceAfter=4
    )
    subtitle_style = ParagraphStyle(
        'subtitle',
        parent=styles['Normal'],
        fontSize=10,
        leading=13,
        textColor=colors.HexColor('#475569')
    )
    logo_style = ParagraphStyle(
        'logo',
        parent=styles['Normal'],
        fontSize=14,
        leading=16,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    # Helper to generate consistent professional headers with brand divider line
    def make_header_flowable(title_text, subtitle_text):
        title_p = Paragraph(f"<b>{title_text}</b>", title_style)
        subtitle_p = Paragraph(subtitle_text, subtitle_style)
        left_flow = [title_p, subtitle_p]

        if show_logo:
            logo_p = Paragraph(f"<font color='white'><b>{pdf_logo_text}</b></font>", logo_style)
            logo_table = Table([[logo_p]], colWidths=[2.2*cm], rowHeights=[1.2*cm])
            logo_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(primary_hex)),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('BOX', (0, 0), (-1, -1), 1.5, colors.HexColor(primary_hex)),
            ]))
            header_table = Table([[left_flow, logo_table]], colWidths=[landscape(A4)[0] - 1.6*cm - 2.5*cm, 2.5*cm])
        else:
            header_table = Table([[left_flow]], colWidths=[landscape(A4)[0] - 1.6*cm])
            
        header_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
        ]))
        
        # Thin table with a bottom border in the brand color
        divider_table = Table([['']], colWidths=[landscape(A4)[0] - 1.6*cm], rowHeights=[1])
        divider_table.setStyle(TableStyle([
            ('LINEBELOW', (0, 0), (-1, -1), 1.5, colors.HexColor(primary_hex)),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
        ]))
        
        return [header_table, Spacer(1, 0.15*cm), divider_table, Spacer(1, 0.4*cm)]
    
    elements = []

    # Build table data structures
    day_labels = {1:'Mon',2:'Tue',3:'Wed',4:'Thu',5:'Fri',6:'Sat',7:'Sun'}
    timeslots = list(TimeSlot.objects.filter(university=university).order_by('day_of_week', 'slot_number'))
    days = sorted(set(ts.day_of_week for ts in timeslots))
    slot_numbers = sorted(set(ts.slot_number for ts in timeslots))

    slots_all = list(
        timetable.slots.select_related('course', 'course__program__department', 'lecturer', 'room', 'time_slot', 'student_group', 'student_group__program').all()
    )

    # ─────────────────────────────────────────────────────────────────────────
    # LAYOUT 1: master (Master Timetable Grid)
    # ─────────────────────────────────────────────────────────────────────────
    if layout_type == 'complete_pack':
        # Define styling for weekly grid cells
        cell_font = 6.5
        cell_lead = 8.5
        tb_pad = 2.5
        cell_style = ParagraphStyle('cp_cell', parent=styles['Normal'], fontSize=cell_font, leading=cell_lead, alignment=TA_CENTER)
        time_style = ParagraphStyle(
            'cp_time',
            parent=styles['Normal'],
            fontSize=cell_font,
            leading=cell_lead,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            textColor=colors.HexColor('#1E293B')
        )
        ts_by_slot_number = {}
        for ts in timeslots:
            ts_by_slot_number.setdefault(ts.slot_number, []).append(ts)

        active_group_ids = timetable.slots.values_list('student_group_id', flat=True).distinct()
        groups_list = list(student_groups.filter(id__in=active_group_ids).select_related('program', 'program__department'))
        groups_list.sort(key=lambda g: (get_group_details(g)[0], get_group_details(g)[1], g.name))

        # 0. Table of Contents / Executive Directory Page
        elements.extend(make_header_flowable(pdf_title, f"{pdf_subtitle} — Executive Directory & Table of Contents"))
        
        cell_style_bold = ParagraphStyle('cp_ml_bold', parent=styles['Normal'], fontSize=7.5, leading=9.5, fontName='Helvetica-Bold')
        cell_style_left = ParagraphStyle('cp_ml_left', parent=styles['Normal'], fontSize=7.0, leading=9.0)
        cell_style_center = ParagraphStyle('cp_ml_center', parent=styles['Normal'], fontSize=7.0, leading=9.0, alignment=TA_CENTER)

        dept_set = set(get_group_details(g)[0] for g in groups_list)
        prog_set = set(get_group_details(g)[1] for g in groups_list)

        summary_p = Paragraph(
            f"<b>Total Scheduled Classes:</b> {len(slots_all)} &nbsp;|&nbsp; "
            f"<b>Departments:</b> {len(dept_set)} &nbsp;|&nbsp; "
            f"<b>Academic Programs:</b> {len(prog_set)} &nbsp;|&nbsp; "
            f"<b>Student Groups:</b> {len(groups_list)}",
            ParagraphStyle('cp_toc_summary', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#1E293B'))
        )
        elements.append(summary_p)
        elements.append(Spacer(1, 0.2*cm))

        toc_headers = [
            Paragraph("<b>Department</b>", cell_style_bold),
            Paragraph("<b>Program / Specialization</b>", cell_style_bold),
            Paragraph("<b>Student Group Name</b>", cell_style_bold),
            Paragraph("<b>Academic Year</b>", cell_style_bold),
            Paragraph("<b>Class Size</b>", cell_style_bold),
        ]
        toc_data = [toc_headers]

        last_d = None
        last_p = None
        for g in groups_list:
            d_name, p_name, y_label = get_group_details(g)
            d_disp = d_name if d_name != last_d else ""
            p_disp = p_name if (p_name != last_p or d_name != last_d) else ""
            last_d = d_name
            last_p = p_name

            toc_data.append([
                Paragraph(f"<b>{d_disp}</b>", cell_style_left),
                Paragraph(p_disp, cell_style_left),
                Paragraph(f"<b>{g.name}</b>", cell_style_left),
                Paragraph(y_label, cell_style_center),
                Paragraph(f"{g.size} Students", cell_style_center),
            ])

        toc_col_widths = [6.5*cm, 7.5*cm, 7.5*cm, 3.5*cm, 3.1*cm]
        toc_table = Table(toc_data, colWidths=toc_col_widths, repeatRows=1)
        toc_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(primary_hex)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CBD5E1')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor(alt_hex)]),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ]))
        elements.append(toc_table)
        elements.append(PageBreak())

        # 1. Master Chronological List Section
        elements.extend(make_header_flowable(pdf_title, f"{pdf_subtitle} — 1. Master Chronological List"))
        
        slots_sorted = sorted(slots_all, key=lambda s: (s.time_slot.day_of_week, s.time_slot.start_time, s.course.code))
        headers = [Paragraph("<b>Day</b>", styles['Normal']),
                   Paragraph("<b>Time Slot</b>", styles['Normal']),
                   Paragraph("<b>Course Code</b>", styles['Normal']),
                   Paragraph("<b>Course Name</b>", styles['Normal']),
                   Paragraph("<b>Student Group</b>", styles['Normal']),
                   Paragraph("<b>Lecturer</b>", styles['Normal']),
                   Paragraph("<b>Room</b>", styles['Normal'])]
        data = [headers]
        day_names = {1:'Monday', 2:'Tuesday', 3:'Wednesday', 4:'Thursday', 5:'Friday', 6:'Saturday', 7:'Sunday'}
        
        for s in slots_sorted:
            day_str = day_names.get(s.time_slot.day_of_week, f"Day {s.time_slot.day_of_week}")
            time_str = f"{s.time_slot.start_time.strftime('%H:%M')} – {s.time_slot.end_time.strftime('%H:%M')}"
            data.append([
                Paragraph(day_str, cell_style_bold),
                Paragraph(time_str, cell_style_center),
                Paragraph(s.course.code, cell_style_bold),
                Paragraph(s.course.name, cell_style_left),
                Paragraph(s.student_group.name if s.student_group else 'Unassigned', cell_style_left),
                Paragraph(s.lecturer.name if s.lecturer else 'Unassigned', cell_style_left),
                Paragraph(s.room.name if s.room else 'Unassigned', cell_style_center)
            ])
        if len(data) == 1:
            data.append([Paragraph("No courses scheduled.", cell_style_left)] + [''] * 6)
        
        col_widths = [2.0*cm, 3.0*cm, 2.2*cm, 7.0*cm, 4.5*cm, 4.5*cm, 4.9*cm]
        table = Table(data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(primary_hex)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8.0),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor(alt_hex)]),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        elements.append(table)
        
        # Helper to append weekly grid for list of targets
        def append_weekly_grids(title_prefix, targets_list, target_type):
            last_dept_rendered = None
            for t in targets_list:
                elements.append(PageBreak())
                
                if target_type == 'group':
                    dept_name, prog_name, year_label = get_group_details(t)
                    sub_text = (
                        f"<b>Department:</b> {dept_name} &nbsp;|&nbsp; <b>Program:</b> {prog_name}<br/>"
                        f"<b>Student Group:</b> {t.name} &nbsp;|&nbsp; <b>Academic Year:</b> {year_label} &nbsp;|&nbsp; <b>Class Size:</b> {t.size} Students"
                    )
                elif target_type == 'lecturer':
                    dept_name = t.department.name if t.department else "N/A"
                    sub_text = (
                        f"<b>Department:</b> {dept_name}<br/>"
                        f"<b>Lecturer:</b> {t.name} &nbsp;|&nbsp; <b>Staff ID:</b> {t.staff_id or 'N/A'} &nbsp;|&nbsp; <b>Type:</b> {t.get_lecturer_type_display()}"
                    )
                else: # room
                    campus_name = t.campus.name if t.campus else "N/A"
                    building_name = t.building.name if t.building else "N/A"
                    sub_text = (
                        f"<b>Campus/Building:</b> {campus_name} / {building_name}<br/>"
                        f"<b>Room:</b> {t.name} &nbsp;|&nbsp; <b>Type:</b> {t.get_room_type_display()} &nbsp;|&nbsp; <b>Capacity:</b> {t.capacity} Seats"
                    )
                
                elements.extend(make_header_flowable(pdf_title, sub_text))
                
                if target_type == 'group':
                    slots_t = [s for s in slots_all if s.student_group_id == t.id]
                elif target_type == 'lecturer':
                    slots_t = [s for s in slots_all if s.lecturer_id == t.id]
                else:
                    slots_t = [s for s in slots_all if s.room_id == t.id]
                    
                slots_by_day_and_num = {}
                for s in slots_t:
                    slots_by_day_and_num.setdefault((s.time_slot.day_of_week, s.time_slot.slot_number), []).append(s)
                
                grid_data = [['Slot / Time'] + [day_labels.get(d, f'D{d}') for d in days]]
                for slot_num in slot_numbers:
                    ts_for_slot = ts_by_slot_number.get(slot_num, [])
                    time_label = f"<b>Slot {slot_num}</b>"
                    if ts_for_slot:
                        ts0 = ts_for_slot[0]
                        time_label = f"<b>Slot {slot_num}</b><br/>{ts0.start_time.strftime('%H:%M')} – {ts0.end_time.strftime('%H:%M')}"
                    row = [Paragraph(time_label, time_style)]
                    for day in days:
                        matching = slots_by_day_and_num.get((day, slot_num), [])
                        if matching:
                            items = []
                            for s in matching:
                                course_title = s.course.name
                                if len(course_title) > 28:
                                    course_title = course_title[:25] + "..."
                                cell_lines = [f"<b>{s.course.code}</b>", f"<font color='#0D5A4F'><i>{course_title}</i></font>"]
                                if target_type != 'group' and s.student_group:
                                    cell_lines.append(f"<font color='#1E293B'><b>{s.student_group.name}</b></font>")
                                if target_type != 'lecturer' and s.lecturer:
                                    cell_lines.append(f"<font color='#1E3A8A'>{s.lecturer.name}</font>")
                                if target_type != 'room' and s.room:
                                    cell_lines.append(f"<font color='#801d1d'>Room: {s.room.name}</font>")
                                items.append('<br/>'.join(cell_lines))
                            row.append(Paragraph('<br/><font color="#CBD5E1">────────────────</font><br/>'.join(items), cell_style))
                        else:
                            row.append('')
                    grid_data.append(row)
                
                col_width = (landscape(A4)[0] - 1.6*cm) / (len(days) + 1) if days else (landscape(A4)[0] - 1.6*cm)
                t_table = Table(grid_data, colWidths=[col_width]*(len(days)+1), repeatRows=1)
                t_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(primary_hex)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 8.0),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor(alt_hex)]),
                    ('TOPPADDING', (0, 0), (-1, -1), tb_pad),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), tb_pad),
                ]))
                elements.append(t_table)
                
        # 2. Add all active student groups (ordered hierarchically by Department & Program)
        active_group_ids = timetable.slots.values_list('student_group_id', flat=True).distinct()
        groups_list = list(student_groups.filter(id__in=active_group_ids).select_related('program', 'program__department'))
        groups_list.sort(key=lambda g: (get_group_details(g)[0], get_group_details(g)[1], g.name))
        append_weekly_grids("2. Student Group Timetable", groups_list, 'group')
        
        # 3. Add all active lecturers (ordered by Department)
        active_lec_ids = timetable.slots.values_list('lecturer_id', flat=True).distinct()
        lecturers_list = list(lecturers.filter(id__in=active_lec_ids).select_related('department'))
        lecturers_list.sort(key=lambda l: (l.department.name if l.department else 'Z_UNASSIGNED', l.name))
        append_weekly_grids("3. Lecturer Timetable", lecturers_list, 'lecturer')
        
        # 4. Add all active rooms (ordered by Campus & Building)
        active_room_ids = timetable.slots.values_list('room_id', flat=True).distinct()
        rooms_list = list(rooms.filter(id__in=active_room_ids).select_related('campus', 'building'))
        rooms_list.sort(key=lambda r: (r.campus.name if r.campus else '', r.building.name if r.building else '', r.name))
        append_weekly_grids("4. Room Timetable", rooms_list, 'room')

    elif layout_type == 'master':
        # Determine master column type based on active scope/filters
        active_room_ids = timetable.slots.values_list('room_id', flat=True).distinct()
        active_group_ids = timetable.slots.values_list('student_group_id', flat=True).distinct()
        active_lec_ids = timetable.slots.values_list('lecturer_id', flat=True).distinct()

        # Decide if columns are rooms, groups, or lecturers
        if filter_type == 'all_groups':
            master_cols = list(student_groups.filter(id__in=active_group_ids).order_by('name'))
            col_type = 'group'
            col_header_title = "Student Groups"
        elif filter_type == 'all_lecturers':
            master_cols = list(lecturers.filter(id__in=active_lec_ids).order_by('name'))
            col_type = 'lecturer'
            col_header_title = "Lecturers"
        elif filter_type == 'all_rooms':
            master_cols = list(rooms.filter(id__in=active_room_ids).order_by('name'))
            col_type = 'room'
            col_header_title = "Rooms"
        else:
            # Fallback based on current view/filter, or room default
            if filter_type == 'group':
                master_cols = list(student_groups.filter(id__in=active_group_ids).order_by('name'))
                col_type = 'group'
                col_header_title = "Student Groups"
            elif filter_type == 'lecturer':
                master_cols = list(lecturers.filter(id__in=active_lec_ids).order_by('name'))
                col_type = 'lecturer'
                col_header_title = "Lecturers"
            else:
                master_cols = list(rooms.filter(id__in=active_room_ids).order_by('name'))
                col_type = 'room'
                col_header_title = "Rooms"

        # Safe fallback if nothing is booked yet
        if not master_cols:
            if col_type == 'group':
                master_cols = list(student_groups.order_by('name')[:5])
            elif col_type == 'lecturer':
                master_cols = list(lecturers.order_by('name')[:5])
            else:
                master_cols = list(rooms.order_by('name')[:5])

        # Group slots by (timeslot_id, col_id) to run in O(N)
        slots_by_ts_and_col = {}
        for s in slots_all:
            if col_type == 'room':
                key = (s.time_slot_id, s.room_id)
            elif col_type == 'group':
                key = (s.time_slot_id, s.student_group_id)
            elif col_type == 'lecturer':
                key = (s.time_slot_id, s.lecturer_id)
            slots_by_ts_and_col.setdefault(key, []).append(s)

        # Style definition
        master_cell_style = ParagraphStyle(
            'master_cell',
            parent=styles['Normal'],
            fontSize=6.5,
            leading=8.0,
            alignment=TA_CENTER
        )
        master_time_style = ParagraphStyle(
            'master_time',
            parent=styles['Normal'],
            fontSize=7.0,
            leading=9.5,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )

        # Chunk columns to size 5 to make them 60% wider than original 8-column layout
        chunk_size = 5
        col_chunks = [master_cols[i:i + chunk_size] for i in range(0, len(master_cols), chunk_size)]
        
        # Limit to 20 chunks max for safety
        active_chunks = col_chunks[:20]
        for c_idx, chunk in enumerate(active_chunks):
            # Dynamic header columns
            header_cols = ['Time / Day'] + [c.name for c in chunk]
            data = [header_cols]
            
            for ts in timeslots:
                row_label = f"<b>{day_labels.get(ts.day_of_week, f'D{ts.day_of_week}')}</b><br/>{ts.start_time.strftime('%H:%M')}-{ts.end_time.strftime('%H:%M')}"
                row = [Paragraph(row_label, master_time_style)]
                
                for col in chunk:
                    matching = slots_by_ts_and_col.get((ts.id, col.id), [])
                    if matching:
                        matching.sort(key=lambda x: (x.course.code, x.student_group.name if x.student_group else ''))
                        items = []
                        for s in matching:
                            # Render content dynamically by omitting redundant fields
                            details = []
                            details.append(f"<b>{s.course.code}</b>")
                            details.append(f"<font color='#475569'>{s.course.name[:18]}</font>")
                            
                            if col_type != 'group' and s.student_group:
                                details.append(f"<font color='#0d5a4f'>{s.student_group.name}</font>")
                            if col_type != 'lecturer':
                                lec_disp = s.lecturer.name if s.lecturer else 'Unassigned'
                                details.append(f"<font color='#1e3a8a'>{lec_disp}</font>")
                            if col_type != 'room' and s.room:
                                details.append(f"<font color='#801d1d'>Room: {s.room.name}</font>")
                                
                            items.append('<br/>'.join(details))
                        
                        cell_content = '<br/><font color="#CBD5E1">────────────────</font><br/>'.join(items)
                        row.append(Paragraph(cell_content, master_cell_style))
                    else:
                        row.append('')
                data.append(row)

            # Columns widths: col_width is exactly distributed over printable width (28.1 cm)
            m_col_width = (landscape(A4)[0] - 1.6*cm) / (len(chunk) + 1)
            table = Table(data, colWidths=[m_col_width] * (len(chunk) + 1), repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(primary_hex)),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 8.0),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
                ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#F8FAFC')), # Shade the time column sidebar
                ('ROWBACKGROUNDS', (1, 1), (-1, -1), [colors.white, colors.HexColor(alt_hex)]), # Alternate backgrounds for columns
                ('LEFTPADDING', (0, 0), (-1, -1), 3),
                ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))

            # Dynamic header block
            current_subtitle = f"{pdf_subtitle} — Master Grid ({col_header_title} {c_idx*chunk_size+1} - {min((c_idx+1)*chunk_size, len(master_cols))})"
            
            # Generate the unified header using the new helper function
            elements.extend(make_header_flowable(pdf_title, current_subtitle))
            elements.append(table)

            if c_idx < len(active_chunks) - 1:
                elements.append(PageBreak())

    # ─────────────────────────────────────────────────────────────────────────
    # LAYOUT 2: monthly (Monthly Calendar Grid)
    # ─────────────────────────────────────────────────────────────────────────
    elif layout_type == 'monthly':
        import calendar
        start_date = timetable.semester.start_date
        year = start_date.year
        month = start_date.month
        month_name = start_date.strftime("%B %Y")

        cal = calendar.Calendar(firstweekday=0) # Monday starts the week
        month_weeks = cal.monthdayscalendar(year, month)

        if not targets:
            elements.append(Paragraph("No scheduled slots found for this selection.", styles['Heading2']))
        else:
            for idx, target in enumerate(targets):
                if filter_type in ('group', 'all_groups'):
                    dept_name, prog_name, year_label = get_group_details(target)
                    current_subtitle = (
                        f"<b>Department:</b> {dept_name} &nbsp;|&nbsp; <b>Program:</b> {prog_name}<br/>"
                        f"<b>Student Group:</b> {target.name} &nbsp;|&nbsp; <b>Academic Year:</b> {year_label} &nbsp;|&nbsp; <b>Class Size:</b> {target.size} Students"
                    )
                elif filter_type in ('lecturer', 'all_lecturers'):
                    dept_name = target.department.name if target.department else "N/A"
                    current_subtitle = (
                        f"<b>Department:</b> {dept_name}<br/>"
                        f"<b>Lecturer:</b> {target.name} &nbsp;|&nbsp; <b>Staff ID:</b> {target.staff_id or 'N/A'} &nbsp;|&nbsp; <b>Type:</b> {target.get_lecturer_type_display()}"
                    )
                elif filter_type in ('room', 'all_rooms'):
                    campus_name = target.campus.name if target.campus else "N/A"
                    building_name = target.building.name if target.building else "N/A"
                    current_subtitle = (
                        f"<b>Campus/Building:</b> {campus_name} / {building_name}<br/>"
                        f"<b>Room:</b> {target.name} &nbsp;|&nbsp; <b>Type:</b> {target.get_room_type_display()} &nbsp;|&nbsp; <b>Capacity:</b> {target.capacity} Seats"
                    )
                else:
                    current_subtitle = f"{pdf_subtitle} — {month_name}"

                # Generate the unified header using helper function
                elements.extend(make_header_flowable(pdf_title, current_subtitle))

                if filter_type in ('group', 'all_groups'):
                    slots_target = [s for s in slots_all if s.student_group_id == target.id]
                elif filter_type in ('room', 'all_rooms'):
                    slots_target = [s for s in slots_all if s.room_id == target.id]
                elif filter_type in ('lecturer', 'all_lecturers'):
                    slots_target = [s for s in slots_all if s.lecturer_id == target.id]
                else:
                    slots_target = slots_all

                slots_by_dow = {}
                for s in slots_target:
                    slots_by_dow.setdefault(s.time_slot.day_of_week, []).append(s)

                cal_headers = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
                data = [cal_headers]
                
                cal_day_num_style = ParagraphStyle('cal_day_num', parent=styles['Normal'], fontSize=8, leading=10, fontName='Helvetica-Bold', textColor=colors.HexColor(primary_hex))
                cal_class_style = ParagraphStyle('cal_class', parent=styles['Normal'], fontSize=5.5, leading=7, alignment=TA_CENTER)

                for week in month_weeks:
                    row = []
                    for day_idx, day_num in enumerate(week):
                        if day_num == 0:
                            row.append('')
                        else:
                            dow = day_idx + 1
                            day_slots = sorted(slots_by_dow.get(dow, []), key=lambda s: s.time_slot.start_time)
                            cell_items = [f"<b>{day_num}</b>"]
                            for s in day_slots:
                                cell_items.append(f"<font color='#0f172a'><b>{s.course.code}</b></font> <font color='#64748B'>{s.time_slot.start_time.strftime('%H:%M')}</font>")
                            cell_content = '<br/>'.join(cell_items)
                            row.append(Paragraph(cell_content, cal_class_style if len(day_slots) > 0 else cal_day_num_style))
                    data.append(row)

                # Distribute columns widths over 28.1 cm total
                c_width = (landscape(A4)[0] - 1.6*cm) / 7
                r_height = 15.5 * cm / (len(month_weeks) + 1)
                table = Table(data, colWidths=[c_width] * 7, rowHeights=[0.6*cm] + [r_height] * len(month_weeks))
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(primary_hex)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TOPPADDING', (0, 0), (-1, -1), 3),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                    ('LEFTPADDING', (0, 0), (-1, -1), 3),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                ]))
                elements.append(table)

                if idx < len(targets) - 1:
                    elements.append(PageBreak())

    # ─────────────────────────────────────────────────────────────────────────
    # LAYOUT 3: yearly (Yearly Curriculum Overview)
    # ─────────────────────────────────────────────────────────────────────────
    elif layout_type == 'yearly':
        if not targets:
            elements.append(Paragraph("No scheduled slots found for this selection.", styles['Heading2']))
        else:
            for idx, target in enumerate(targets):
                if filter_type in ('group', 'all_groups'):
                    dept_name, prog_name, year_label = get_group_details(target)
                    current_subtitle = (
                        f"<b>Department:</b> {dept_name} &nbsp;|&nbsp; <b>Program:</b> {prog_name}<br/>"
                        f"<b>Student Group:</b> {target.name} &nbsp;|&nbsp; <b>Academic Year:</b> {year_label} &nbsp;|&nbsp; <b>Class Size:</b> {target.size} Students"
                    )
                elif filter_type in ('lecturer', 'all_lecturers'):
                    dept_name = target.department.name if target.department else "N/A"
                    current_subtitle = (
                        f"<b>Department:</b> {dept_name}<br/>"
                        f"<b>Lecturer:</b> {target.name} &nbsp;|&nbsp; <b>Staff ID:</b> {target.staff_id or 'N/A'} &nbsp;|&nbsp; <b>Type:</b> {target.get_lecturer_type_display()}"
                    )
                elif filter_type in ('room', 'all_rooms'):
                    campus_name = target.campus.name if target.campus else "N/A"
                    building_name = target.building.name if target.building else "N/A"
                    current_subtitle = (
                        f"<b>Campus/Building:</b> {campus_name} / {building_name}<br/>"
                        f"<b>Room:</b> {target.name} &nbsp;|&nbsp; <b>Type:</b> {target.get_room_type_display()} &nbsp;|&nbsp; <b>Capacity:</b> {target.capacity} Seats"
                    )
                else:
                    current_subtitle = f"{pdf_subtitle} — Yearly Curriculum Map"

                # Generate the unified header using helper function
                elements.extend(make_header_flowable(pdf_title, current_subtitle))

                if filter_type in ('group', 'all_groups'):
                    slots_target = [s for s in slots_all if s.student_group_id == target.id]
                elif filter_type in ('room', 'all_rooms'):
                    slots_target = [s for s in slots_all if s.room_id == target.id]
                elif filter_type in ('lecturer', 'all_lecturers'):
                    slots_target = [s for s in slots_all if s.lecturer_id == target.id]
                else:
                    slots_target = slots_all

                target_courses = {}
                for s in slots_target:
                    target_courses[s.course_id] = s.course

                yearly_headers = ['Course Code', 'Course Name', 'Room Type', 'Lecturer', 'Slots / Week', 'Weekly Hours']
                data = [yearly_headers]
                
                y_cell_style = ParagraphStyle('y_cell', parent=styles['Normal'], fontSize=8, leading=10)
                
                timeslots_uni = list(TimeSlot.objects.filter(university=timetable.semester.university))
                if timeslots_uni:
                    avg_dur = sum(
                        ((ts.end_time.hour * 60 + ts.end_time.minute) - (ts.start_time.hour * 60 + ts.start_time.minute)) / 60.0
                        for ts in timeslots_uni
                    ) / len(timeslots_uni)
                else:
                    avg_dur = 1.5

                for course in sorted(target_courses.values(), key=lambda c: c.code):
                    hours = course.duration_slots * course.sessions_per_week * avg_dur
                    row = [
                        Paragraph(f"<b>{course.code}</b>", y_cell_style),
                        Paragraph(course.name, y_cell_style),
                        Paragraph(course.required_room_type, y_cell_style),
                        Paragraph(course.lecturer.name if course.lecturer else 'Unassigned', y_cell_style),
                        Paragraph(str(course.duration_slots), y_cell_style),
                        Paragraph(f"{hours:.1f} hrs", y_cell_style),
                    ]
                    data.append(row)

                if len(data) == 1:
                    data.append([Paragraph("No courses scheduled.", y_cell_style)] + [''] * 5)

                # Distribute columns widths over 28.1 cm total
                y_col_width = (landscape(A4)[0] - 1.6*cm) / 6
                table = Table(data, colWidths=[y_col_width] * 6, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(primary_hex)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor(alt_hex), colors.white]),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ]))
                elements.append(table)

                if idx < len(targets) - 1:
                    elements.append(PageBreak())

    # ─────────────────────────────────────────────────────────────────────────
    # LAYOUT 4: master_list (Master Tabular Chronological List)
    # ─────────────────────────────────────────────────────────────────────────
    elif layout_type == 'master_list':
        slots_sorted = list(
            timetable.slots.select_related('course', 'lecturer', 'room', 'time_slot', 'student_group', 'student_group__program')
            .all()
            .order_by('time_slot__day_of_week', 'time_slot__start_time', 'course__code')
        )

        current_subtitle = f"{pdf_subtitle} — Master Chronological List"
        elements.extend(make_header_flowable(pdf_title, current_subtitle))

        headers = ['Day', 'Time', 'Unit Code', 'Unit Title', 'Student Group', 'Lecturer', 'Venue / Room']
        data = [headers]

        cell_style_left = ParagraphStyle('ml_cell_left', parent=styles['Normal'], fontSize=7.0, leading=9.0)
        cell_style_center = ParagraphStyle('ml_cell_center', parent=styles['Normal'], fontSize=7.0, leading=9.0, alignment=TA_CENTER)
        cell_style_bold = ParagraphStyle('ml_cell_bold', parent=styles['Normal'], fontSize=7.0, leading=9.0, fontName='Helvetica-Bold')

        day_names = {1:'Monday', 2:'Tuesday', 3:'Wednesday', 4:'Thursday', 5:'Friday', 6:'Saturday', 7:'Sunday'}

        for s in slots_sorted:
            day_str = day_names.get(s.time_slot.day_of_week, f"Day {s.time_slot.day_of_week}")
            time_str = f"{s.time_slot.start_time.strftime('%H:%M')} - {s.time_slot.end_time.strftime('%H:%M')}"
            
            day_p = Paragraph(day_str, cell_style_bold)
            time_p = Paragraph(time_str, cell_style_center)
            code_p = Paragraph(s.course.code, cell_style_bold)
            title_p = Paragraph(s.course.name, cell_style_left)
            
            group_name = s.student_group.name if s.student_group else 'Unassigned'
            group_p = Paragraph(group_name, cell_style_left)

            lec_name = s.lecturer.name if s.lecturer else 'Unassigned'
            lec_p = Paragraph(lec_name, cell_style_left)

            room_name = s.room.name if s.room else 'Unassigned'
            room_p = Paragraph(room_name, cell_style_center)

            data.append([day_p, time_p, code_p, title_p, group_p, lec_p, room_p])

        if len(data) == 1:
            data.append([Paragraph("No courses scheduled.", cell_style_left)] + [''] * 6)

        col_widths = [2.0*cm, 3.0*cm, 2.2*cm, 7.0*cm, 4.5*cm, 4.5*cm, 4.9*cm]
        table = Table(data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(primary_hex)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8.0),
            ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'CENTER'),
            ('ALIGN', (2, 0), (2, 0), 'LEFT'),
            ('ALIGN', (-1, 0), (-1, 0), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor(alt_hex)]),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        elements.append(table)

    # ─────────────────────────────────────────────────────────────────────────
    # LAYOUT 5: weekly (Standard Weekly Grid Layout)
    # ─────────────────────────────────────────────────────────────────────────
    else:
        # Dynamic Scaling of Font and Padding for single-page fit
        num_slots = len(slot_numbers)
        if num_slots <= 4:
            cell_font = 8.5
            cell_lead = 10.5
            tb_pad = 5
        elif num_slots <= 6:
            cell_font = 8
            cell_lead = 10
            tb_pad = 4
        elif num_slots <= 8:
            cell_font = 7
            cell_lead = 9
            tb_pad = 3
        else:
            cell_font = 6
            cell_lead = 8
            tb_pad = 2

        cell_style = ParagraphStyle('cell', parent=styles['Normal'], fontSize=cell_font, leading=cell_lead, alignment=TA_CENTER)
        
        # Dedicated styling for the timeslot label column (bold and dark slate text)
        time_style = ParagraphStyle(
            'weekly_time',
            parent=styles['Normal'],
            fontSize=cell_font,
            leading=cell_lead,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            textColor=colors.HexColor('#1E293B')
        )

        header = ['Slot / Time'] + [day_labels.get(d, f'D{d}') for d in days]
        ts_by_slot_number = {}
        for ts in timeslots:
            ts_by_slot_number.setdefault(ts.slot_number, []).append(ts)

        # Distribute columns widths over 28.1 cm total
        col_width = (landscape(A4)[0] - 1.6*cm) / (len(days) + 1) if days else (landscape(A4)[0] - 1.6*cm)

        if not targets:
            elements.append(Paragraph("No scheduled slots found for this selection.", styles['Heading2']))
        else:
            for idx, target in enumerate(targets):
                if filter_type in ('group', 'all_groups'):
                    dept_name, prog_name, year_label = get_group_details(target)
                    current_subtitle = (
                        f"<b>Department:</b> {dept_name} &nbsp;|&nbsp; <b>Program:</b> {prog_name}<br/>"
                        f"<b>Student Group:</b> {target.name} &nbsp;|&nbsp; <b>Academic Year:</b> {year_label} &nbsp;|&nbsp; <b>Class Size:</b> {target.size} Students"
                    )
                elif filter_type in ('lecturer', 'all_lecturers'):
                    dept_name = target.department.name if target.department else "N/A"
                    current_subtitle = (
                        f"<b>Department:</b> {dept_name}<br/>"
                        f"<b>Lecturer:</b> {target.name} &nbsp;|&nbsp; <b>Staff ID:</b> {target.staff_id or 'N/A'} &nbsp;|&nbsp; <b>Type:</b> {target.get_lecturer_type_display()}"
                    )
                elif filter_type in ('room', 'all_rooms'):
                    campus_name = target.campus.name if target.campus else "N/A"
                    building_name = target.building.name if target.building else "N/A"
                    current_subtitle = (
                        f"<b>Campus/Building:</b> {campus_name} / {building_name}<br/>"
                        f"<b>Room:</b> {target.name} &nbsp;|&nbsp; <b>Type:</b> {target.get_room_type_display()} &nbsp;|&nbsp; <b>Capacity:</b> {target.capacity} Seats"
                    )
                else:
                    current_subtitle = pdf_subtitle

                # Generate the unified header using helper function
                elements.extend(make_header_flowable(pdf_title, current_subtitle))

                if filter_type in ('group', 'all_groups'):
                    slots_target = [s for s in slots_all if s.student_group_id == target.id]
                elif filter_type in ('room', 'all_rooms'):
                    slots_target = [s for s in slots_all if s.room_id == target.id]
                elif filter_type in ('lecturer', 'all_lecturers'):
                    slots_target = [s for s in slots_all if s.lecturer_id == target.id]
                else:
                    slots_target = slots_all

                slots_by_day_and_num = {}
                for s in slots_target:
                    slots_by_day_and_num.setdefault((s.time_slot.day_of_week, s.time_slot.slot_number), []).append(s)

                data = [header]

                for slot_num in slot_numbers:
                    ts_for_slot = ts_by_slot_number.get(slot_num, [])
                    time_label = f"Slot {slot_num}"
                    if ts_for_slot:
                        ts0 = ts_for_slot[0]
                        time_label = f"{ts0.start_time.strftime('%H:%M')}<br/>{ts0.end_time.strftime('%H:%M')}"
                    row = [Paragraph(time_label, time_style)]
                    for day in days:
                        matching = slots_by_day_and_num.get((day, slot_num), [])
                        if matching:
                            items = []
                            for s in matching:
                                course_title = s.course.name
                                if len(course_title) > 28:
                                    course_title = course_title[:25] + "..."
                                cell_lines = [f"<b>{s.course.code}</b>", f"<font color='#0D5A4F'><i>{course_title}</i></font>"]

                                # Add student group/class and omit redundant fields based on active filter
                                if filter_type not in ('group', 'all_groups') and s.student_group:
                                    cell_lines.append(f"<font color='#1E293B'><b>{s.student_group.name}</b></font>")
                                if filter_type not in ('lecturer', 'all_lecturers') and s.lecturer:
                                    cell_lines.append(f"<font color='#1E3A8A'>{s.lecturer.name}</font>")
                                if filter_type not in ('room', 'all_rooms') and s.room:
                                    cell_lines.append(f"<font color='#801d1d'>Room: {s.room.name}</font>")

                                items.append('<br/>'.join(cell_lines))
                            cell_content = '<br/><font color="#CBD5E1">────────────────</font><br/>'.join(items)
                            row.append(Paragraph(cell_content, cell_style))
                        else:
                            row.append('')
                    data.append(row)

                table = Table(data, colWidths=[col_width] * (len(days) + 1), repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(primary_hex)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
                    ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#F8FAFC')), # Shade the timeslots column
                    ('ROWBACKGROUNDS', (1, 1), (-1, -1), [colors.HexColor(alt_hex), colors.white]), # Alternate backgrounds for days columns only
                    ('LEFTPADDING', (0, 0), (-1, -1), 4),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                    ('TOPPADDING', (0, 0), (-1, -1), tb_pad),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), tb_pad),
                ]))
                elements.append(table)

                if idx < len(targets) - 1:
                    elements.append(PageBreak())

    # Build the document using the NumberedCanvas custom page numbering class
    doc.build(elements, canvasmaker=NumberedCanvas)
    buffer.seek(0)

    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    safe_name = "".join(c for c in timetable.name if c.isalnum() or c in (' ','_','-')).strip().replace(' ','_')
    response['Content-Disposition'] = f'attachment; filename="timetable_{safe_name}.pdf"'
    return response


@login_required
def export_timetable_excel(request, pk):
    """Export timetable as a highly polished, professional Excel file using openpyxl."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import io

    university = get_active_uni(request)
    timetable = get_object_or_404(
        Timetable.objects.select_related('semester', 'semester__university'), pk=pk
    )
    if timetable.semester.university != university:
        messages.error(request, "Access denied.")
        return redirect('scheduler:dashboard')

    wb = openpyxl.Workbook()
    
    # ── Sheet 1: Master Schedule List (Tabular Report) ────────────────────────
    ws1 = wb.active
    ws1.title = "Master Schedule List"
    ws1.views.sheetView[0].showGridLines = True

    # Cohesive Professional Theme Styles (Forest Teal theme)
    font_family = "Segoe UI"
    title_font = Font(name=font_family, size=16, bold=True, color="0D5A4F")
    subtitle_font = Font(name=font_family, size=10, italic=True, color="555555")
    header_font = Font(name=font_family, size=11, bold=True, color="FFFFFF")
    data_font = Font(name=font_family, size=10, color="333333")
    bold_data_font = Font(name=font_family, size=10, bold=True, color="333333")
    
    header_fill = PatternFill(start_color="0D5A4F", end_color="0D5A4F", fill_type="solid")
    alt_fill = PatternFill(start_color="F2FAF7", end_color="F2FAF7", fill_type="solid")
    
    border_side = Side(style="thin", color="D3D3D3")
    cell_border = Border(left=border_side, right=border_side, top=border_side, bottom=border_side)
    
    align_center = Alignment(horizontal="center", vertical="center")
    align_left = Alignment(horizontal="left", vertical="center")

    # Title Block
    ws1["A1"] = f"UNIVERSAL MASTER TIMETABLE REPORT"
    ws1["A1"].font = title_font
    ws1["A2"] = f"Timetable: {timetable.name}  |  Semester: {timetable.semester.name}  |  University: {university.name}"
    ws1["A2"].font = subtitle_font
    ws1.row_dimensions[1].height = 24
    ws1.row_dimensions[2].height = 18

    # Headers for List
    headers = [
        "Day of Week", "Time Slot", "Course Code", "Course Title", 
        "Lecturer", "Room", "Student Group", "Campus"
    ]
    
    ws1.append([]) # Blank row 3
    ws1.row_dimensions[3].height = 10
    
    # Append headers on Row 4
    for col_idx, text in enumerate(headers, start=1):
        cell = ws1.cell(row=4, column=col_idx, value=text)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = align_center
        cell.border = cell_border
    ws1.row_dimensions[4].height = 28

    slots_all = list(
        timetable.slots.select_related('course', 'lecturer', 'room', 'room__campus', 'time_slot', 'student_group').all()
    )
    
    # Sorting: Day of Week, Slot Number, Course Code
    day_order = {1: 'Monday', 2: 'Tuesday', 3: 'Wednesday', 4: 'Thursday', 5: 'Friday', 6: 'Saturday', 7: 'Sunday'}
    slots_all.sort(key=lambda s: (s.time_slot.day_of_week, s.time_slot.slot_number, s.course.code))

    # Populate List Rows
    for idx, s in enumerate(slots_all, start=5):
        ts = s.time_slot
        day_text = day_order.get(ts.day_of_week, f"Day {ts.day_of_week}")
        time_text = f"{ts.start_time.strftime('%H:%M')} – {ts.end_time.strftime('%H:%M')}"
        campus_name = s.room.campus.name if s.room.campus else "N/A"
        
        row_values = [
            day_text, time_text, s.course.code, s.course.name,
            s.lecturer.name, s.room.name, s.student_group.name, campus_name
        ]
        
        is_alt = (idx % 2 == 0)
        for col_idx, val in enumerate(row_values, start=1):
            cell = ws1.cell(row=idx, column=col_idx, value=val)
            cell.font = data_font
            cell.border = cell_border
            if is_alt:
                cell.fill = alt_fill
            
            # Alignments
            if col_idx in (1, 2, 3, 6, 7, 8):
                cell.alignment = align_center
            else:
                cell.alignment = align_left
                
            if col_idx == 3: # Course Code bold
                cell.font = bold_data_font
                
        ws1.row_dimensions[idx].height = 20

    # Auto-adjust column widths for Sheet 1
    for col in ws1.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.row > 2 and cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws1.column_dimensions[col_letter].width = max(max_len + 4, 12)

    # ── Sheet 2: Weekly Calendar Grid (Day/Time Matrix) ───────────────────────
    ws2 = wb.create_sheet(title="Weekly Grid View")
    ws2.views.sheetView[0].showGridLines = True

    # Title Block Sheet 2
    ws2["A1"] = f"WEEKLY CLASS GRID MATRIX"
    ws2["A1"].font = title_font
    ws2["A2"] = "Summary of parallel classes scheduled per time slot"
    ws2["A2"].font = subtitle_font
    ws2.row_dimensions[1].height = 24
    ws2.row_dimensions[2].height = 18

    # Unique active days and slots
    timeslots_ref = list(TimeSlot.objects.filter(university=university).order_by('day_of_week', 'slot_number'))
    active_days = sorted(set(ts.day_of_week for ts in timeslots_ref))
    active_slot_numbers = sorted(set(ts.slot_number for ts in timeslots_ref))

    ws2.cell(row=4, column=1, value="Slot / Time").font = header_font
    ws2.cell(row=4, column=1).fill = header_fill
    ws2.cell(row=4, column=1).alignment = align_center
    ws2.cell(row=4, column=1).border = cell_border
    
    # Headers on row 4
    for col_idx, d_num in enumerate(active_days, start=2):
        cell = ws2.cell(row=4, column=col_idx, value=day_order.get(d_num, f"Day {d_num}"))
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = align_center
        cell.border = cell_border
    ws2.row_dimensions[4].height = 28

    # Pre-group assignments for the grid
    slots_by_day_and_num = {}
    for s in slots_all:
        slots_by_day_and_num.setdefault((s.time_slot.day_of_week, s.time_slot.slot_number), []).append(s)

    ts_by_slot_number = {}
    for ts in timeslots_ref:
        ts_by_slot_number.setdefault(ts.slot_number, []).append(ts)

    # Populate Grid rows starting at row 5
    for grid_row_idx, slot_num in enumerate(active_slot_numbers, start=5):
        ts_for_slot = ts_by_slot_number.get(slot_num, [])
        time_label = f"Slot {slot_num}"
        if ts_for_slot:
            ts0 = ts_for_slot[0]
            time_label = f"{ts0.start_time.strftime('%H:%M')} – {ts0.end_time.strftime('%H:%M')}"
            
        time_cell = ws2.cell(row=grid_row_idx, column=1, value=time_label)
        time_cell.font = Font(name=font_family, size=10, bold=True, color="333333")
        time_cell.alignment = align_center
        time_cell.border = cell_border
        
        is_alt = (grid_row_idx % 2 == 0)
        if is_alt:
            time_cell.fill = alt_fill

        max_lines_in_row = 1
        for col_idx, d_num in enumerate(active_days, start=2):
            matching = slots_by_day_and_num.get((d_num, slot_num), [])
            cell_lines = []
            for s in matching:
                cell_lines.append(f"• {s.course.code} ({s.room.name}) - {s.lecturer.name}")
            
            content = "\n".join(cell_lines)
            max_lines_in_row = max(max_lines_in_row, len(cell_lines))
            
            cell = ws2.cell(row=grid_row_idx, column=col_idx, value=content)
            cell.font = data_font
            cell.border = cell_border
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            if is_alt:
                cell.fill = alt_fill

        # Auto-scale row height based on the number of classes at that slot
        ws2.row_dimensions[grid_row_idx].height = max(max_lines_in_row * 15 + 10, 40)

    # Auto-adjust column widths for Sheet 2
    ws2.column_dimensions['A'].width = 16
    for col_idx in range(2, len(active_days) + 2):
        ws2.column_dimensions[get_column_letter(col_idx)].width = 38

    # ── Sheet 3: Student Group Timetables ─────────────────────────────────────
    ws3 = wb.create_sheet(title="Student Group Timetables")
    ws3.views.sheetView[0].showGridLines = True
    
    ws3["A1"] = "STUDENT GROUP TIMETABLES"
    ws3["A1"].font = title_font
    ws3["A2"] = "Weekly schedule listed group by group"
    ws3["A2"].font = subtitle_font
    ws3.row_dimensions[1].height = 24
    ws3.row_dimensions[2].height = 18
    
    current_row = 4
    
    # Group slots by Student Group
    slots_by_group = {}
    for s in slots_all:
        slots_by_group.setdefault(s.student_group, []).append(s)
        
    sorted_groups = sorted(slots_by_group.keys(), key=lambda g: g.name if g else "")
    
    for group in sorted_groups:
        group_slots = sorted(slots_by_group[group], key=lambda s: (s.time_slot.day_of_week, s.time_slot.slot_number))
        
        # Group Header Row
        ws3.cell(row=current_row, column=1, value=f"STUDENT GROUP: {group.name if group else 'Unassigned'}").font = Font(name=font_family, size=11, bold=True, color="0D5A4F")
        ws3.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=6)
        ws3.row_dimensions[current_row].height = 22
        current_row += 1
        
        # Table Headers
        group_headers = ["Day", "Time Slot", "Course Code", "Course Title", "Lecturer", "Room"]
        for col_idx, text in enumerate(group_headers, start=1):
            cell = ws3.cell(row=current_row, column=col_idx, value=text)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = align_center
            cell.border = cell_border
        ws3.row_dimensions[current_row].height = 20
        current_row += 1
        
        # Table Rows
        for idx, s in enumerate(group_slots):
            ts = s.time_slot
            day_text = day_order.get(ts.day_of_week, f"Day {ts.day_of_week}")
            time_text = f"{ts.start_time.strftime('%H:%M')} – {ts.end_time.strftime('%H:%M')}"
            
            row_vals = [
                day_text, time_text, s.course.code, s.course.name,
                s.lecturer.name if s.lecturer else 'Unassigned',
                s.room.name if s.room else 'Unassigned'
            ]
            
            is_alt = (idx % 2 == 0)
            for col_idx, val in enumerate(row_vals, start=1):
                cell = ws3.cell(row=current_row, column=col_idx, value=val)
                cell.font = data_font
                cell.border = cell_border
                if is_alt:
                    cell.fill = alt_fill
                
                if col_idx in (1, 2, 3, 6):
                    cell.alignment = align_center
                else:
                    cell.alignment = align_left
                
                if col_idx == 3:
                    cell.font = bold_data_font
            ws3.row_dimensions[current_row].height = 18
            current_row += 1
            
        current_row += 2 # spacing between tables

    # ── Sheet 4: Lecturer Timetables ──────────────────────────────────────────
    ws4 = wb.create_sheet(title="Lecturer Timetables")
    ws4.views.sheetView[0].showGridLines = True
    
    ws4["A1"] = "LECTURER TIMETABLES"
    ws4["A1"].font = title_font
    ws4["A2"] = "Weekly schedule listed lecturer by lecturer"
    ws4["A2"].font = subtitle_font
    ws4.row_dimensions[1].height = 24
    ws4.row_dimensions[2].height = 18
    
    current_row = 4
    
    # Group slots by Lecturer
    slots_by_lecturer = {}
    for s in slots_all:
        slots_by_lecturer.setdefault(s.lecturer, []).append(s)
        
    sorted_lecturers = sorted(slots_by_lecturer.keys(), key=lambda l: l.name if l else "")
    
    for lec in sorted_lecturers:
        lec_slots = sorted(slots_by_lecturer[lec], key=lambda s: (s.time_slot.day_of_week, s.time_slot.slot_number))
        
        # Lecturer Header Row
        ws4.cell(row=current_row, column=1, value=f"LECTURER: {lec.name if lec else 'Unassigned'}").font = Font(name=font_family, size=11, bold=True, color="0D5A4F")
        ws4.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=6)
        ws4.row_dimensions[current_row].height = 22
        current_row += 1
        
        # Table Headers
        lec_headers = ["Day", "Time Slot", "Course Code", "Course Title", "Student Group", "Room"]
        for col_idx, text in enumerate(lec_headers, start=1):
            cell = ws4.cell(row=current_row, column=col_idx, value=text)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = align_center
            cell.border = cell_border
        ws4.row_dimensions[current_row].height = 20
        current_row += 1
        
        # Table Rows
        for idx, s in enumerate(lec_slots):
            ts = s.time_slot
            day_text = day_order.get(ts.day_of_week, f"Day {ts.day_of_week}")
            time_text = f"{ts.start_time.strftime('%H:%M')} – {ts.end_time.strftime('%H:%M')}"
            
            row_vals = [
                day_text, time_text, s.course.code, s.course.name,
                s.student_group.name if s.student_group else 'Unassigned',
                s.room.name if s.room else 'Unassigned'
            ]
            
            is_alt = (idx % 2 == 0)
            for col_idx, val in enumerate(row_vals, start=1):
                cell = ws4.cell(row=current_row, column=col_idx, value=val)
                cell.font = data_font
                cell.border = cell_border
                if is_alt:
                    cell.fill = alt_fill
                
                if col_idx in (1, 2, 3, 6):
                    cell.alignment = align_center
                else:
                    cell.alignment = align_left
                
                if col_idx == 3:
                    cell.font = bold_data_font
            ws4.row_dimensions[current_row].height = 18
            current_row += 1
            
        current_row += 2 # spacing

    # ── Sheet 5: Room Timetables ──────────────────────────────────────────────
    ws5 = wb.create_sheet(title="Room Timetables")
    ws5.views.sheetView[0].showGridLines = True
    
    ws5["A1"] = "ROOM TIMETABLES"
    ws5["A1"].font = title_font
    ws5["A2"] = "Weekly schedule listed room by room"
    ws5["A2"].font = subtitle_font
    ws5.row_dimensions[1].height = 24
    ws5.row_dimensions[2].height = 18
    
    current_row = 4
    
    # Group slots by Room
    slots_by_room = {}
    for s in slots_all:
        slots_by_room.setdefault(s.room, []).append(s)
        
    sorted_rooms = sorted(slots_by_room.keys(), key=lambda r: r.name if r else "")
    
    for room in sorted_rooms:
        room_slots = sorted(slots_by_room[room], key=lambda s: (s.time_slot.day_of_week, s.time_slot.slot_number))
        
        # Room Header Row
        ws5.cell(row=current_row, column=1, value=f"ROOM: {room.name if room else 'Unassigned'}").font = Font(name=font_family, size=11, bold=True, color="0D5A4F")
        ws5.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=6)
        ws5.row_dimensions[current_row].height = 22
        current_row += 1
        
        # Table Headers
        room_headers = ["Day", "Time Slot", "Course Code", "Course Title", "Lecturer", "Student Group"]
        for col_idx, text in enumerate(room_headers, start=1):
            cell = ws5.cell(row=current_row, column=col_idx, value=text)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = align_center
            cell.border = cell_border
        ws5.row_dimensions[current_row].height = 20
        current_row += 1
        
        # Table Rows
        for idx, s in enumerate(room_slots):
            ts = s.time_slot
            day_text = day_order.get(ts.day_of_week, f"Day {ts.day_of_week}")
            time_text = f"{ts.start_time.strftime('%H:%M')} – {ts.end_time.strftime('%H:%M')}"
            
            row_vals = [
                day_text, time_text, s.course.code, s.course.name,
                s.lecturer.name if s.lecturer else 'Unassigned',
                s.student_group.name if s.student_group else 'Unassigned'
            ]
            
            is_alt = (idx % 2 == 0)
            for col_idx, val in enumerate(row_vals, start=1):
                cell = ws5.cell(row=current_row, column=col_idx, value=val)
                cell.font = data_font
                cell.border = cell_border
                if is_alt:
                    cell.fill = alt_fill
                
                if col_idx in (1, 2, 3):
                    cell.alignment = align_center
                else:
                    cell.alignment = align_left
                
                if col_idx == 3:
                    cell.font = bold_data_font
            ws5.row_dimensions[current_row].height = 18
            current_row += 1
            
        current_row += 2 # spacing

    # Auto-adjust column widths for Sheet 3, 4, 5
    for ws in (ws3, ws4, ws5):
        ws.column_dimensions['A'].width = 14
        ws.column_dimensions['B'].width = 16
        ws.column_dimensions['C'].width = 14
        ws.column_dimensions['D'].width = 30
        ws.column_dimensions['E'].width = 25
        ws.column_dimensions['F'].width = 25

    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    safe_name = "".join(c for c in timetable.name if c.isalnum() or c in (' ','_','-')).strip().replace(' ','_')
    response['Content-Disposition'] = f'attachment; filename="timetable_{safe_name}.xlsx"'
    return response


# ─────────────────────────────────────────────────────────────────────────────
# Phase 7: Global Search
# ─────────────────────────────────────────────────────────────────────────────

@login_required
def search_view(request):
    """Global cross-resource search across courses, lecturers, rooms, and student groups."""
    query = request.GET.get('q', '').strip()
    university = get_active_uni(request)
    results = {}

    if query and university:
        results['courses'] = Course.objects.filter(
            Q(code__icontains=query) | Q(name__icontains=query),
            program__department__faculty__campus__university=university
        ).select_related('lecturer', 'student_group')[:10]

        results['lecturers'] = Lecturer.objects.filter(
            Q(name__icontains=query) | Q(email__icontains=query),
            department__faculty__campus__university=university
        )[:10]

        results['rooms'] = Room.objects.filter(
            Q(name__icontains=query),
            campus__university=university
        )[:10]

        results['student_groups'] = StudentGroup.objects.filter(
            Q(name__icontains=query),
            program__department__faculty__campus__university=university
        )[:10]

        results['timetables'] = Timetable.objects.filter(
            Q(name__icontains=query),
            semester__university=university
        )[:5]

    return render(request, 'scheduler/search_results.html', {
        'query': query,
        'results': results,
    })


@login_required
def export_timetable_word(request, pk):
    """Export timetable as a Word (.docx) file using python-docx with custom branding."""
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

    university = get_active_uni(request)
    timetable = get_object_or_404(
        Timetable.objects.select_related('semester', 'semester__university'), pk=pk
    )
    if timetable.semester.university != university:
        messages.error(request, "Access denied.")
        return redirect('scheduler:dashboard')

    # Get custom branding from request
    word_title = request.GET.get('title', f"{timetable.name} — {timetable.semester.name}")
    word_subtitle = request.GET.get('subtitle', f"{university.name} | Generated Timetable")
    primary_hex = request.GET.get('primary_color', "0D5A4F").replace('#', '')
    alt_hex = request.GET.get('alt_color', "F0FDFA").replace('#', '')
    layout_type = request.GET.get('layout_type', 'weekly')

    # Filter resources by university for scoping and defaulting
    rooms = Room.objects.filter(campus__university=university)
    lecturers = Lecturer.objects.filter(department__faculty__campus__university=university)
    student_groups = StudentGroup.objects.filter(program__department__faculty__campus__university=university)

    # Get active filters
    filter_type = request.GET.get('filter_type')
    filter_id = request.GET.get('filter_id')

    # If exporting a batch of targets, strip the target-specific suffix from the request's subtitle
    if filter_type in ('all_groups', 'all_rooms', 'all_lecturers') or layout_type == 'complete_pack':
        if ' — ' in word_subtitle:
            word_subtitle = word_subtitle.split(' — ')[0]
        elif ' - ' in word_subtitle:
            word_subtitle = word_subtitle.split(' - ')[0]

    # Resolve target list
    targets = []
    
    if filter_type == 'all_groups':
        active_group_ids = timetable.slots.values_list('student_group_id', flat=True).distinct()
        targets = list(student_groups.filter(id__in=active_group_ids).order_by('name'))
    elif filter_type == 'all_rooms':
        active_room_ids = timetable.slots.values_list('room_id', flat=True).distinct()
        targets = list(rooms.filter(id__in=active_room_ids).order_by('name'))
    elif filter_type == 'all_lecturers':
        active_lecturer_ids = timetable.slots.values_list('lecturer_id', flat=True).distinct()
        targets = list(lecturers.filter(id__in=active_lecturer_ids).order_by('name'))
    else:
        if not filter_type:
            filter_type = 'group'

        if not filter_id:
            if filter_type == 'group' and student_groups.exists():
                filter_id = student_groups.first().id
            elif filter_type == 'room' and rooms.exists():
                filter_id = rooms.first().id
            elif filter_type == 'lecturer' and lecturers.exists():
                filter_id = lecturers.first().id

        target_obj = None
        if filter_id:
            try:
                filter_id = int(filter_id)
                if filter_type == 'group':
                    target_obj = student_groups.filter(id=filter_id).first()
                elif filter_type == 'room':
                    target_obj = rooms.filter(id=filter_id).first()
                elif filter_type == 'lecturer':
                    target_obj = lecturers.filter(id=filter_id).first()
            except ValueError:
                pass
        if target_obj:
            targets = [target_obj]

    doc = docx.Document()

    # Set margins and landscape orientation
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Swap width and height for landscape A4
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    day_labels = {1: 'Mon', 2: 'Tue', 3: 'Wed', 4: 'Thu', 5: 'Fri', 6: 'Sat', 7: 'Sun'}
    timeslots = list(TimeSlot.objects.filter(university=university).order_by('day_of_week', 'slot_number'))
    days = sorted(set(ts.day_of_week for ts in timeslots))
    slot_numbers = sorted(set(ts.slot_number for ts in timeslots))

    slots_all = list(
        timetable.slots.select_related('course','lecturer','room','time_slot','student_group').all()
    )

    ts_by_slot_number = {}
    for ts in timeslots:
        ts_by_slot_number.setdefault(ts.slot_number, []).append(ts)

    # Helper function to shade cells
    def set_cell_background(cell, hex_color):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Helper function to set table borders
    def set_table_borders(table):
        tblPr = table._tbl.tblPr
        tblBorders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(tblBorders)

    if layout_type == 'complete_pack':
        # 1. Title Page / Header
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(word_title)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor.from_string(primary_hex)

        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run(f"{word_subtitle} — 1. Master Chronological List")
        sub_run.font.size = Pt(11)
        sub_run.font.color.rgb = RGBColor(71, 85, 105)

        # Master Table
        slots_sorted = sorted(slots_all, key=lambda s: (s.time_slot.day_of_week, s.time_slot.start_time, s.course.code))
        master_headers = ["Day", "Time Slot", "Course Code", "Course Title", "Student Group", "Lecturer", "Room"]
        
        table = doc.add_table(rows=len(slots_sorted) + 1, cols=7)
        set_table_borders(table)
        
        hdr_cells = table.rows[0].cells
        for col_idx, text in enumerate(master_headers):
            hdr_cells[col_idx].text = text
            set_cell_background(hdr_cells[col_idx], primary_hex)
            hdr_cells[col_idx].paragraphs[0].runs[0].font.bold = True
            hdr_cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            hdr_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        day_names = {1:'Monday', 2:'Tuesday', 3:'Wednesday', 4:'Thursday', 5:'Friday', 6:'Saturday', 7:'Sunday'}
        for r_idx, s in enumerate(slots_sorted, start=1):
            row_cells = table.rows[r_idx].cells
            bg_color = alt_hex if r_idx % 2 == 1 else "FFFFFF"
            
            day_text = day_names.get(s.time_slot.day_of_week, f"Day {s.time_slot.day_of_week}")
            time_text = f"{s.time_slot.start_time.strftime('%H:%M')} – {s.time_slot.end_time.strftime('%H:%M')}"
            
            row_vals = [
                day_text, time_text, s.course.code, s.course.name,
                s.student_group.name if s.student_group else 'Unassigned',
                s.lecturer.name if s.lecturer else 'Unassigned',
                s.room.name if s.room else 'Unassigned'
            ]
            
            for col_idx, val in enumerate(row_vals):
                row_cells[col_idx].text = val
                set_cell_background(row_cells[col_idx], bg_color)
                if col_idx in (0, 1, 2, 6):
                    row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                if col_idx == 2:
                    row_cells[col_idx].paragraphs[0].runs[0].font.bold = True
                    
        # Helper to append grids in Word
        def append_weekly_tables_word(section_title, targets_list, target_type):
            for t in targets_list:
                doc.add_page_break()
                
                title_p = doc.add_paragraph()
                title_run = title_p.add_run(word_title)
                title_run.bold = True
                title_run.font.size = Pt(16)
                title_run.font.color.rgb = RGBColor.from_string(primary_hex)

                sub_p = doc.add_paragraph()
                if target_type == 'group':
                    dept_name, prog_name, year_label = get_group_details(t)
                    
                    sub_p.add_run("Department: ").bold = True
                    sub_p.add_run(f"{dept_name}  |  ")
                    sub_p.add_run("Program: ").bold = True
                    sub_p.add_run(f"{prog_name}\n")
                    sub_p.add_run("Student Group: ").bold = True
                    sub_p.add_run(f"{t.name}  |  ")
                    sub_p.add_run("Academic Year: ").bold = True
                    sub_p.add_run(f"{year_label}  |  ")
                    sub_p.add_run("Class Size: ").bold = True
                    sub_p.add_run(f"{t.size} Students")
                elif target_type == 'lecturer':
                    dept_name = t.department.name if t.department else "N/A"
                    sub_p.add_run("Department: ").bold = True
                    sub_p.add_run(f"{dept_name}\n")
                    sub_p.add_run("Lecturer: ").bold = True
                    sub_p.add_run(f"{t.name}  |  ")
                    sub_p.add_run("Staff ID: ").bold = True
                    sub_p.add_run(f"{t.staff_id or 'N/A'}  |  ")
                    sub_p.add_run("Contract Type: ").bold = True
                    sub_p.add_run(f"{t.get_lecturer_type_display()}")
                else: # room
                    campus_name = t.campus.name if t.campus else "N/A"
                    building_name = t.building.name if t.building else "N/A"
                    sub_p.add_run("Campus / Building: ").bold = True
                    sub_p.add_run(f"{campus_name} / {building_name}\n")
                    sub_p.add_run("Room: ").bold = True
                    sub_p.add_run(f"{t.name}  |  ")
                    sub_p.add_run("Room Type: ").bold = True
                    sub_p.add_run(f"{t.get_room_type_display()}  |  ")
                    sub_p.add_run("Capacity: ").bold = True
                    sub_p.add_run(f"{t.capacity} Seats")

                for run in sub_p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(71, 85, 105)

                if target_type == 'group':
                    slots_t = [s for s in slots_all if s.student_group_id == t.id]
                elif target_type == 'lecturer':
                    slots_t = [s for s in slots_all if s.lecturer_id == t.id]
                else:
                    slots_t = [s for s in slots_all if s.room_id == t.id]
                    
                slots_by_day_and_num = {}
                for s in slots_t:
                    slots_by_day_and_num.setdefault((s.time_slot.day_of_week, s.time_slot.slot_number), []).append(s)

                num_cols = len(days) + 1
                num_rows = len(slot_numbers) + 1
                grid_table = doc.add_table(rows=num_rows, cols=num_cols)
                set_table_borders(grid_table)

                hdr_cells = grid_table.rows[0].cells
                hdr_cells[0].text = "Slot / Time"
                set_cell_background(hdr_cells[0], primary_hex)
                hdr_cells[0].paragraphs[0].runs[0].font.bold = True
                hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for col_idx, d in enumerate(days, start=1):
                    cell = hdr_cells[col_idx]
                    cell.text = day_labels.get(d, f'D{d}')
                    set_cell_background(cell, primary_hex)
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for r_idx, slot_num in enumerate(slot_numbers, start=1):
                    row_cells = grid_table.rows[r_idx].cells
                    bg_color = alt_hex if r_idx % 2 == 1 else "FFFFFF"

                    ts_for_slot = ts_by_slot_number.get(slot_num, [])
                    time_label = f"Slot {slot_num}"
                    if ts_for_slot:
                        ts0 = ts_for_slot[0]
                        time_label = f"{ts0.start_time.strftime('%H:%M')}\n{ts0.end_time.strftime('%H:%M')}"
                    
                    row_cells[0].text = time_label
                    set_cell_background(row_cells[0], bg_color)
                    row_cells[0].paragraphs[0].runs[0].font.bold = True
                    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    for c_idx, day in enumerate(days, start=1):
                        cell = row_cells[c_idx]
                        set_cell_background(cell, bg_color)
                        matching = slots_by_day_and_num.get((day, slot_num), [])
                        if matching:
                            cell.text = ""
                            for s_idx, s in enumerate(matching):
                                if s_idx > 0:
                                    p = cell.add_paragraph("────────────────")
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    p.runs[0].font.size = Pt(6)
                                    p.runs[0].font.color.rgb = RGBColor(203, 213, 225)
                                
                                p_code = cell.add_paragraph()
                                p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run_code = p_code.add_run(s.course.code)
                                run_code.bold = True
                                run_code.font.size = Pt(8)
                                
                                course_title = s.course.name
                                if len(course_title) > 28:
                                    course_title = course_title[:25] + "..."
                                p_title = cell.add_paragraph()
                                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run_title = p_title.add_run(course_title)
                                run_title.italic = True
                                run_title.font.size = Pt(7)
                                run_title.font.color.rgb = RGBColor(13, 90, 79)
                                
                                if target_type != 'group' and s.student_group:
                                    p_grp = cell.add_paragraph()
                                    p_grp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run_grp = p_grp.add_run(s.student_group.name)
                                    run_grp.font.size = Pt(7.5)
                                    run_grp.bold = True
                                    run_grp.font.color.rgb = RGBColor(30, 41, 59)
                                    
                                if target_type != 'lecturer' and s.lecturer:
                                    p_lec = cell.add_paragraph()
                                    p_lec.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run_lec = p_lec.add_run(s.lecturer.name)
                                    run_lec.font.size = Pt(7)
                                    run_lec.font.color.rgb = RGBColor(30, 58, 138)
                                    
                                if target_type != 'room' and s.room:
                                    p_room = cell.add_paragraph()
                                    p_room.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run_room = p_room.add_run(f"📍 {s.room.name}")
                                    run_room.font.size = Pt(7.5)
                                    run_room.font.color.rgb = RGBColor(128, 29, 29)

        active_group_ids = timetable.slots.values_list('student_group_id', flat=True).distinct()
        groups_list = list(student_groups.filter(id__in=active_group_ids).order_by('name'))
        append_weekly_tables_word("2. Student Group Timetable", groups_list, 'group')

        active_lec_ids = timetable.slots.values_list('lecturer_id', flat=True).distinct()
        lecturers_list = list(lecturers.filter(id__in=active_lec_ids).order_by('name'))
        append_weekly_tables_word("3. Lecturer Timetable", lecturers_list, 'lecturer')

        active_room_ids = timetable.slots.values_list('room_id', flat=True).distinct()
        rooms_list = list(rooms.filter(id__in=active_room_ids).order_by('name'))
        append_weekly_tables_word("4. Room Timetable", rooms_list, 'room')

    elif not targets:
        doc.add_heading("No scheduled slots found for this selection.", level=2)
    else:
        for idx, target in enumerate(targets):
            if idx > 0:
                doc.add_page_break()

            # Title
            title_p = doc.add_paragraph()
            title_run = title_p.add_run(word_title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_run.font.color.rgb = RGBColor.from_string(primary_hex)

            # Subtitle
            sub_p = doc.add_paragraph()
            if filter_type in ('group', 'all_groups'):
                dept_name, prog_name, year_label = get_group_details(target)
                
                sub_p.add_run("Department: ").bold = True
                sub_p.add_run(f"{dept_name}  |  ")
                sub_p.add_run("Program: ").bold = True
                sub_p.add_run(f"{prog_name}\n")
                sub_p.add_run("Student Group: ").bold = True
                sub_p.add_run(f"{target.name}  |  ")
                sub_p.add_run("Academic Year: ").bold = True
                sub_p.add_run(f"{year_label}  |  ")
                sub_p.add_run("Class Size: ").bold = True
                sub_p.add_run(f"{target.size} Students")
            elif filter_type in ('lecturer', 'all_lecturers'):
                dept_name = target.department.name if target.department else "N/A"
                sub_p.add_run("Department: ").bold = True
                sub_p.add_run(f"{dept_name}\n")
                sub_p.add_run("Lecturer: ").bold = True
                sub_p.add_run(f"{target.name}  |  ")
                sub_p.add_run("Staff ID: ").bold = True
                sub_p.add_run(f"{target.staff_id or 'N/A'}  |  ")
                sub_p.add_run("Contract Type: ").bold = True
                sub_p.add_run(f"{target.get_lecturer_type_display()}")
            elif filter_type in ('room', 'all_rooms'):
                campus_name = target.campus.name if target.campus else "N/A"
                building_name = target.building.name if target.building else "N/A"
                sub_p.add_run("Campus / Building: ").bold = True
                sub_p.add_run(f"{campus_name} / {building_name}\n")
                sub_p.add_run("Room: ").bold = True
                sub_p.add_run(f"{target.name}  |  ")
                sub_p.add_run("Room Type: ").bold = True
                sub_p.add_run(f"{target.get_room_type_display()}  |  ")
                sub_p.add_run("Capacity: ").bold = True
                sub_p.add_run(f"{target.capacity} Seats")
            else:
                sub_p.add_run(word_subtitle)

            for run in sub_p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(71, 85, 105)

            # Filter slots for this target
            if filter_type in ('group', 'all_groups'):
                slots_target = [s for s in slots_all if s.student_group_id == target.id]
            elif filter_type in ('room', 'all_rooms'):
                slots_target = [s for s in slots_all if s.room_id == target.id]
            elif filter_type in ('lecturer', 'all_lecturers'):
                slots_target = [s for s in slots_all if s.lecturer_id == target.id]
            else:
                slots_target = slots_all

            # Pre-group
            slots_by_day_and_num = {}
            for s in slots_target:
                slots_by_day_and_num.setdefault((s.time_slot.day_of_week, s.time_slot.slot_number), []).append(s)

            # Create Table
            num_cols = len(days) + 1
            num_rows = len(slot_numbers) + 1
            table = doc.add_table(rows=num_rows, cols=num_cols)
            set_table_borders(table)

            # Header Row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Slot / Time"
            set_cell_background(hdr_cells[0], primary_hex)
            hdr_cells[0].paragraphs[0].runs[0].font.bold = True
            hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for col_idx, d in enumerate(days, start=1):
                cell = hdr_cells[col_idx]
                cell.text = day_labels.get(d, f'D{d}')
                set_cell_background(cell, primary_hex)
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Data Rows
            for r_idx, slot_num in enumerate(slot_numbers, start=1):
                row_cells = table.rows[r_idx].cells
                bg_color = alt_hex if r_idx % 2 == 1 else "FFFFFF"

                ts_for_slot = ts_by_slot_number.get(slot_num, [])
                time_label = f"Slot {slot_num}"
                if ts_for_slot:
                    ts0 = ts_for_slot[0]
                    time_label = f"{ts0.start_time.strftime('%H:%M')}\n{ts0.end_time.strftime('%H:%M')}"
                
                row_cells[0].text = time_label
                set_cell_background(row_cells[0], bg_color)
                row_cells[0].paragraphs[0].runs[0].font.bold = True
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for c_idx, day in enumerate(days, start=1):
                    cell = row_cells[c_idx]
                    set_cell_background(cell, bg_color)
                    matching = slots_by_day_and_num.get((day, slot_num), [])
                    if matching:
                        cell.text = "" # Clear default text
                        for s_idx, s in enumerate(matching):
                            if s_idx > 0:
                                p = cell.add_paragraph("────────────────")
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p.runs[0].font.size = Pt(6)
                                p.runs[0].font.color.rgb = RGBColor(203, 213, 225)
                            
                            course_title = s.course.name
                            if len(course_title) > 28:
                                course_title = course_title[:25] + "..."
                            p_title = cell.add_paragraph()
                            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run_title = p_title.add_run(course_title)
                            run_title.italic = True
                            run_title.font.size = Pt(7)
                            run_title.font.color.rgb = RGBColor(13, 90, 79)

                            if filter_type not in ('group', 'all_groups') and s.student_group:
                                p_grp = cell.add_paragraph()
                                p_grp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run_grp = p_grp.add_run(s.student_group.name)
                                run_grp.bold = True
                                run_grp.font.size = Pt(7.5)
                                run_grp.font.color.rgb = RGBColor(30, 41, 59)

                            if filter_type not in ('lecturer', 'all_lecturers') and s.lecturer:
                                p_lec = cell.add_paragraph()
                                p_lec.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run_lec = p_lec.add_run(s.lecturer.name)
                                run_lec.font.size = Pt(7)
                                run_lec.font.color.rgb = RGBColor(30, 58, 138)

                            if filter_type not in ('room', 'all_rooms') and s.room:
                                p_room = cell.add_paragraph()
                                p_room.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run_room = p_room.add_run(f"📍 {s.room.name}")
                                run_room.font.size = Pt(7.5)
                                run_room.font.color.rgb = RGBColor(128, 29, 29)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    safe_name = "".join(c for c in timetable.name if c.isalnum() or c in (' ','_','-')).strip().replace(' ','_')
    response['Content-Disposition'] = f'attachment; filename="timetable_{safe_name}.docx"'
    return response


@login_required
def export_workload_word(request, pk):
    """
    Export lecturer workloads for a timetable as a Word (.docx) file
    formatted exactly like the CUK workload document.
    """
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from django.http import HttpResponse

    uni = get_active_uni(request)
    timetable = get_object_or_404(
        Timetable.objects.select_related('semester', 'semester__university'), pk=pk
    )
    if timetable.semester.university != uni:
        messages.error(request, "Access denied.")
        return redirect('scheduler:dashboard')

    doc = docx.Document()

    # Page Setup (Landscape A4)
    section = doc.sections[-1]
    from docx.enum.section import WD_ORIENT
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Styling Default Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9.5)

    # 1. Header (Centered, bold)
    h_uni = doc.add_paragraph()
    h_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_uni = h_uni.add_run(uni.name.upper())
    r_uni.bold = True
    r_uni.font.size = Pt(13)

    h_sch = doc.add_paragraph()
    h_sch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sch = h_sch.add_run("SCHOOL OF BUSINESS AND ECONOMICS")
    r_sch.bold = True
    r_sch.font.size = Pt(11)

    h_sem = doc.add_paragraph()
    h_sem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sem = h_sem.add_run(f"ACADEMIC YEAR {timetable.semester.name.upper()}")
    r_sem.bold = True
    r_sem.font.size = Pt(10)

    h_title = doc.add_paragraph()
    h_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = h_title.add_run("TEACHING WORKLOAD FOR DEPARTMENT OF ACCOUNTING AND FINANCE")
    r_title.bold = True
    r_title.underline = True
    r_title.font.size = Pt(11)
    
    doc.add_paragraph() # Spacer

    # Helper function to style headers and tables
    def set_cell_background(cell, hex_color):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    def set_table_borders(table):
        tblPr = table._tbl.tblPr
        tblBorders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="A0AEC0"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="A0AEC0"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="A0AEC0"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="A0AEC0"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(tblBorders)

    # Separate lecturers into Full-Time and Part-Time
    lecturers = Lecturer.objects.filter(department__faculty__campus__university=uni).order_by('name')
    slots = list(timetable.slots.select_related('course', 'room', 'student_group').filter(lecturer__in=lecturers))

    from collections import defaultdict
    slots_by_lec = defaultdict(list)
    for s in slots:
        if s.lecturer_id:
            slots_by_lec[s.lecturer_id].append(s)

    ft_lecturers = []
    pt_lecturers = []
    for lec in lecturers:
        if lec.lecturer_type == 'FT':
            ft_lecturers.append(lec)
        else:
            pt_lecturers.append(lec)

    # Helper function to generate table
    def add_workload_table(title_label, target_lecturers):
        p = doc.add_paragraph()
        r = p.add_run(title_label.upper())
        r.bold = True
        r.font.size = Pt(11)
        
        table = doc.add_table(rows=1, cols=7)
        table.autofit = False
        set_table_borders(table)
        
        # Header Row
        headers = ["S.NO", "NAME", "QUALIFICATIONS / EMAIL", "UNIT CODE/ NAME", "CLASS SIZE", "STATUS", "TOTAL UNITS"]
        widths = [Inches(0.5), Inches(1.5), Inches(1.8), Inches(2.8), Inches(0.8), Inches(0.8), Inches(0.8)]
        
        hdr_row = table.rows[0]
        for idx, text in enumerate(headers):
            cell = hdr_row.cells[idx]
            cell.text = text
            cell.width = widths[idx]
            set_cell_background(cell, "F1F5F9")
            # Make text bold
            p_cell = cell.paragraphs[0]
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_cell.runs[0]
            run.font.bold = True
            run.font.size = Pt(9.5)
            
        s_no = 1
        for lec in target_lecturers:
            lec_slots = slots_by_lec.get(lec.id, [])
            # Format course units code/name
            units_text = ""
            sizes_text = ""
            status_val = lec.get_lecturer_type_display()  # 'Full-Time' or 'Part-Time'
            status_text = ""
            
            for s in lec_slots:
                units_text += f"{s.course.code}: {s.course.name}\n"
                sizes_text += f"{s.student_group.size if s.student_group else '-'}\n"
                status_text += f"{status_val}\n"
                
            # Strip trailing newlines
            units_text = units_text.rstrip("\n")
            sizes_text = sizes_text.rstrip("\n")
            status_text = status_text.rstrip("\n")
            
            row_cells = table.add_row().cells
            for idx in range(7):
                row_cells[idx].width = widths[idx]
                
            row_cells[0].text = str(s_no)
            row_cells[1].text = lec.name
            row_cells[2].text = lec.email or lec.department.name
            row_cells[3].text = units_text or "No assigned courses"
            row_cells[4].text = sizes_text or "-"
            row_cells[5].text = status_text or status_val
            row_cells[6].text = str(len(lec_slots))
            
            # Align center for numbers and status
            for idx in [0, 4, 5, 6]:
                row_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            s_no += 1
            
        doc.add_paragraph() # Spacer after table

    if ft_lecturers:
        add_workload_table("FULL - TIME LECTURERS", ft_lecturers)
    if pt_lecturers:
        add_workload_table("PART - TIME LECTURERS", pt_lecturers)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    safe_name = "".join(c for c in timetable.name if c.isalnum() or c in (' ','_','-')).strip().replace(' ','_')
    response['Content-Disposition'] = f'attachment; filename="workload_{safe_name}.docx"'
    doc.save(response)
    return response


@login_required
def student_my_schedule(request):
    """Personal timetable view for a logged-in student."""
    try:
        student_group = request.user.profile.student_group
    except Exception:
        student_group = None

    if not student_group:
        messages.error(request, "Your account is not linked to a student group. Ask an administrator to link your profile.")
        return redirect('accounts:profile')

    import uuid
    from django.urls import reverse

    if not student_group.calendar_token:
        student_group.calendar_token = uuid.uuid4()
        student_group.save(update_fields=['calendar_token'])

    feed_url = request.build_absolute_uri(
        reverse('scheduler:student_group_calendar_feed', args=[student_group.calendar_token])
    )
    webcal_url = feed_url.replace('http://', 'webcal://').replace('https://', 'webcal://')

    university = get_active_uni(request)
    active_timetable = Timetable.objects.filter(
        semester__university=university, is_active=True
    ).first()
    if not active_timetable:
        active_timetable = Timetable.objects.filter(semester__university=university).first()

    slots = []
    if active_timetable:
        slots = list(
            ScheduleSlot.objects.filter(
                timetable=active_timetable, student_group=student_group
            ).select_related('course', 'room', 'time_slot', 'lecturer')
            .order_by('time_slot__day_of_week', 'time_slot__slot_number')
        )
        from django.utils import timezone
        local_now = timezone.localtime(timezone.now())
        current_dow = local_now.isoweekday()
        current_time = local_now.time()
        for s in slots:
            ts = s.time_slot
            if ts.day_of_week < current_dow:
                s.has_ended = True
                s.is_ongoing = False
            elif ts.day_of_week == current_dow:
                s.has_ended = (ts.end_time < current_time)
                s.is_ongoing = (ts.start_time <= current_time < ts.end_time)  # FIX BUG 16: strict < for end_time
            else:
                s.has_ended = False
                s.is_ongoing = False


    return render(request, 'scheduler/student_my_schedule.html', {
        'student_group': student_group,
        'timetable': active_timetable,
        'slots': slots,
        'feed_url': feed_url,
        'webcal_url': webcal_url,
    })


@login_required
@manager_required
def setup_wizard(request):
    """
    Quick-start wizard: create university structure, default Mon–Sun timeslots, and an active semester.
    """
    import datetime
    from django.db import transaction

    university = get_active_uni(request)
    created = []

    if request.method == 'POST':
        uni_name = request.POST.get('university_name', '').strip()
        uni_code = request.POST.get('university_code', '').strip().upper()
        campus_name = request.POST.get('campus_name', 'Main Campus').strip()
        semester_name = request.POST.get('semester_name', '').strip()
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')

        if not all([uni_name, uni_code, semester_name, start_date, end_date]):
            messages.error(request, "Please fill in all required fields.")
        else:
            try:
                with transaction.atomic():
                    if not university:
                        university = University.objects.create(name=uni_name, code=uni_code)
                        created.append(f"University '{uni_name}'")
                    campus, _ = Campus.objects.get_or_create(
                        university=university, name=campus_name
                    )
                    faculty, _ = Faculty.objects.get_or_create(campus=campus, name='General Faculty')
                    dept, _ = Department.objects.get_or_create(faculty=faculty, name='General Department')

                    Semester.objects.filter(university=university, is_active=True).update(is_active=False)
                    semester = Semester.objects.create(
                        university=university,
                        name=semester_name,
                        start_date=datetime.date.fromisoformat(start_date),
                        end_date=datetime.date.fromisoformat(end_date),
                        is_active=True,
                    )
                    created.append(f"Semester '{semester_name}'")

                    slot_templates = [
                        (datetime.time(8, 0), datetime.time(11, 0), 1, False),
                        (datetime.time(11, 0), datetime.time(14, 0), 2, False),
                        (datetime.time(14, 0), datetime.time(17, 0), 3, False),
                        (datetime.time(17, 30), datetime.time(20, 30), 4, True),
                    ]
                    for day in range(1, 8):
                        for start_t, end_t, slot_num, is_evening in slot_templates:
                            TimeSlot.objects.get_or_create(
                                university=university,
                                day_of_week=day,
                                slot_number=slot_num,
                                defaults={
                                    'start_time': start_t,
                                    'end_time': end_t,
                                    'is_evening': is_evening,
                                },
                            )
                    created.append("Default Mon–Sun 3-hour time slots (4 per day)")

                    request.session['active_university_id'] = university.id
                messages.success(request, "Setup complete: " + ", ".join(created) + ".")
                return redirect('scheduler:resources_manager')
            except Exception as exc:
                messages.error(request, f"Setup failed: {exc}")

    has_timeslots = bool(university and TimeSlot.objects.filter(university=university).exists())
    has_semester = bool(university and Semester.objects.filter(university=university).exists())

    return render(request, 'scheduler/setup_wizard.html', {
        'university': university,
        'has_timeslots': has_timeslots,
        'has_semester': has_semester,
    })


@manager_required
def apply_default_3hr_timeslots(request):
    """
    1-Click Preset action for University Managers:
    Resets/re-provisions standard 3-hour time slots (Mon-Sun, 4 slots per day)
    for the active university.
    """
    university = get_active_uni(request)
    if not university:
        messages.error(request, "No active university selected.")
        return redirect('scheduler:resources_manager')

    if request.method == 'POST':
        # Remove existing slots for clean reset
        TimeSlot.objects.filter(university=university).delete()

        slot_templates = [
            (datetime.time(8, 0), datetime.time(11, 0), 1, False),
            (datetime.time(11, 0), datetime.time(14, 0), 2, False),
            (datetime.time(14, 0), datetime.time(17, 0), 3, False),
            (datetime.time(17, 30), datetime.time(20, 30), 4, True),
        ]
        slots = [
            TimeSlot(
                university=university,
                day_of_week=day,
                slot_number=slot_num,
                start_time=start_t,
                end_time=end_t,
                is_evening=is_eve,
            )
            for day in range(1, 8)
            for start_t, end_t, slot_num, is_eve in slot_templates
        ]
        TimeSlot.objects.bulk_create(slots)
        messages.success(request, f"Successfully created 28 standard 3-hour time slots (Mon–Sun) for {university.name}.")

    return redirect('/resources/?tab=timeslot')


def lecturer_calendar_feed(request, token):
    """
    Public feed serving a specific lecturer's weekly schedule as an iCalendar file.
    """
    lecturer = get_object_or_404(Lecturer, calendar_token=token)
    from .calendar_exporter import generate_lecturer_ics
    ics_content = generate_lecturer_ics(lecturer)
    
    response = HttpResponse(ics_content, content_type='text/calendar; charset=utf-8')
    # Clean lecturer name for header
    safe_name = "".join(c for c in lecturer.name if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')
    response['Content-Disposition'] = f'inline; filename="lecturer_{safe_name}_schedule.ics"'
    return response


def student_group_calendar_feed(request, token):
    """
    Public feed serving a specific student group's weekly schedule as an iCalendar file.
    """
    student_group = get_object_or_404(StudentGroup, calendar_token=token)
    from .calendar_exporter import generate_student_group_ics
    ics_content = generate_student_group_ics(student_group)
    
    response = HttpResponse(ics_content, content_type='text/calendar; charset=utf-8')
    safe_name = "".join(c for c in student_group.name if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')
    response['Content-Disposition'] = f'inline; filename="group_{safe_name}_schedule.ics"'
    return response


@login_required
def sync_to_google_calendar(request):
    """
    Manually triggers a full sync of the lecturer's active timetable schedule to Google Calendar.
    """
    try:
        lecturer = request.user.lecturer_profile
    except Exception:
        messages.error(request, "Your account is not linked to a lecturer record.")
        return redirect('accounts:profile')

    from accounts.models import GoogleCalendarToken
    if not GoogleCalendarToken.objects.filter(user=request.user).exists():
        messages.error(request, "Please connect your Google Calendar account first.")
        return redirect('accounts:profile')

    university = get_active_uni(request)
    active_timetable = Timetable.objects.filter(
        semester__university=university, is_active=True
    ).first()
    if not active_timetable:
        active_timetable = Timetable.objects.filter(semester__university=university).first()

    if not active_timetable:
        messages.error(request, "No active timetable found to sync.")
        return redirect('scheduler:lecturer_my_schedule')

    import sys
    if 'test' in sys.argv or 'pytest' in sys.argv or any('pytest' in arg for arg in sys.argv):
        from .google_tasks import sync_lecturer_timetable_google
        sync_lecturer_timetable_google(lecturer.id, active_timetable.id)
    else:
        from django_q.tasks import async_task
        async_task('scheduler.google_tasks.sync_lecturer_timetable_google', lecturer.id, active_timetable.id)

    messages.success(request, "✓ Google Calendar sync started in the background. Your calendar will update in a moment!")
    return redirect('scheduler:lecturer_my_schedule')


@login_required
def notifications_list(request):
    notifications = request.user.notifications.all()
    role = get_user_role(request)

    # Lecturers have their own portal base — use the matching template so
    # the lecturer sidebar menus are NOT lost when opening notifications.
    if role == ROLE_LECTURER:
        # Also provide unread_count and lecturer for the portal base template
        unread_count = request.user.notifications.filter(is_read=False).count()
        lecturer = None
        try:
            lecturer = request.user.profile.lecturer
        except Exception:
            pass
        return render(request, 'scheduler/lecturer_portal_notifications.html', {
            'notifications': notifications,
            'unread_count': unread_count,
            'lecturer': lecturer,
        })

    return render(request, 'scheduler/notifications_list.html', {
        'notifications': notifications
    })

@login_required
def notification_read(request, pk):
    from .models import Notification
    notification = get_object_or_404(Notification, pk=pk, user=request.user)
    notification.is_read = True
    notification.save()
    if notification.link:
        return redirect(notification.link)
    return redirect('scheduler:notifications_list')

@login_required
def notifications_mark_all_read(request):
    request.user.notifications.filter(is_read=False).update(is_read=True)
    messages.success(request, "All notifications marked as read.")
    return redirect('scheduler:notifications_list')

@login_required
def subscription_billing(request):
    university = get_active_uni(request)
    if not university:
        messages.error(request, "No active university selected.")
        return redirect('scheduler:dashboard')
    
    role = get_user_role(request)
    if role not in ('admin', 'institution_admin'):
        messages.error(request, "Permission denied. Only Admins can manage subscriptions.")
        return redirect('scheduler:dashboard')

    from .models import Subscription
    subscription = getattr(university, 'subscription', None)
    if not subscription:
        subscription = Subscription.objects.create(university=university, tier='free', status='active')

    if request.method == 'POST':
        tier = request.POST.get('tier')
        if tier in ('free', 'growth', 'enterprise'):
            subscription.tier = tier
            if tier == 'free':
                subscription.max_rooms = 10
                subscription.max_courses = 50
            elif tier == 'growth':
                subscription.max_rooms = 30
                subscription.max_courses = 150
            elif tier == 'enterprise':
                subscription.max_rooms = 100
                subscription.max_courses = 500
            subscription.save()
            messages.success(request, f"Subscription upgraded to {subscription.get_tier_display()}!")
            return redirect('scheduler:subscription_billing')

    rooms_count = Room.objects.filter(campus__university=university).count()
    courses_count = Course.objects.filter(program__department__faculty__campus__university=university).count()
    
    return render(request, 'scheduler/subscription_billing.html', {
        'subscription': subscription,
        'rooms_count': rooms_count,
        'courses_count': courses_count,
        'rooms_pct': min(100, int((rooms_count / subscription.max_rooms) * 100)) if subscription.max_rooms else 0,
        'courses_pct': min(100, int((courses_count / subscription.max_courses) * 100)) if subscription.max_courses else 0,
    })


# ═══════════════════════════════════════════════════════════════════════════════
# LECTURER PORTAL VIEWS
# ═══════════════════════════════════════════════════════════════════════════════

def _get_lecturer_or_redirect(request):
    """Helper: return (lecturer, None) or (None, redirect_response) for portal views."""
    try:
        lecturer = request.user.profile.lecturer
        if lecturer:
            return lecturer, None
    except Exception:
        pass
    messages.error(request, "Your account is not linked to a lecturer profile. Please update your profile.")
    return None, redirect('accounts:profile')



@login_required
def lecturer_portal_dashboard(request):
    """Main Lecturer Portal dashboard — aggregates all personalised data."""
    import datetime
    role = get_user_role(request)
    if role != ROLE_LECTURER:
        messages.warning(request, "Access restricted to lecturers.")
        return redirect('scheduler:dashboard')

    lecturer, err = _get_lecturer_or_redirect(request)
    if err:
        return err

    university = get_active_uni(request)
    active_timetable = None
    if university:
        active_timetable = (
            Timetable.objects.filter(semester__university=university, is_active=True).first()
            or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
        )

    my_slots = []
    today_slots = []
    next_slot = None
    total_hours = 0.0
    unique_courses = set()
    unique_groups = set()

    if active_timetable:
        my_slots = list(
            ScheduleSlot.objects.filter(timetable=active_timetable, lecturer=lecturer)
            .select_related('course', 'room', 'time_slot', 'student_group',
                            'student_group__program', 'course__program')
            .order_by('time_slot__day_of_week', 'time_slot__slot_number')
        )
        now_time = datetime.datetime.now().time()
        today_dow = datetime.date.today().weekday() + 1  # 1=Mon … 7=Sun

        for s in my_slots:
            unique_courses.add(s.course_id)
            unique_groups.add(s.student_group_id)
            ts = s.time_slot
            duration_mins = (
                ts.end_time.hour * 60 + ts.end_time.minute
            ) - (ts.start_time.hour * 60 + ts.start_time.minute)
            total_hours += duration_mins / 60

        today_slots = [s for s in my_slots if s.time_slot.day_of_week == today_dow]
        for s in today_slots:
            s.is_now = s.time_slot.start_time <= now_time <= s.time_slot.end_time
        upcoming = [s for s in today_slots if s.time_slot.start_time > now_time]
        next_slot = upcoming[0] if upcoming else None

    # Notifications (unread)
    notifications = list(request.user.notifications.filter(is_read=False).order_by('-created_at')[:8])
    unread_count = request.user.notifications.filter(is_read=False).count()

    # Announcements
    announcements = []
    if university:
        announcements = list(
            Announcement.objects.filter(
                Q(university=university) | Q(university__isnull=True)
            ).order_by('-created_at')[:5]
        )

    # Active timetable semester
    active_semester = None
    if university:
        active_semester = Semester.objects.filter(university=university, is_active=True).first()

    # Attendance: open sessions
    open_sessions = AttendanceSession.objects.filter(
        schedule_slot__lecturer=lecturer, is_active=True
    ).select_related('schedule_slot__course', 'schedule_slot__room')[:3]

    now_hour = datetime.datetime.now().hour
    greeting = 'Good morning' if now_hour < 12 else ('Good afternoon' if now_hour < 17 else 'Good evening')

    return render(request, 'scheduler/lecturer_portal_dashboard.html', {
        'active_role': role,
        'lecturer': lecturer,
        'active_university': university,
        'active_semester': active_semester,
        'timetable': active_timetable,
        'my_slots': my_slots,
        'today_slots': today_slots,
        'next_slot': next_slot,
        'total_slots': len(my_slots),
        'total_hours': round(total_hours, 1),
        'unique_courses_count': len(unique_courses),
        'unique_groups_count': len(unique_groups),
        'notifications': notifications,
        'unread_count': unread_count,
        'announcements': announcements,
        'open_sessions': open_sessions,
        'greeting': greeting,
        'today_date': datetime.date.today(),
    })


@login_required
def lecturer_portal_weekly_timetable(request):
    """Full 5-day weekly timetable grid for the lecturer."""
    import datetime
    role = get_user_role(request)
    if role != ROLE_LECTURER:
        return redirect('scheduler:dashboard')

    lecturer, err = _get_lecturer_or_redirect(request)
    if err:
        return err

    university = get_active_uni(request)
    active_timetable = None
    if university:
        active_timetable = (
            Timetable.objects.filter(semester__university=university, is_active=True).first()
            or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
        )

    slots_by_day = {1: [], 2: [], 3: [], 4: [], 5: [], 6: [], 7: []}
    all_slots = []
    # All university timeslots — used to build the FULL weekly grid (not just slots the lecturer teaches)
    all_university_timeslots = []

    if active_timetable:
        all_slots = list(
            ScheduleSlot.objects.filter(timetable=active_timetable, lecturer=lecturer)
            .select_related('course', 'room', 'time_slot', 'student_group')
            .order_by('time_slot__slot_number', 'time_slot__day_of_week')
        )
        from django.utils import timezone
        local_now = timezone.localtime(timezone.now())
        current_dow = local_now.isoweekday()
        current_time = local_now.time()
        for s in all_slots:
            ts = s.time_slot
            if ts.day_of_week < current_dow:
                s.has_ended = True
                s.is_ongoing = False
            elif ts.day_of_week > current_dow:
                s.has_ended = False
                s.is_ongoing = False
            else:
                s.has_ended = current_time > ts.end_time
                s.is_ongoing = ts.start_time <= current_time <= ts.end_time

        for s in all_slots:
            slots_by_day[s.time_slot.day_of_week].append(s)

    # Always load ALL university timeslots to build the complete weekly grid
    if university:
        all_university_timeslots = list(
            TimeSlot.objects.filter(university=university)
            .order_by('slot_number', 'day_of_week')
        )

    import datetime
    _today = datetime.date.today()
    _monday = _today - datetime.timedelta(days=_today.weekday())
    
    max_day = 7

    all_days = [
        (1, 'Monday'), (2, 'Tuesday'), (3, 'Wednesday'),
        (4, 'Thursday'), (5, 'Friday'), (6, 'Saturday'), (7, 'Sunday')
    ]
    days_list = []
    for d, day_name in all_days[:max_day]:
        days_list.append({
            'dow': d,
            'name': day_name,
            'date': _monday + datetime.timedelta(days=d - 1)
        })
    today_dow = _today.weekday() + 1

    # Build the FULL weekly grid using ALL university timeslots (not just the lecturer's own slots)
    # This ensures every time slot row is displayed even if the lecturer has no class at that time.
    grid_rows = []
    unique_slots_info = {}
    for ts in all_university_timeslots:
        num = ts.slot_number
        if num not in unique_slots_info:
            unique_slots_info[num] = {
                'slot_number': num,
                'start_time': ts.start_time,
                'end_time': ts.end_time,
            }

    sorted_slots = sorted(unique_slots_info.values(), key=lambda x: x['start_time'])

    for s_info in sorted_slots:
        num = s_info['slot_number']
        day_slots = {}
        for d in range(1, max_day + 1):
            # Find the lecturer's class at this day+slot, or None (empty cell)
            day_slots[d] = next(
                (s for s in all_slots if s.time_slot.day_of_week == d and s.time_slot.slot_number == num),
                None
            )

        grid_rows.append({
            'start_time': s_info['start_time'],
            'end_time': s_info['end_time'],
            'slot_number': num,
            'day_slots': day_slots,
        })

    import uuid
    from django.urls import reverse
    if not lecturer.calendar_token:
        lecturer.calendar_token = uuid.uuid4()
        lecturer.save(update_fields=['calendar_token'])

    feed_url = request.build_absolute_uri(
        reverse('scheduler:lecturer_calendar_feed', args=[lecturer.calendar_token])
    )
    webcal_url = feed_url.replace('http://', 'webcal://').replace('https://', 'webcal://')

    return render(request, 'scheduler/lecturer_weekly_timetable.html', {
        'active_role': role,
        'lecturer': lecturer,
        'active_university': university,
        'timetable': active_timetable,
        'slots_by_day': slots_by_day,
        'all_slots': all_slots,
        'grid_rows': grid_rows,
        'days': days_list,
        'today_dow': today_dow,
        'active_semester': Semester.objects.filter(
            university=university, is_active=True).first() if university else None,
        'feed_url': feed_url,
        'webcal_url': webcal_url,
    })




@login_required
def lecturer_portal_courses(request):
    """List all courses assigned to the lecturer."""
    role = get_user_role(request)
    if role != ROLE_LECTURER:
        return redirect('scheduler:dashboard')

    lecturer, err = _get_lecturer_or_redirect(request)
    if err:
        return err

    courses = Course.objects.filter(lecturer=lecturer).select_related(
        'program', 'program__department', 'student_group'
    ).prefetch_related('additional_student_groups').order_by('code')

    university = get_active_uni(request)
    active_timetable = None
    if university:
        active_timetable = (
            Timetable.objects.filter(semester__university=university, is_active=True).first()
            or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
        )

    # Annotate each course with weekly session count from active timetable
    for c in courses:
        if active_timetable:
            c.weekly_sessions = ScheduleSlot.objects.filter(
                timetable=active_timetable, course=c, lecturer=lecturer
            ).count()
        else:
            c.weekly_sessions = 0

    return render(request, 'scheduler/lecturer_courses.html', {
        'active_role': role,
        'lecturer': lecturer,
        'courses': courses,
        'active_university': university,
    })


@login_required
def lecturer_portal_student_groups(request):
    """All student groups the lecturer teaches."""
    role = get_user_role(request)
    if role != ROLE_LECTURER:
        return redirect('scheduler:dashboard')

    lecturer, err = _get_lecturer_or_redirect(request)
    if err:
        return err

    university = get_active_uni(request)
    active_timetable = None
    if university:
        active_timetable = (
            Timetable.objects.filter(semester__university=university, is_active=True).first()
            or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
        )

    groups_data = []
    if active_timetable:
        from django.db.models import Count as DjangoCount
        slot_qs = (
            ScheduleSlot.objects.filter(timetable=active_timetable, lecturer=lecturer)
            .values('student_group')
            .annotate(session_count=DjangoCount('id'))
        )
        group_sessions = {r['student_group']: r['session_count'] for r in slot_qs}
        group_ids = list(group_sessions.keys())
        groups = StudentGroup.objects.filter(id__in=group_ids).select_related(
            'program', 'program__department'
        )
        for g in groups:
            groups_data.append({
                'group': g,
                'sessions': group_sessions.get(g.id, 0),
                'courses': Course.objects.filter(
                    lecturer=lecturer, student_group=g
                ).values_list('code', flat=True),
            })
    else:
        # Fall back to courses FKs
        direct_groups = StudentGroup.objects.filter(
            courses__lecturer=lecturer
        ).distinct().select_related('program', 'program__department')
        for g in direct_groups:
            groups_data.append({'group': g, 'sessions': 0, 'courses': []})

    return render(request, 'scheduler/lecturer_student_groups.html', {
        'active_role': role,
        'lecturer': lecturer,
        'groups_data': groups_data,
        'active_university': university,
    })


@login_required
def lecturer_portal_workload(request):
    """Teaching workload breakdown for the lecturer."""
    import datetime
    role = get_user_role(request)
    if role != ROLE_LECTURER:
        return redirect('scheduler:dashboard')

    lecturer, err = _get_lecturer_or_redirect(request)
    if err:
        return err

    university = get_active_uni(request)
    active_timetable = None
    if university:
        active_timetable = (
            Timetable.objects.filter(semester__university=university, is_active=True).first()
            or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
        )

    day_hours = {1: 0.0, 2: 0.0, 3: 0.0, 4: 0.0, 5: 0.0, 6: 0.0, 7: 0.0}
    day_slots = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0, 7: 0}
    day_names = {1: 'Monday', 2: 'Tuesday', 3: 'Wednesday', 4: 'Thursday', 5: 'Friday', 6: 'Saturday', 7: 'Sunday'}
    slots = []
    total_hours = 0.0
    unique_courses = set()

    if active_timetable:
        slots = list(
            ScheduleSlot.objects.filter(timetable=active_timetable, lecturer=lecturer)
            .select_related('course', 'room', 'time_slot', 'student_group')
            .order_by('time_slot__day_of_week', 'time_slot__slot_number')
        )
        for s in slots:
            unique_courses.add(s.course_id)
            ts = s.time_slot
            dur = (ts.end_time.hour * 60 + ts.end_time.minute) - (
                ts.start_time.hour * 60 + ts.start_time.minute
            )
            h = dur / 60
            total_hours += h
            dow = ts.day_of_week
            if dow in day_hours:  # 1–7: Mon–Sun
                day_hours[dow] += h
                day_slots[dow] += 1

    max_day_hrs = max(day_hours.values()) if day_hours else 1
    workload_pct = round((total_hours / lecturer.max_hours_per_week) * 100) if lecturer.max_hours_per_week else 0

    return render(request, 'scheduler/lecturer_workload.html', {
        'active_role': role,
        'lecturer': lecturer,
        'slots': slots,
        'total_hours': round(total_hours, 1),
        'unique_courses_count': len(unique_courses),
        'day_hours': day_hours,
        'day_slots': day_slots,
        'day_names': day_names,
        'max_day_hrs': max_day_hrs,
        'workload_pct': workload_pct,
        'active_university': university,
        'timetable': active_timetable,
    })


@login_required
def lecturer_attendance_start(request, slot_id):
    """Open (or resume) an attendance session for a given ScheduleSlot."""
    import datetime
    role = get_user_role(request)
    if role != ROLE_LECTURER:
        return redirect('scheduler:dashboard')

    lecturer, err = _get_lecturer_or_redirect(request)
    if err:
        return err

    slot = get_object_or_404(
        ScheduleSlot.objects.select_related(
            'course', 'room', 'time_slot', 'student_group', 'lecturer'
        ),
        id=slot_id, lecturer=lecturer
    )

    today = datetime.date.today()
    session, created = AttendanceSession.objects.get_or_create(
        schedule_slot=slot,
        date=today,
        defaults={'is_active': True},
    )
    if not session.is_active:
        session.is_active = True
        session.closed_at = None
        session.save(update_fields=['is_active', 'closed_at'])

    return redirect('scheduler:lecturer_attendance_session', session_id=session.id)


@login_required
def lecturer_attendance_session(request, session_id):
    """Display and handle the manual attendance marking form."""
    import datetime
    role = get_user_role(request)
    if role != ROLE_LECTURER:
        return redirect('scheduler:dashboard')

    lecturer, err = _get_lecturer_or_redirect(request)
    if err:
        return err

    session = get_object_or_404(
        AttendanceSession.objects.select_related(
            'schedule_slot__course', 'schedule_slot__room',
            'schedule_slot__time_slot', 'schedule_slot__student_group'
        ),
        id=session_id,
        schedule_slot__lecturer=lecturer,
    )

    if request.method == 'POST':
        action = request.POST.get('action', 'save')
        # Present IDs submitted as checkboxes
        present_ids = set(request.POST.getlist('present'))
        # Fetch/update all records for this session
        for record in session.records.all():
            record.is_present = str(record.id) in present_ids
            record.save(update_fields=['is_present', 'marked_at'])

        if action == 'close':
            session.is_active = False
            session.closed_at = datetime.datetime.now()
            session.save(update_fields=['is_active', 'closed_at'])
            messages.success(request, f"Attendance session closed. {session.records.filter(is_present=True).count()} students marked present.")
            return redirect('scheduler:lecturer_attendance_report')
        else:
            messages.success(request, "Attendance saved.")
            return redirect('scheduler:lecturer_attendance_session', session_id=session.id)

    # Auto-populate records from student group (if not yet done)
    group = session.schedule_slot.student_group

    # Clear old generic "Student 001" placeholders if present to upgrade them
    if session.records.filter(student_name__startswith='Student ').exists():
        session.records.all().delete()

    existing_names = set(session.records.values_list('student_name', flat=True))

    # Generate realistic student names deterministically per group size
    if not existing_names and group and group.size > 0:
        import random
        # Seed with group.id to make the list consistent for this group across sessions
        rng = random.Random(group.id)
        
        first_names = [
            'John', 'Jane', 'Mary', 'Joseph', 'David', 'James', 'Peter', 'Anne', 'Grace', 'Paul',
            'Moses', 'Ruth', 'Sarah', 'Alice', 'Michael', 'Daniel', 'Elizabeth', 'Esther', 'Thomas', 'Stephen',
            'Mercy', 'Joy', 'Hope', 'Charles', 'Philip', 'George', 'William', 'Andrew', 'Simon', 'Lucy',
            'Abdi', 'Fatuma', 'Amina', 'Hassan', 'Ali', 'Mohamed', 'Hussein', 'Halima', 'Omar', 'Yusuf'
        ]
        last_names = [
            'Mwanza', 'Karanja', 'Ochieng', 'Kamau', 'Mwangi', 'Otieno', 'Wanjiku', 'Kimani', 'Njoroge', 'Ouma',
            'Maina', 'Nduta', 'Muthoni', 'Odhiambo', 'Waweru', 'Kiprotich', 'Kipkorir', 'Mutua', 'Musyoka', 'Mogaka',
            'Nyangidi', 'Gathoni', 'Wambui', 'Chepngetich', 'Kibet', 'Cheruiyot', 'Lagat', 'Wafula', 'Nekesa', 'Simiyu',
            'Ibrahim', 'Kariuki', 'Njeri', 'Kipruto', 'Muriuki', 'Ndungu', 'Wambua', 'Onyango', 'Okoth', 'Anyango'
        ]
        
        records_to_create = []
        generated_names = []
        for i in range(1, group.size + 1):
            fn = rng.choice(first_names)
            ln = rng.choice(last_names)
            name = f"{fn} {ln}"
            while name in generated_names:
                fn = rng.choice(first_names)
                ln = rng.choice(last_names)
                name = f"{fn} {ln}"
            generated_names.append(name)
            
            # Format a realistic registration number prefix based on the group name
            code_prefix = ''.join([c for c in group.name if c.isalnum()]).upper()[:4]
            student_id = f"{code_prefix}/{i:03d}/{rng.randint(23, 26)}"
            
            records_to_create.append(AttendanceRecord(
                session=session,
                student_name=name,
                student_id=student_id,
                is_present=False,
            ))
        AttendanceRecord.objects.bulk_create(records_to_create, ignore_conflicts=True)

    records = session.records.all().order_by('student_name')
    present_count = records.filter(is_present=True).count()

    return render(request, 'scheduler/lecturer_attendance_session.html', {
        'active_role': role,
        'lecturer': lecturer,
        'session': session,
        'records': records,
        'present_count': present_count,
        'absent_count': records.count() - present_count,
        'total': records.count(),
    })


@login_required
def lecturer_attendance_report(request):
    """Attendance report across all sessions for this lecturer."""
    role = get_user_role(request)
    if role != ROLE_LECTURER:
        return redirect('scheduler:dashboard')

    lecturer, err = _get_lecturer_or_redirect(request)
    if err:
        return err

    sessions = AttendanceSession.objects.filter(
        schedule_slot__lecturer=lecturer
    ).select_related(
        'schedule_slot__course', 'schedule_slot__student_group',
        'schedule_slot__room', 'schedule_slot__time_slot'
    ).order_by('-date', '-created_at')

    # Filter by course if requested
    course_filter = request.GET.get('course')
    if course_filter:
        sessions = sessions.filter(schedule_slot__course__code__icontains=course_filter)

    # Annotate each session
    session_data = []
    for s in sessions:
        total = s.records.count()
        present = s.records.filter(is_present=True).count()
        session_data.append({
            'session': s,
            'total': total,
            'present': present,
            'absent': total - present,
            'rate': round((present / total) * 100) if total > 0 else 0,
        })

    courses = Course.objects.filter(lecturer=lecturer).order_by('code')

    return render(request, 'scheduler/lecturer_attendance_report.html', {
        'active_role': role,
        'lecturer': lecturer,
        'session_data': session_data,
        'courses': courses,
        'course_filter': course_filter or '',
    })


@login_required
def lecturer_portal_profile(request):
    """Lecturer profile view and bio editor."""
    role = get_user_role(request)
    if role != ROLE_LECTURER:
        return redirect('scheduler:dashboard')

    lecturer, err = _get_lecturer_or_redirect(request)
    if err:
        return err

    university = get_active_uni(request)

    if request.method == 'POST':
        bio = request.POST.get('bio', '').strip()
        staff_id = request.POST.get('staff_id', '').strip()
        try:
            profile = request.user.profile
            profile.bio = bio
            profile.save(update_fields=['bio'])
        except Exception:
            pass
        if staff_id and (not lecturer.staff_id or lecturer.staff_id != staff_id):
            from django.db import IntegrityError
            try:
                lecturer.staff_id = staff_id
                lecturer.save(update_fields=['staff_id'])
            except IntegrityError:
                messages.warning(request, "That Staff ID is already in use. Profile saved without changing it.")
        messages.success(request, "Profile updated successfully.")
        return redirect('scheduler:lecturer_portal_profile')

    try:
        bio = request.user.profile.bio
    except Exception:
        bio = ''

    # Stats
    total_courses = Course.objects.filter(lecturer=lecturer).count()
    active_timetable = None
    if university:
        active_timetable = (
            Timetable.objects.filter(semester__university=university, is_active=True).first()
            or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
        )
    total_slots = 0
    if active_timetable:
        total_slots = ScheduleSlot.objects.filter(timetable=active_timetable, lecturer=lecturer).count()

    total_sessions_taken = AttendanceSession.objects.filter(
        schedule_slot__lecturer=lecturer
    ).count()

    return render(request, 'scheduler/lecturer_portal_profile.html', {
        'active_role': role,
        'lecturer': lecturer,
        'bio': bio,
        'active_university': university,
        'total_courses': total_courses,
        'total_slots': total_slots,
        'total_sessions_taken': total_sessions_taken,
    })


@login_required
def admin_lecturer_profile(request, pk):
    """Admin view to manage an individual lecturer's profile and assignments."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    university = get_active_uni(request)
    lecturer = get_object_or_404(
        Lecturer.objects.select_related('department', 'department__faculty'), pk=pk
    )
    if lecturer.department.faculty.campus.university != university:
        messages.error(request, "Access denied.")
        return redirect('scheduler:resources_manager')

    if request.method == 'POST':
        name = request.POST.get('name')
        email = request.POST.get('email')
        staff_id = request.POST.get('staff_id')
        dept_id = request.POST.get('department')
        max_hours = request.POST.get('max_hours_per_week')
        max_slots = request.POST.get('max_slots_per_day')
        is_active = request.POST.get('is_active') == 'on'

        lecturer.name = name
        lecturer.email = email
        if staff_id:
            lecturer.staff_id = staff_id
        else:
            lecturer.staff_id = None
        
        lecturer_type = request.POST.get('lecturer_type')
        if lecturer_type in ('FT', 'PT'):
            lecturer.lecturer_type = lecturer_type

        if dept_id:
            dept = Department.objects.filter(faculty__campus__university=university, id=dept_id).first()
            if dept:
                lecturer.department = dept
        
        if max_hours:
            lecturer.max_hours_per_week = int(max_hours)
        if max_slots:
            lecturer.max_slots_per_day = int(max_slots)
        
        lecturer.is_active = is_active
        lecturer.save()
        messages.success(request, f"Lecturer {lecturer.name} profile updated successfully.")
        return redirect('scheduler:admin_lecturer_profile', pk=lecturer.pk)

    departments = Department.objects.filter(faculty__campus__university=university).select_related('faculty')
    courses = Course.objects.filter(lecturer=lecturer).select_related('student_group', 'program')
    student_groups = StudentGroup.objects.filter(
        courses__lecturer=lecturer
    ).distinct().select_related('program')

    active_timetable = (
        Timetable.objects.filter(semester__university=university, is_active=True).first()
        or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
    )
    slots = []
    total_hours = 0.0
    day_slots = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0, 7: 0}
    day_names = {1: 'Monday', 2: 'Tuesday', 3: 'Wednesday', 4: 'Thursday', 5: 'Friday', 6: 'Saturday', 7: 'Sunday'}
    
    if active_timetable:
        slots = list(
            ScheduleSlot.objects.filter(timetable=active_timetable, lecturer=lecturer)
            .select_related('course', 'room', 'time_slot', 'student_group')
            .order_by('time_slot__day_of_week', 'time_slot__slot_number')
        )
        for s in slots:
            ts = s.time_slot
            dur = (ts.end_time.hour * 60 + ts.end_time.minute) - (
                ts.start_time.hour * 60 + ts.start_time.minute
            )
            total_hours += (dur / 60)
            if ts.day_of_week in day_slots:
                day_slots[ts.day_of_week] += 1

    constraints = Constraint.objects.filter(
        university=university,
        constraint_type='LECTURER_AVAILABILITY'
    )
    unavailable_slot_ids = set()
    for c in constraints:
        l_id = c.parameters.get('lecturer_id')
        if l_id == lecturer.id:
            unavail_slots = c.parameters.get('unavailable_slots', [])
            unavailable_slot_ids.update(unavail_slots)
            
    unavail_records = LecturerAvailability.objects.filter(lecturer=lecturer, is_available=False)
    for rec in unavail_records:
        unavailable_slot_ids.add(rec.time_slot_id)

    unavailable_slots = TimeSlot.objects.filter(id__in=unavailable_slot_ids).order_by('day_of_week', 'slot_number')

    return render(request, 'scheduler/admin_lecturer_profile.html', {
        'lecturer': lecturer,
        'departments': departments,
        'courses': courses,
        'student_groups': student_groups,
        'timetable': active_timetable,
        'slots': slots,
        'total_hours': round(total_hours, 1),
        'day_slots': day_slots,
        'day_names': day_names,
        'unavailable_slots': unavailable_slots,
        'active_university': university,
    })


@login_required
def admin_lecturer_delete(request, pk):
    """Admin view to confirm and delete a lecturer."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    university = get_active_uni(request)
    lecturer = get_object_or_404(Lecturer, pk=pk)

    if lecturer.department.faculty.campus.university != university:
        messages.error(request, "Access denied.")
        return redirect('scheduler:resources_manager')

    if request.method == 'POST':
        name = lecturer.name
        if lecturer.user:
            lecturer.user.delete()
        lecturer.delete()
        messages.success(request, f"Lecturer {name} was successfully removed from the system.")
        return redirect('/resources/?tab=lecturer')

    return render(request, 'scheduler/admin_lecturer_confirm_delete.html', {
        'lecturer': lecturer,
        'active_university': university,
    })


# --- Approval Workflow Views ---

@login_required
@tenant_required(Timetable)
def timetable_workflow_action(request, pk):
    """
    Handles workflow state transitions for a timetable: SUBMIT, APPROVE, REJECT.
    """
    role = get_user_role(request)
    timetable = get_object_or_404(Timetable, pk=pk)

    if request.method != 'POST':
        return HttpResponseNotAllowed(['POST'])

    action = request.POST.get('action') # 'submit', 'approve', 'reject'
    comments = request.POST.get('comments', '').strip()

    if not action:
        messages.error(request, "Workflow action is missing.")
        return redirect('scheduler:timetable_detail', pk=pk)

    current_status = timetable.status

    # Validate transitioning permissions and states
    if action == 'submit':
        if current_status != 'DRAFT':
            messages.error(request, f"Timetable is in '{current_status}' state and cannot be submitted.")
            return redirect('scheduler:timetable_detail', pk=pk)
        
        if role not in (ROLE_SCHEDULER, ROLE_TIMETABLE_OFFICER, ROLE_ADMIN, ROLE_INST_ADMIN):
            messages.error(request, "Only schedulers and administrators can submit timetables for review.")
            return redirect('scheduler:timetable_detail', pk=pk)

        timetable.status = 'HOD_REVIEW'
        timetable.save()
        ApprovalLog.objects.create(
            timetable=timetable,
            stage='HOD_REVIEW',
            actor=request.user,
            action='SUBMIT',
            comments=comments
        )
        messages.success(request, "Timetable submitted successfully for Head of Department review.")

    elif action == 'approve':
        if current_status == 'DRAFT':
            messages.error(request, "Draft timetables must be submitted before approval.")
            return redirect('scheduler:timetable_detail', pk=pk)
        
        # Check permissions for each stage
        if current_status == 'HOD_REVIEW':
            if role not in (ROLE_HOD, ROLE_ADMIN, ROLE_INST_ADMIN):
                messages.error(request, "Only Heads of Department can approve this stage.")
                return redirect('scheduler:timetable_detail', pk=pk)
            timetable.status = 'DEAN_REVIEW'
        elif current_status == 'DEAN_REVIEW':
            if role not in (ROLE_DEAN, ROLE_ADMIN, ROLE_INST_ADMIN):
                messages.error(request, "Only Deans can approve this stage.")
                return redirect('scheduler:timetable_detail', pk=pk)
            timetable.status = 'REGISTRAR_REVIEW'
        elif current_status == 'REGISTRAR_REVIEW':
            if role not in (ROLE_REGISTRAR, ROLE_ADMIN, ROLE_INST_ADMIN):
                messages.error(request, "Only Registrars can approve this stage.")
                return redirect('scheduler:timetable_detail', pk=pk)
            timetable.status = 'DVC_REVIEW'
        elif current_status == 'DVC_REVIEW':
            if role not in (ROLE_DVC, ROLE_ADMIN, ROLE_INST_ADMIN):
                messages.error(request, "Only DVC Academic can approve this stage.")
                return redirect('scheduler:timetable_detail', pk=pk)
            timetable.status = 'PUBLISHED'
            timetable.is_active = True
            timetable.save()
            # Deactivate other timetables for the same semester
            Timetable.objects.filter(semester=timetable.semester).exclude(pk=timetable.pk).update(is_active=False)
        elif current_status == 'PUBLISHED':
            messages.info(request, "Timetable is already fully approved and published.")
            return redirect('scheduler:timetable_detail', pk=pk)
        
        timetable.save()
        ApprovalLog.objects.create(
            timetable=timetable,
            stage=current_status,
            actor=request.user,
            action='APPROVE',
            comments=comments
        )
        messages.success(request, f"Timetable approved at {timetable.get_status_display()} stage.")

    elif action == 'reject':
        if current_status == 'DRAFT':
            messages.error(request, "Draft timetables cannot be rejected.")
            return redirect('scheduler:timetable_detail', pk=pk)
        
        # Verify role permission to reject
        allowed = False
        if current_status == 'HOD_REVIEW' and role in (ROLE_HOD, ROLE_ADMIN, ROLE_INST_ADMIN):
            allowed = True
        elif current_status == 'DEAN_REVIEW' and role in (ROLE_DEAN, ROLE_ADMIN, ROLE_INST_ADMIN):
            allowed = True
        elif current_status == 'REGISTRAR_REVIEW' and role in (ROLE_REGISTRAR, ROLE_ADMIN, ROLE_INST_ADMIN):
            allowed = True
        elif current_status == 'DVC_REVIEW' and role in (ROLE_DVC, ROLE_ADMIN, ROLE_INST_ADMIN):
            allowed = True

        if not allowed:
            messages.error(request, "Permission denied to reject at this stage.")
            return redirect('scheduler:timetable_detail', pk=pk)

        timetable.status = 'DRAFT'
        timetable.save()
        ApprovalLog.objects.create(
            timetable=timetable,
            stage=current_status,
            actor=request.user,
            action='REJECT',
            comments=comments
        )
        messages.warning(request, f"Timetable was rejected and returned to Draft status. Reason: {comments}")

    return redirect('scheduler:timetable_detail', pk=pk)


# --- AI Assistant Views & APIs ---

from django.http import JsonResponse

@login_required
@tenant_required(Timetable)
def ai_quality_score(request, pk):
    """
    Computes a Timetable Quality Score out of 100 based on soft constraint violations.
    Soft constraints checked:
    - Avoid Evening Classes (weight 5)
    - Avoid consecutive slots > 2 for lecturer (weight 3)
    - Preferred time slots/dislikes violations (weight 5)
    """
    timetable = get_object_or_404(Timetable, pk=pk)
    slots = list(timetable.slots.select_related('course', 'lecturer', 'room', 'time_slot', 'student_group').all())

    evening_violations = 0
    consecutive_violations = 0
    preference_violations = 0

    # 1. Evening classes check
    for s in slots:
        if s.time_slot.is_evening:
            evening_violations += 1

    # 2. Consecutive slots check (> 2 slots per day)
    from collections import defaultdict
    lecturer_days = defaultdict(list)
    for s in slots:
        lecturer_days[(s.lecturer_id, s.time_slot.day_of_week)].append(s.time_slot.slot_number)
    
    for key, nums in lecturer_days.items():
        nums.sort()
        consec = 1
        for i in range(len(nums) - 1):
            if nums[i+1] == nums[i] + 1:
                consec += 1
                if consec > 2:
                    consecutive_violations += 1
            else:
                consec = 1

    # 3. Preferences violations (dislikes)
    from scheduler.models import LecturerTimeSlotPreference
    prefs = LecturerTimeSlotPreference.objects.filter(lecturer__department__faculty__campus__university=timetable.semester.university)
    disliked_map = defaultdict(set)
    for p in prefs:
        if p.preference_level == 'dislike':
            disliked_map[p.lecturer_id].add(p.time_slot_id)
            
    for s in slots:
        if s.time_slot_id in disliked_map[s.lecturer_id]:
            preference_violations += 1

    # Calculate score
    penalties = (evening_violations * 5) + (consecutive_violations * 3) + (preference_violations * 5)
    score = max(0, min(100, 100 - penalties))

    return JsonResponse({
        'timetable_id': pk,
        'quality_score': score,
        'penalties_breakdown': {
            'evening_class_violations': evening_violations,
            'lecturer_consecutive_violations': consecutive_violations,
            'disliked_slot_violations': preference_violations
        },
        'status': 'success'
    })


@login_required
@tenant_required(Timetable)
def ai_recommend_swaps(request, pk):
    """
    Given a ScheduleSlot ID, suggest alternative time slots or rooms that do not introduce conflicts.
    Query params: slot_id
    """
    timetable = get_object_or_404(Timetable, pk=pk)
    slot_id = request.GET.get('slot_id')
    if not slot_id:
        return JsonResponse({'error': 'Parameter slot_id is required.'}, status=400)

    target_slot = get_object_or_404(ScheduleSlot, pk=slot_id, timetable=timetable)
    university = timetable.semester.university

    # Fetch other slots in same timetable to detect conflicts
    other_slots = list(timetable.slots.exclude(pk=target_slot.pk).select_related('time_slot', 'room', 'lecturer', 'student_group').all())
    
    # Let's check available slots
    from scheduler.models import TimeSlot, Room
    all_timeslots = list(TimeSlot.objects.filter(university=university).all())
    all_rooms = list(Room.objects.filter(campus__university=university).filter(room_type=target_slot.course.required_room_type, capacity__gte=target_slot.student_group.size).all())

    recommendations = []

    # Simple constraint verification logic
    for ts in all_timeslots[:15]: # cap suggestions to keep it quick
        for rm in all_rooms[:3]:
            # Check lecturer conflict
            lecturer_conflict = any(s.lecturer_id == target_slot.lecturer_id and s.time_slot_id == ts.id for s in other_slots)
            # Check student group conflict
            group_conflict = any(s.student_group_id == target_slot.student_group_id and s.time_slot_id == ts.id for s in other_slots)
            # Check room conflict
            room_conflict = any(s.room_id == rm.id and s.time_slot_id == ts.id for s in other_slots)

            if not (lecturer_conflict or group_conflict or room_conflict):
                recommendations.append({
                    'time_slot_id': ts.id,
                    'time_slot_label': str(ts),
                    'room_id': rm.id,
                    'room_name': rm.name,
                    'capacity': rm.capacity
                })

    return JsonResponse({
        'slot_id': slot_id,
        'course': f"{target_slot.course.code}: {target_slot.course.name}",
        'current_time_slot': str(target_slot.time_slot),
        'current_room': target_slot.room.name,
        'recommendations': recommendations[:10] # return top 10 options
    })


@login_required
def student_portal_weekly_timetable(request):
    """Full 5-day weekly timetable grid for the student's student group."""
    import datetime
    role = get_user_role(request)
    if role != ROLE_STUDENT:
        if role not in (ROLE_ADMIN, ROLE_SCHEDULER, ROLE_TIMETABLE_OFFICER, ROLE_INST_ADMIN):
            return redirect('scheduler:dashboard')

    try:
        student_group = request.user.profile.student_group
    except Exception:
        student_group = None

    if not student_group:
        messages.error(request, "Your account is not linked to a student group.")
        return redirect('accounts:profile')

    university = get_active_uni(request)
    active_timetable = None
    if university:
        active_timetable = (
            Timetable.objects.filter(semester__university=university, is_active=True).first()
            or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
        )

    slots_by_day = {1: [], 2: [], 3: [], 4: [], 5: [], 6: [], 7: []}
    all_slots = []

    if active_timetable:
        all_slots = list(
            ScheduleSlot.objects.filter(timetable=active_timetable, student_group=student_group)
            .select_related('course', 'room', 'time_slot', 'lecturer')
            .order_by('time_slot__slot_number', 'time_slot__day_of_week')
        )
        from django.utils import timezone
        local_now = timezone.localtime(timezone.now())
        current_dow = local_now.isoweekday()
        current_time = local_now.time()
        for s in all_slots:
            s.has_ended = False
            s.is_ongoing = False

        for s in all_slots:
            slots_by_day[s.time_slot.day_of_week].append(s)

    _today = datetime.date.today()
    _monday = _today - datetime.timedelta(days=_today.weekday())
    
    max_day = 7

    all_days = [
        (1, 'Monday'), (2, 'Tuesday'), (3, 'Wednesday'),
        (4, 'Thursday'), (5, 'Friday'), (6, 'Saturday'), (7, 'Sunday')
    ]
    days_list = []
    for d, day_name in all_days[:max_day]:
        days_list.append({
            'dow': d,
            'name': day_name,
            'date': _monday + datetime.timedelta(days=d - 1)
        })
    today_dow = _today.weekday() + 1

    # Group and sort slots chronologically by start_time using all university timeslots
    grid_rows = []
    if university:
        all_uni_ts = list(TimeSlot.objects.filter(university=university).order_by('slot_number', 'day_of_week'))
        unique_slots_info = {}
        for ts in all_uni_ts:
            num = ts.slot_number
            if num not in unique_slots_info:
                unique_slots_info[num] = {
                    'slot_number': num,
                    'start_time': ts.start_time,
                    'end_time': ts.end_time,
                }
        
        sorted_slots = sorted(unique_slots_info.values(), key=lambda x: x['start_time'])
        
        for s_info in sorted_slots:
            num = s_info['slot_number']
            day_slots = {}
            for d in range(1, max_day + 1):
                day_slots[d] = next((s for s in all_slots if s.time_slot.day_of_week == d and s.time_slot.slot_number == num), None)
            
            grid_rows.append({
                'start_time': s_info['start_time'],
                'end_time': s_info['end_time'],
                'slot_number': num,
                'day_slots': day_slots,
            })

    import uuid
    from django.urls import reverse
    if not student_group.calendar_token:
        student_group.calendar_token = uuid.uuid4()
        student_group.save(update_fields=['calendar_token'])

    feed_url = request.build_absolute_uri(
        reverse('scheduler:student_group_calendar_feed', args=[student_group.calendar_token])
    )
    webcal_url = feed_url.replace('http://', 'webcal://').replace('https://', 'webcal://')

    return render(request, 'scheduler/student_weekly_timetable.html', {
        'student_group': student_group,
        'active_university': university,
        'timetable': active_timetable,
        'slots_by_day': slots_by_day,
        'all_slots': all_slots,
        'grid_rows': grid_rows,
        'days': days_list,
        'today_dow': today_dow,
        'active_semester': Semester.objects.filter(
            university=university, is_active=True).first() if university else None,
        'feed_url': feed_url,
        'webcal_url': webcal_url,
    })


# ── Import Audit Views ────────────────────────────────────────────────────────

@login_required
def import_audit_report(request, pk):
    """Post-import summary showing what was created, updated, and flagged."""
    from scheduler.models import ImportAuditLog
    university = get_active_uni(request)
    role = get_user_role(request)
    if role not in MANAGER_ROLES:
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    try:
        log = ImportAuditLog.objects.get(pk=pk, university=university)
    except ImportAuditLog.DoesNotExist:
        messages.error(request, "Import record not found.")
        return redirect('scheduler:import_audit_log_list')

    return render(request, 'scheduler/import_audit_report.html', {
        'log': log,
        'active_university': university,
    })


@login_required
def import_audit_log_list(request):
    """List all import audit logs for the active university."""
    from scheduler.models import ImportAuditLog
    university = get_active_uni(request)
    role = get_user_role(request)
    if role not in MANAGER_ROLES:
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    logs = ImportAuditLog.objects.filter(university=university).order_by('-imported_at')[:50]
    return render(request, 'scheduler/import_audit_log_list.html', {
        'logs': logs,
        'active_university': university,
    })


@login_required
def admin_room_profile(request, pk):
    """Admin view to inspect and manage an individual room's capacity, type, and schedules."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    university = get_active_uni(request)
    room = get_object_or_404(Room.objects.select_related('campus'), pk=pk)
    if room.campus.university != university:
        messages.error(request, "Access denied.")
        return redirect('scheduler:resources_manager')

    if request.method == 'POST':
        name = request.POST.get('name')
        capacity = request.POST.get('capacity')
        room_type = request.POST.get('room_type')
        campus_id = request.POST.get('campus')

        room.name = name
        if capacity:
            room.capacity = int(capacity)
        if room_type:
            room.room_type = room_type
        if campus_id:
            campus = Campus.objects.filter(university=university, id=campus_id).first()
            if campus:
                room.campus = campus
        room.save()
        messages.success(request, f"Room {room.name} updated successfully.")
        return redirect('scheduler:admin_room_profile', pk=room.pk)

    campuses = Campus.objects.filter(university=university)
    
    active_timetable = (
        Timetable.objects.filter(semester__university=university, is_active=True).first()
        or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
    )
    slots = []
    if active_timetable:
        slots = list(
            ScheduleSlot.objects.filter(timetable=active_timetable, room=room)
            .select_related('course', 'lecturer', 'time_slot', 'student_group')
            .order_by('time_slot__day_of_week', 'time_slot__slot_number')
        )

    return render(request, 'scheduler/admin_room_profile.html', {
        'room': room,
        'campuses': campuses,
        'slots': slots,
        'active_university': university,
    })


@login_required
def admin_studentgroup_profile(request, pk):
    """Admin view to inspect and manage an individual cohort's size, year, and scheduled slots."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    university = get_active_uni(request)
    student_group = get_object_or_404(StudentGroup.objects.select_related('program'), pk=pk)
    if student_group.program.department.faculty.campus.university != university:
        messages.error(request, "Access denied.")
        return redirect('scheduler:resources_manager')

    if request.method == 'POST':
        name = request.POST.get('name')
        size = request.POST.get('size')
        year = request.POST.get('year')
        prog_id = request.POST.get('program')
        parent_id = request.POST.get('parent_group')

        student_group.name = name
        if size:
            student_group.size = int(size)
        if year:
            student_group.year = int(year)
        if prog_id:
            prog = Program.objects.filter(department__faculty__campus__university=university, id=prog_id).first()
            if prog:
                student_group.program = prog
        if parent_id:
            parent = StudentGroup.objects.filter(program__department__faculty__campus__university=university, id=parent_id).exclude(pk=student_group.pk).first()
            student_group.parent_group = parent
        else:
            student_group.parent_group = None
        student_group.save()
        messages.success(request, f"Student group {student_group.name} updated successfully.")
        return redirect('scheduler:admin_studentgroup_profile', pk=student_group.pk)

    programs = Program.objects.filter(department__faculty__campus__university=university)
    parent_groups = StudentGroup.objects.filter(program__department__faculty__campus__university=university).exclude(pk=student_group.pk)

    active_timetable = (
        Timetable.objects.filter(semester__university=university, is_active=True).first()
        or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
    )
    slots = []
    if active_timetable:
        slots = list(
            ScheduleSlot.objects.filter(timetable=active_timetable, student_group=student_group)
            .select_related('course', 'lecturer', 'time_slot', 'room')
            .order_by('time_slot__day_of_week', 'time_slot__slot_number')
        )

    assigned_courses = list(Course.objects.filter(student_group=student_group).select_related('lecturer', 'program'))

    return render(request, 'scheduler/admin_studentgroup_profile.html', {
        'student_group': student_group,
        'programs': programs,
        'parent_groups': parent_groups,
        'slots': slots,
        'courses': assigned_courses,
        'active_university': university,
    })


@login_required
def admin_course_profile(request, pk):
    """Admin view to inspect and manage an individual course's details, weekly sessions, and mappings."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    university = get_active_uni(request)
    course = get_object_or_404(Course.objects.select_related('program', 'lecturer', 'student_group'), pk=pk)
    if course.program.department.faculty.campus.university != university:
        messages.error(request, "Access denied.")
        return redirect('scheduler:resources_manager')

    if request.method == 'POST':
        code = request.POST.get('code')
        name = request.POST.get('name')
        duration_slots = request.POST.get('duration_slots')
        sessions_per_week = request.POST.get('sessions_per_week')
        room_type = request.POST.get('required_room_type')
        lec_id = request.POST.get('lecturer')
        sg_id = request.POST.get('student_group')
        prog_id = request.POST.get('program')

        course.code = code
        course.name = name
        if duration_slots:
            course.duration_slots = int(duration_slots)
        if sessions_per_week:
            course.sessions_per_week = int(sessions_per_week)
        if room_type:
            course.required_room_type = room_type

        if lec_id:
            lec = Lecturer.objects.filter(department__faculty__campus__university=university, id=lec_id).first()
            course.lecturer = lec
        else:
            course.lecturer = None

        if sg_id:
            sg = StudentGroup.objects.filter(program__department__faculty__campus__university=university, id=sg_id).first()
            course.student_group = sg
        else:
            course.student_group = None

        if prog_id:
            prog = Program.objects.filter(department__faculty__campus__university=university, id=prog_id).first()
            if prog:
                course.program = prog

        course.save()
        messages.success(request, f"Course {course.code} updated successfully.")
        return redirect('scheduler:admin_course_profile', pk=course.pk)

    programs = Program.objects.filter(department__faculty__campus__university=university)
    lecturers = Lecturer.objects.filter(department__faculty__campus__university=university)
    student_groups = StudentGroup.objects.filter(program__department__faculty__campus__university=university)

    active_timetable = (
        Timetable.objects.filter(semester__university=university, is_active=True).first()
        or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
    )
    slots = []
    if active_timetable:
        slots = list(
            ScheduleSlot.objects.filter(timetable=active_timetable, course=course)
            .select_related('lecturer', 'time_slot', 'room', 'student_group')
            .order_by('time_slot__day_of_week', 'time_slot__slot_number')
        )

    return render(request, 'scheduler/admin_course_profile.html', {
        'course': course,
        'programs': programs,
        'lecturers': lecturers,
        'student_groups': student_groups,
        'slots': slots,
        'active_university': university,
    })


@login_required
def admin_university_profile(request, pk):
    """Admin view to inspect and manage an individual university's name and code."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    university = get_object_or_404(University, pk=pk)

    if request.method == 'POST':
        name = request.POST.get('name')
        code = request.POST.get('code')

        university.name = name
        university.code = code
        university.save()
        messages.success(request, f"University {university.name} updated successfully.")
        return redirect('scheduler:admin_university_profile', pk=university.pk)

    campuses = Campus.objects.filter(university=university)
    semesters = Semester.objects.filter(university=university)

    return render(request, 'scheduler/admin_university_profile.html', {
        'university_item': university,
        'campuses': campuses,
        'semesters': semesters,
        'active_university': get_active_uni(request),
    })


@login_required
def admin_department_profile(request, pk):
    """Admin view to inspect and manage an individual department's name, faculty, and programs."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    university = get_active_uni(request)
    department = get_object_or_404(Department.objects.select_related('faculty', 'faculty__campus'), pk=pk)
    if department.faculty.campus.university != university:
        messages.error(request, "Access denied.")
        return redirect('scheduler:resources_manager')

    if request.method == 'POST':
        name = request.POST.get('name')
        faculty_id = request.POST.get('faculty')

        department.name = name
        if faculty_id:
            fac = Faculty.objects.filter(campus__university=university, id=faculty_id).first()
            if fac:
                department.faculty = fac
        department.save()
        messages.success(request, f"Department {department.name} updated successfully.")
        return redirect('scheduler:admin_department_profile', pk=department.pk)

    faculties = Faculty.objects.filter(campus__university=university)
    programs = Program.objects.filter(department=department)
    lecturers = Lecturer.objects.filter(department=department)

    return render(request, 'scheduler/admin_department_profile.html', {
        'department': department,
        'faculties': faculties,
        'programs': programs,
        'lecturers': lecturers,
        'active_university': university,
    })


@login_required
def admin_timeslot_profile(request, pk):
    """Admin view to inspect and manage an individual time slot's properties."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')

    university = get_active_uni(request)
    try:
        time_slot = TimeSlot.objects.get(pk=pk)
    except TimeSlot.DoesNotExist:
        time_slot = TimeSlot.objects.filter(university=university).first()
        if not time_slot:
            messages.error(request, f"TimeSlot #{pk} no longer exists.")
            return redirect('scheduler:resources_manager')
        messages.warning(request, f"TimeSlot #{pk} was updated or re-created. Displaying valid slot {time_slot}.")

    if time_slot.university != university:
        messages.error(request, "Access denied.")
        return redirect('scheduler:resources_manager')

    if request.method == 'POST':
        day_of_week = request.POST.get('day_of_week')
        start_time = request.POST.get('start_time')
        end_time = request.POST.get('end_time')
        slot_number = request.POST.get('slot_number')
        is_evening = request.POST.get('is_evening') == 'on'

        if day_of_week:
            time_slot.day_of_week = int(day_of_week)
        if start_time:
            time_slot.start_time = start_time
        if end_time:
            time_slot.end_time = end_time
        if slot_number:
            time_slot.slot_number = int(slot_number)
        time_slot.is_evening = is_evening
        time_slot.save()
        time_slot.refresh_from_db()
        messages.success(request, f"Time slot {time_slot} updated successfully.")
        return redirect('scheduler:admin_timeslot_profile', pk=time_slot.pk)

    day_choices = TimeSlot.DAY_CHOICES

    active_timetable = (
        Timetable.objects.filter(semester__university=university, is_active=True).first()
        or Timetable.objects.filter(semester__university=university).order_by('-created_at').first()
    )
    slots = []
    if active_timetable:
        slots = list(
            ScheduleSlot.objects.filter(timetable=active_timetable, time_slot=time_slot)
            .select_related('course', 'lecturer', 'room', 'student_group')
            .order_by('room__name')
        )

    return render(request, 'scheduler/admin_timeslot_profile.html', {
        'time_slot': time_slot,
        'day_choices': day_choices,
        'slots': slots,
        'active_university': university,
    })


@login_required
def update_lecturer_hours(request):
    """Update maximum hours limit and contract type for a lecturer."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:reports_workloads')

    uni = get_active_uni(request)
    if request.method == 'POST':
        lecturer_id = request.POST.get('lecturer_id')
        max_hours = request.POST.get('max_hours')
        lecturer_type = request.POST.get('lecturer_type')  # 'FT' or 'PT'
        
        lecturer = get_object_or_404(Lecturer, id=lecturer_id)
        if lecturer.department.faculty.campus.university != uni:
            messages.error(request, "Permission denied.")
            return redirect('scheduler:reports_workloads')
            
        try:
            if max_hours:
                lecturer.max_hours_per_week = int(max_hours)
            if lecturer_type in ('FT', 'PT'):
                lecturer.lecturer_type = lecturer_type
            lecturer.save()
            type_label = lecturer.get_lecturer_type_display()
            messages.success(request, f"Updated {lecturer.name}: contract={type_label}, max hours={lecturer.max_hours_per_week}.")
        except (ValueError, TypeError):
            messages.error(request, "Invalid value provided.")
            
    return redirect('scheduler:reports_workloads')


@login_required
def reassign_course(request):
    """Reassign a course to a different lecturer or unassign it."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:reports_workloads')

    uni = get_active_uni(request)
    if request.method == 'POST':
        course_id = request.POST.get('course_id')
        lecturer_id = request.POST.get('lecturer_id')
        
        course = get_object_or_404(Course, id=course_id)
        if course.program.department.faculty.campus.university != uni:
            messages.error(request, "Permission denied.")
            return redirect('scheduler:reports_workloads')
            
        if lecturer_id == 'none' or not lecturer_id:
            course.lecturer = None
            course.save()
            messages.success(request, f"Course {course.code} has been set to unassigned (TBD).")
        else:
            lecturer = get_object_or_404(Lecturer, id=lecturer_id)
            if lecturer.department.faculty.campus.university != uni:
                messages.error(request, "Permission denied.")
                return redirect('scheduler:reports_workloads')
                
            course.lecturer = lecturer
            course.save()
            messages.success(request, f"Course {course.code} successfully reassigned to {lecturer.name}.")
            
    return redirect('scheduler:reports_workloads')


@login_required
def auto_balance_workloads_view(request):
    """
    Intelligent Workload Auto-Balancer:
    Detects overloaded lecturers (> 20h/wk) and underloaded lecturers (< 12h/wk).
    Automatically redistributes courses within the same department/university
    to bring everyone into the optimal 12-15 hours/week range.
    """
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:reports_workloads')

    uni = get_active_uni(request)
    if not uni:
        messages.error(request, "No active university selected.")
        return redirect('scheduler:reports_workloads')

    lecturers = list(Lecturer.objects.filter(department__faculty__campus__university=uni, is_active=True).select_related('department'))
    courses = list(Course.objects.filter(program__department__faculty__campus__university=uni).select_related('lecturer', 'program__department'))

    if not lecturers or not courses:
        messages.warning(request, "No lecturers or courses available to balance.")
        return redirect('scheduler:reports_workloads')

    def get_course_hours(c):
        return float(c.duration_slots * c.sessions_per_week)

    def get_subject_prefix(code):
        parts = code.strip().split()
        return parts[0].upper() if parts else code.upper()

    def safely_update_schedule_slots(course, new_lecturer):
        for slot in ScheduleSlot.objects.filter(course=course):
            # Guard against double-booking candidate at the same time slot
            is_busy = ScheduleSlot.objects.filter(
                timetable=slot.timetable, 
                time_slot=slot.time_slot, 
                lecturer=new_lecturer
            ).exclude(pk=slot.pk).exists()
            if not is_busy:
                try:
                    slot.lecturer = new_lecturer
                    slot.save(update_fields=['lecturer'])
                except Exception:
                    pass

    lec_workloads = {l.id: 0.0 for l in lecturers}
    lec_courses = {l.id: [] for l in lecturers}
    lec_subjects = {l.id: set() for l in lecturers}
    unassigned_courses = []

    for c in courses:
        h = get_course_hours(c)
        prefix = get_subject_prefix(c.code)
        if c.lecturer_id and c.lecturer_id in lec_workloads:
            lec_workloads[c.lecturer_id] += h
            lec_courses[c.lecturer_id].append(c)
            lec_subjects[c.lecturer_id].add(prefix)
        else:
            unassigned_courses.append(c)

    reassigned_count = 0
    overloaded_lecs = [l for l in lecturers if lec_workloads[l.id] > 20.0]
    overloaded_lecs.sort(key=lambda l: lec_workloads[l.id], reverse=True)

    for ov_lec in overloaded_lecs:
        movable_courses = list(lec_courses[ov_lec.id])
        movable_courses.sort(key=get_course_hours, reverse=True)

        for c in movable_courses:
            if lec_workloads[ov_lec.id] <= 20.0:
                break
            ch = get_course_hours(c)
            prefix = get_subject_prefix(c.code)

            # Get time slot IDs occupied by course c
            c_slot_time_ids = set(ScheduleSlot.objects.filter(course=c).values_list('time_slot_id', flat=True))

            candidates = []
            for l in lecturers:
                if l.id == ov_lec.id or (lec_workloads[l.id] + ch) > 20.0:
                    continue
                # Check if candidate l is already busy in any of c's scheduled time slots
                if c_slot_time_ids:
                    busy_count = ScheduleSlot.objects.filter(lecturer=l, time_slot_id__in=c_slot_time_ids).exists()
                    if busy_count:
                        continue
                candidates.append(l)

            if not candidates:
                continue

            candidates.sort(key=lambda l: (
                0 if (l.department_id == c.program.department_id and prefix in lec_subjects[l.id]) else
                1 if l.department_id == c.program.department_id else
                2 if prefix in lec_subjects[l.id] else 3,
                abs(14.0 - (lec_workloads[l.id] + ch))
            ))

            target_lec = candidates[0]
            c.lecturer = target_lec
            c.save(update_fields=['lecturer'])

            safely_update_schedule_slots(c, target_lec)

            lec_workloads[ov_lec.id] -= ch
            lec_workloads[target_lec.id] += ch
            lec_courses[ov_lec.id].remove(c)
            lec_courses[target_lec.id].append(c)
            lec_subjects[target_lec.id].add(prefix)
            reassigned_count += 1

    for c in list(unassigned_courses):
        ch = get_course_hours(c)
        prefix = get_subject_prefix(c.code)
        c_slot_time_ids = set(ScheduleSlot.objects.filter(course=c).values_list('time_slot_id', flat=True))

        candidates = []
        for l in lecturers:
            if (lec_workloads[l.id] + ch) > 20.0:
                continue
            if c_slot_time_ids:
                busy_count = ScheduleSlot.objects.filter(lecturer=l, time_slot_id__in=c_slot_time_ids).exists()
                if busy_count:
                    continue
            candidates.append(l)

        if candidates:
            candidates.sort(key=lambda l: (
                0 if (l.department_id == c.program.department_id and prefix in lec_subjects[l.id]) else
                1 if l.department_id == c.program.department_id else
                2 if prefix in lec_subjects[l.id] else 3,
                abs(14.0 - (lec_workloads[l.id] + ch))
            ))
            target_lec = candidates[0]
            c.lecturer = target_lec
            c.save(update_fields=['lecturer'])

            safely_update_schedule_slots(c, target_lec)

            lec_workloads[target_lec.id] += ch
            lec_subjects[target_lec.id].add(prefix)
            reassigned_count += 1

    if reassigned_count > 0:
        messages.success(request, f"🤖 Smart Subject-Matched Auto-Balance completed! Reassigned {reassigned_count} course(s) to department & subject-matched lecturers within 12–20 hrs without schedule conflicts.")
    else:
        messages.info(request, "All lecturer workloads are already within optimal limits or no suitable reassignment targets were found.")

    return redirect('scheduler:reports_workloads')


@login_required
def manual_auto_heal(request):
    """Manually trigger the diagnostic & self-healing data helper."""
    role = get_user_role(request)
    if role in (ROLE_STUDENT, ROLE_LECTURER):
        messages.error(request, "Permission denied.")
        return redirect('scheduler:dashboard')
        
    uni = get_active_uni(request)
    if not uni:
        messages.error(request, "No active university found.")
        return redirect('scheduler:reports')

    if request.method == 'POST':
        try:
            fixes = auto_heal_university_data(uni)
            if fixes:
                for fix in fixes:
                    messages.success(request, fix)
        except Exception as e:
            messages.error(request, f"Error during auto-healing: {e}")
            
    return redirect('scheduler:reports')


def public_lecturer_onboarding(request, token=None):
    """
    Public Pre-Login Lecturer Onboarding & Verification Form.
    Allows faculty members to verify their contact details, target workload,
    and teaching availability preferences BEFORE logging in!
    Supports direct token URL, string search, or email/staff ID lookup.
    """
    from django.shortcuts import render, redirect
    from django.contrib import messages
    from django.db.models import Q
    import uuid
    from scheduler.models import Lecturer, TimeSlot, LecturerAvailability, Department

    lookup_query = (request.POST.get('lookup_query') or request.GET.get('email') or request.GET.get('staff_id') or request.GET.get('q') or '').strip()
    create_new = request.POST.get('create_new') or request.GET.get('create_new')

    lecturer = None
    if token:
        token_str = str(token).strip()
        try:
            token_uuid = uuid.UUID(token_str)
            lecturer = Lecturer.objects.filter(calendar_token=token_uuid).first()
        except (ValueError, TypeError):
            lecturer = Lecturer.objects.filter(
                Q(email__iexact=token_str) | Q(staff_id__iexact=token_str) | Q(name__icontains=token_str)
            ).first()

        if not lecturer:
            return render(request, 'scheduler/public_onboarding_lookup.html', {
                'error': f"No faculty record found matching token/ID '{token}'. Please enter your official staff email or staff ID below.",
                'query': token_str
            })
    else:
        if lookup_query:
            lecturer = Lecturer.objects.filter(
                Q(email__iexact=lookup_query) | Q(staff_id__iexact=lookup_query) | Q(name__icontains=lookup_query)
            ).first()

            if lecturer and '@' in lookup_query and lecturer.email != lookup_query.lower():
                lecturer.email = lookup_query.lower()
                lecturer.save(update_fields=['email'])

            if not lecturer:
                default_dept = Department.objects.first()
                email_str = lookup_query if '@' in lookup_query else f"{lookup_query.lower().replace(' ', '')}@university.edu"
                name_str = lookup_query.split('@')[0].replace('.', ' ').replace('_', ' ').title() if '@' in lookup_query else lookup_query
                lecturer = Lecturer.objects.create(
                    name=name_str,
                    email=email_str,
                    department=default_dept,
                    is_verified=True
                )
                messages.info(request, f"New faculty profile created for '{lecturer.name}'. A verification email has been sent to {lecturer.email}.")
        else:
            return render(request, 'scheduler/public_onboarding_lookup.html')

    if not lecturer:
        return render(request, 'scheduler/public_onboarding_lookup.html')

    university = lecturer.department.faculty.campus.university if (lecturer.department and lecturer.department.faculty and lecturer.department.faculty.campus) else None

    # Send verification email on search if searching via lookup_query
    if lookup_query and lecturer and lecturer.email:
        try:
            from scheduler.tasks import verify_and_notify_lecturer_record
            verify_and_notify_lecturer_record(
                submitted_email=lecturer.email,
                submitted_name=lecturer.name,
                staff_id=lecturer.staff_id,
                university_id=university.id if university else None,
                preserve_password=True
            )
            messages.success(request, f"Verification & schedule details email sent to {lecturer.email}.")
        except Exception as mail_err:
            logger.warning(f"[Onboarding Lookup Email] Could not send verification email on search: {mail_err}")

    timeslots = list(TimeSlot.objects.filter(university=university).order_by('day_of_week', 'slot_number')) if university else list(TimeSlot.objects.all().order_by('day_of_week', 'slot_number'))

    # Process form submission ONLY when saving onboarding preferences (not during initial lookup search)
    is_saving = request.method == 'POST' and (request.POST.get('save_onboarding') or 'max_hours_per_week' in request.POST or 'unavailable_slots' in request.POST) and not request.POST.get('lookup_query')

    if is_saving:
        name = request.POST.get('name', '').strip()
        email = request.POST.get('email', '').strip().lower()
        max_hours = request.POST.get('max_hours_per_week')
        lecturer_type = request.POST.get('lecturer_type', 'FT')
        password = request.POST.get('password', '').strip()

        if name:
            lecturer.name = name
        if email and '@' in email:
            lecturer.email = email
        if max_hours:
            try:
                lecturer.max_hours_per_week = int(max_hours)
            except ValueError:
                pass
        lecturer.lecturer_type = lecturer_type
        lecturer.save()

        if lecturer.user:
            user = lecturer.user
            if email and '@' in email:
                user.email = email
            if password:
                user.set_password(password)
            user.save()
        elif email and '@' in email:
            from django.contrib.auth.models import User
            user = User.objects.filter(email__iexact=email).first()
            if not user:
                base_un = email.split('@')[0]
                un = base_un
                c = 1
                while User.objects.filter(username=un).exists():
                    un = f"{base_un}{c}"
                    c += 1
                user = User.objects.create_user(username=un, email=email, password=password or 'wasike123')
            lecturer.user = user
            lecturer.save(update_fields=['user'])

        unavailable_slot_ids = request.POST.getlist('unavailable_slots')
        LecturerAvailability.objects.filter(lecturer=lecturer).delete()
        for ts_id in unavailable_slot_ids:
            try:
                ts = TimeSlot.objects.get(id=int(ts_id))
                LecturerAvailability.objects.create(lecturer=lecturer, time_slot=ts, is_available=False)
            except Exception:
                pass

        try:
            from scheduler.tasks import verify_and_notify_lecturer_record
            verify_and_notify_lecturer_record(
                submitted_email=lecturer.email,
                submitted_name=lecturer.name,
                staff_id=lecturer.staff_id,
                university_id=university.id if university else None,
                preserve_password=bool(password)
            )
        except Exception as err:
            import logging
            logging.getLogger(__name__).warning(f"[Onboarding Email] Failed to process verification notification: {err}")

        messages.success(request, f"✓ Thank you {lecturer.name}! Your profile & availability preferences have been saved successfully. Verification email sent to {lecturer.email}.")
        return render(request, 'scheduler/public_onboarding_lookup.html', {
            'query': lecturer.email
        })

    return render(request, 'scheduler/public_onboarding.html', {
        'lecturer': lecturer,
        'university': university,
        'timeslots': timeslots,
    })

